from django.forms import BooleanField
from collections import defaultdict
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_GET
from docx.shared import Pt, RGBColor

from .utils import get_subject_id
from django.db.models import Q, Case, When, BooleanField, Value, F, Min
from django.db.models import Case, When, IntegerField, Count, Prefetch
from docx import Document
from .models import Teacher, Lesson, Replacement, ClassSchedule, ActivityLog, SpecialReplacement, DocxImportTask
import json
import logging
import re
import io
import base64
from django.shortcuts import render
from datetime import datetime, timedelta
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.conf import settings
from django.core.cache import cache

from django.db import transaction, IntegrityError, DatabaseError, OperationalError, ProgrammingError
from django.views.decorators.http import require_POST
from django_ratelimit.decorators import ratelimit

import calendar as _pycalendar
from types import SimpleNamespace

from .audit import log_activity
from accounts.icon_service import get_icon_for_display_name

from .scheduling import (
    day_short_from_date,
    extract_grade,
    overlaps,
    infer_shift_by_grade,
    SECOND_SHIFT_GRADES,
    FIRST_SHIFT_GRADES_SECONDARY,
)
from .services.permissions_parsing import _use_gsheets_backend


def _is_vacancy_teacher_name(full_name: str | None) -> bool:
    n = (full_name or "").strip().lower()
    return ("вакан" in n) or ("vakans" in n)


def _name_matches_term_ci(full_name: str | None, term: str | None) -> bool:
    """Case-insensitive substring match that works reliably for Cyrillic."""
    t = (term or "").strip()
    if not t:
        return True
    return t.casefold() in (full_name or "").casefold()


def _is_guest_user(user) -> bool:
    return bool(getattr(user, "is_guest", False))


def _is_teacher_user(user) -> bool:
    return bool(getattr(user, "is_teacher", False))


def _resolve_teacher_for_user(user) -> Teacher | None:
    full_name = str(getattr(user, "full_name", "") or "").strip()
    if not full_name:
        return None
    return Teacher.objects.filter(full_name__iexact=full_name).first()


def _deny_guest_write_json(request):
    if _is_guest_user(request.user):
        return JsonResponse({"error": "Гостевой доступ: только просмотр"}, status=403)
    return None


def _can_calendar_read(user) -> bool:
    return bool(
        user.is_superuser
        or _is_guest_user(user)
        or _is_teacher_user(user)
        or getattr(user, "can_calendar", False)
        or getattr(user, "can_calls", False)
    )


def _can_calendar_write(user) -> bool:
    return bool(user.is_superuser or getattr(user, "can_calendar", False))


def _api_error(message: str, *, status: int = 400, code: str = "bad_request", **extra):
    payload = {
        "error": {
            "code": code,
            "message": message,
        }
    }
    if extra:
        payload.update(extra)
    return JsonResponse(payload, status=status)


def as_bool(v) -> bool:
    s = str(v or "").strip().lower()
    return s in {"1", "true", "t", "yes", "y"}


def as_int(v, default=None):
    if v is None:
        return default
    s = str(v).strip()
    if not s:
        return default
    try:
        return int(float(s))
    except (TypeError, ValueError):
        return default


def _gs_shift_boundary(selected_date):
    schedule_rows = _gs_store().get_table_dicts("replacements_class_schedule")
    shift1_max_end = None
    shift2_min_start = None
    for row in schedule_rows:
        sh = as_int(row.get("shift"))
        st = _gs_parse_time(row.get("start_time"))
        en = _gs_parse_time(row.get("end_time"))
        if sh == 1 and en:
            shift1_max_end = en if shift1_max_end is None or en > shift1_max_end else shift1_max_end
        if sh == 2 and st:
            shift2_min_start = st if shift2_min_start is None or st < shift2_min_start else shift2_min_start

    boundary = shift2_min_start
    if shift1_max_end and shift2_min_start:
        try:
            dt1 = datetime.combine(selected_date, shift1_max_end)
            dt2 = datetime.combine(selected_date, shift2_min_start)
            boundary = (dt1 + (dt2 - dt1) / 2).time() if dt2 > dt1 else shift2_min_start
        except Exception:
            boundary = shift2_min_start
    return shift1_max_end, shift2_min_start, boundary


def _gs_save_lessons_from_parsed(parsed_data: list[dict]) -> tuple[int, int]:
    """Apply parsed schedule as active schedule in Google Sheets backend.

    Returns:
      (deactivated_count, saved_count)
    """
    store = _gs_store()
    teacher_rows = store.get_table_dicts("replacements_teacher")
    subject_rows = store.get_table_dicts("replacements_subject")
    lesson_rows = store.get_table_dicts("replacements_lesson")

                                                                                       
    deactivated = 0
    for l in lesson_rows:
        if as_bool(l.get("is_active")):
            l["is_active"] = 0
            deactivated += 1

                                                              
    subject_by_name = {
        str(s.get("name") or "").strip(): as_int(s.get("id_subject"))
        for s in subject_rows
        if str(s.get("name") or "").strip() and as_int(s.get("id_subject")) is not None
    }
    teacher_by_name = {
        str(t.get("full_name") or "").strip(): as_int(t.get("teacher_id"))
        for t in teacher_rows
        if str(t.get("full_name") or "").strip() and as_int(t.get("teacher_id")) is not None
    }

                                                                   
    subjects_by_teacher: dict[str, set[str]] = {}
    for item in parsed_data:
        t_name = str(item.get("teacher") or "").strip()
        s_name = str(item.get("subject") or "").strip()
        if not t_name or not s_name:
            continue
        subjects_by_teacher.setdefault(t_name, set()).add(s_name)

    for subj_name in sorted({str(i.get("subject") or "").strip() for i in parsed_data if str(i.get("subject") or "").strip()}):
        if subj_name not in subject_by_name:
            new_sid = _gs_next_id(subject_rows, "id_subject")
            subject_rows.append({"id_subject": new_sid, "name": subj_name})
            subject_by_name[subj_name] = new_sid

    for teacher_name in sorted({str(i.get("teacher") or "").strip() for i in parsed_data if str(i.get("teacher") or "").strip()}):
        if teacher_name not in teacher_by_name:
            new_tid = _gs_next_id(teacher_rows, "teacher_id")
            teacher_rows.append({
                "teacher_id": new_tid,
                "full_name": teacher_name,
                "specialization": "",
                "hours_per_week": 0,
            })
            teacher_by_name[teacher_name] = new_tid

                                                 
    for t in teacher_rows:
        t_name = str(t.get("full_name") or "").strip()
        parsed_subjects = subjects_by_teacher.get(t_name, set())
        if not parsed_subjects:
            continue
        spec_ids = sorted({
            subject_by_name.get(s_name)
            for s_name in parsed_subjects
            if subject_by_name.get(s_name) is not None
        })
        t["specialization"] = ",".join(str(sid) for sid in spec_ids) if spec_ids else ""

    saved_count = 0
    for item in parsed_data:
        teacher_name = str(item.get("teacher") or "").strip()
        subject_name = str(item.get("subject") or "").strip()
        class_group = str(item.get("class_group") or "").strip()
        lesson_number = as_int(item.get("lesson_number"))
        day_of_week = str(item.get("day_of_week") or "").strip()
        shift = as_int(item.get("shift"))
        start_time = _gs_time_str(item.get("start_time"))
        end_time = _gs_time_str(item.get("end_time"))
        classroom = str(item.get("room") or "---").strip() or "---"

        teacher_id = teacher_by_name.get(teacher_name)
        subject_id = subject_by_name.get(subject_name)
        if teacher_id is None or subject_id is None or lesson_number is None or shift is None:
            continue

        existing_idx = None
        for idx, l in enumerate(lesson_rows):
            if (
                as_int(l.get("teacher_id")) == teacher_id
                and as_int(l.get("subject_id_subject")) == subject_id
                and str(l.get("class_group") or "") == class_group
                and as_int(l.get("lesson_number")) == lesson_number
                and str(l.get("day_of_week") or "") == day_of_week
                and _gs_time_str(l.get("start_time")) == start_time
                and _gs_time_str(l.get("end_time")) == end_time
                and as_int(l.get("shift")) == shift
                and str(l.get("classroom") or "---") == classroom
            ):
                existing_idx = idx
                break

        if existing_idx is not None:
            if not as_bool(lesson_rows[existing_idx].get("is_active")):
                lesson_rows[existing_idx]["is_active"] = 1
            saved_count += 1
        else:
            lesson_rows.append({
                "lesson_id": _gs_next_id(lesson_rows, "lesson_id"),
                "lesson_number": lesson_number,
                "class_group": class_group,
                "classroom": classroom,
                "start_time": start_time,
                "end_time": end_time,
                "shift": shift,
                "day_of_week": day_of_week,
                "subject_id_subject": subject_id,
                "teacher_id": teacher_id,
                "is_active": 1,
            })
            saved_count += 1

    store.replace_table_dicts("replacements_subject", subject_rows)
    store.replace_table_dicts("replacements_teacher", teacher_rows)
    store.replace_table_dicts("replacements_lesson", lesson_rows)
    return deactivated, saved_count


@login_required
@require_GET
def available_rooms(request):
    """
    Возвращает список кабинетов для замещения на указанный урок.
    В список попадают все уникальные кабинеты из расписания и сохранённых замен.
    Для каждого кабинета вычисляется статус: свободен или занят (есть пересечение по времени).
    """
    if not (request.user.is_superuser or _is_guest_user(request.user) or getattr(request.user, "can_calendar", False) or getattr(request.user, "can_calls", False)):
        return HttpResponse("Forbidden", status=403)

    try:
        lesson_id = request.GET.get('lesson_id')
        date_str = request.GET.get('date')
        day = request.GET.get('day')
        start_str = request.GET.get('start')
        end_str = request.GET.get('end')

        if _use_gsheets_backend():
            active_lessons = _gs_active_lessons_rows()
            repl_rows = _gs_store().get_table_dicts("replacements_replacement")
            lessons_by_id = {as_int(l.get("lesson_id")): l for l in active_lessons if as_int(l.get("lesson_id")) is not None}

            rooms_set = set()
            for l in active_lessons:
                r = str(l.get("classroom") or "").strip()
                if r:
                    rooms_set.add(r)
            for r in repl_rows:
                rr = str(r.get("replacement_classroom") or "").strip()
                if rr:
                    rooms_set.add(rr)
            rooms = sorted(rooms_set)

            if not (date_str and day and start_str and end_str and lesson_id):
                return JsonResponse({'rooms': [{'value': room, 'status': ''} for room in rooms]})

            try:
                start_t = datetime.strptime(start_str, "%H:%M").time()
                end_t = datetime.strptime(end_str, "%H:%M").time()
            except Exception:
                return JsonResponse({'rooms': [{'value': room, 'status': ''} for room in rooms]})

            lesson_obj = lessons_by_id.get(as_int(lesson_id))

            def is_parallel_allowed(item):
                if not lesson_obj:
                    return False
                return (
                    str(item.get('class_group') or "") == str(lesson_obj.get("class_group") or "")
                    and as_int(item.get('lesson_number')) == as_int(lesson_obj.get("lesson_number"))
                )

            result = []
            for room in rooms:
                busy = False
                for l in active_lessons:
                    if str(l.get("classroom") or "") != room:
                        continue
                    if str(l.get("day_of_week") or "") != str(day):
                        continue
                    l_start = _gs_parse_time(l.get("start_time"))
                    l_end = _gs_parse_time(l.get("end_time"))
                    if not l_start or not l_end:
                        continue
                    if not overlaps(start_t, end_t, l_start, l_end):
                        continue
                    if lesson_obj and as_int(l.get("lesson_id")) == as_int(lesson_obj.get("lesson_id")):
                        continue
                    if not is_parallel_allowed(l):
                        busy = True
                        break
                if busy:
                    result.append({'value': room, 'status': 'Занят'})
                    continue

                for repl in repl_rows:
                    if str(repl.get("date") or "") != str(date_str):
                        continue
                    if str(repl.get("replacement_classroom") or "") != room:
                        continue
                    l2 = lessons_by_id.get(as_int(repl.get("lesson_id")))
                    if not l2:
                        continue
                    if lesson_obj and as_int(l2.get("lesson_id")) == as_int(lesson_obj.get("lesson_id")):
                        continue
                    if str(l2.get("day_of_week") or "") != str(day):
                        continue
                    s2 = _gs_parse_time(l2.get("start_time"))
                    e2 = _gs_parse_time(l2.get("end_time"))
                    if not s2 or not e2:
                        continue
                    if overlaps(start_t, end_t, s2, e2) and not is_parallel_allowed(l2):
                        busy = True
                        break
                result.append({'value': room, 'status': 'Занят' if busy else 'Свободен'})
            return JsonResponse({'rooms': result})

                                                                          
        rooms_set = set()
        for r in Lesson.objects.exclude(classroom__isnull=True).exclude(classroom='').values_list('classroom', flat=True).distinct():
            rooms_set.add(str(r))
        for r in Replacement.objects.exclude(replacement_classroom__isnull=True).exclude(replacement_classroom='').values_list('replacement_classroom', flat=True).distinct():
            rooms_set.add(str(r))

                                          
        rooms = sorted(rooms_set)

                                                                                          
        if not (date_str and day and start_str and end_str and lesson_id):
            return JsonResponse({
                'rooms': [{'value': room, 'status': ''} for room in rooms]
            })

                                         
        try:
            selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
            start_t = datetime.strptime(start_str, "%H:%M").time()
            end_t = datetime.strptime(end_str, "%H:%M").time()
        except Exception:
            return JsonResponse({'rooms': [{'value': room, 'status': ''} for room in rooms]})

                                                           
        try:
            lesson_obj = Lesson.objects.filter(id=int(lesson_id)).first()
        except Exception:
            lesson_obj = None

        def is_parallel_allowed(item):
            if not lesson_obj:
                return False
            return (
                item.get('class_group') == lesson_obj.class_group and
                item.get('lesson_number') == lesson_obj.lesson_number
            )

        result = []
        for room in rooms:
            busy = False

                                                   
            for l in _active_lessons().filter(
                classroom=room,
                day_of_week=day,
                start_time__lt=end_t,
                end_time__gt=start_t
            ).values('start_time', 'end_time', 'class_group', 'lesson_number', 'id'):
                if lesson_obj and l.get('id') == lesson_obj.id:
                    continue
                if not is_parallel_allowed(l):
                    busy = True
                    break
            if busy:
                result.append({'value': room, 'status': 'Занят'})
                continue

                                                      
            for repl in Replacement.objects.filter(
                date=selected_date,
                replacement_classroom=room
            ).select_related('lesson'):
                l2 = repl.lesson
                if not l2:
                    continue
                if lesson_obj and l2.id == lesson_obj.id:
                    continue
                if l2.day_of_week != day:
                    continue
                if l2.start_time < end_t and l2.end_time > start_t:
                    item = {
                        'class_group': l2.class_group,
                        'lesson_number': l2.lesson_number
                    }
                    if not is_parallel_allowed(item):
                        busy = True
                        break
            if busy:
                result.append({'value': room, 'status': 'Занят'})
            else:
                result.append({'value': room, 'status': 'Свободен'})

        return JsonResponse({'rooms': result})

    except Exception:
        return _api_error(
            "Не удалось проверить занятость кабинетов",
            status=500,
            code="internal_error",
            rooms=[],
        )


@login_required
@require_GET
def room_conflicts_api(request):
    """
    Возвращает список конфликтов занятости выбранного кабинета по времени.
    Используется на фронте для показа предупреждений и требования подтверждения при выборе кабинета.
    """
    if not (request.user.is_superuser or _is_guest_user(request.user) or getattr(request.user, "can_calendar", False) or getattr(request.user, "can_calls", False)):
        return HttpResponse("Forbidden", status=403)

    classroom = (request.GET.get('classroom') or '').strip()
    date_str = request.GET.get('date')
    day = request.GET.get('day')
    start_str = request.GET.get('start')
    end_str = request.GET.get('end')
    lesson_id = request.GET.get('lesson_id')

                                          
    if not (classroom and date_str and day and start_str and end_str):
        return JsonResponse({'conflicts': []})

    try:
        selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
        start_t = datetime.strptime(start_str, "%H:%M").time()
        end_t = datetime.strptime(end_str, "%H:%M").time()
    except Exception:
        return JsonResponse({'conflicts': []})

    if _use_gsheets_backend():
        active_lessons = _gs_active_lessons_rows()
        repl_rows = _gs_store().get_table_dicts("replacements_replacement")
        lessons_by_id = {as_int(l.get("lesson_id")): l for l in active_lessons if as_int(l.get("lesson_id")) is not None}
        target_lesson = lessons_by_id.get(as_int(lesson_id)) if lesson_id else None

        def is_parallel_allowed(item):
            if not target_lesson:
                return False
            return (
                str(item.get('class_group') or "") == str(target_lesson.get("class_group") or "")
                and as_int(item.get('lesson_number')) == as_int(target_lesson.get("lesson_number"))
            )

        conflicts = []
        for l in active_lessons:
            if str(l.get("classroom") or "") != classroom:
                continue
            if str(l.get("day_of_week") or "") != str(day):
                continue
            s = _gs_parse_time(l.get("start_time"))
            e = _gs_parse_time(l.get("end_time"))
            if not s or not e or not overlaps(start_t, end_t, s, e):
                continue
            if target_lesson and as_int(l.get("lesson_id")) == as_int(target_lesson.get("lesson_id")):
                continue
            if not is_parallel_allowed(l):
                conflicts.append({
                    'type': 'room_lesson',
                    'start': _gs_time_str(s),
                    'end': _gs_time_str(e),
                    'class_group': l.get('class_group'),
                    'lesson_number': as_int(l.get('lesson_number')),
                })

        for repl in repl_rows:
            if str(repl.get("date") or "") != selected_date.strftime("%Y-%m-%d"):
                continue
            if str(repl.get("replacement_classroom") or "") != classroom:
                continue
            l2 = lessons_by_id.get(as_int(repl.get("lesson_id")))
            if not l2:
                continue
            if target_lesson and as_int(l2.get("lesson_id")) == as_int(target_lesson.get("lesson_id")):
                continue
            if str(l2.get("day_of_week") or "") != str(day):
                continue
            s2 = _gs_parse_time(l2.get("start_time"))
            e2 = _gs_parse_time(l2.get("end_time"))
            if not s2 or not e2 or not overlaps(start_t, end_t, s2, e2):
                continue
            if not is_parallel_allowed(l2):
                conflicts.append({
                    'type': 'room_replacement',
                    'start': _gs_time_str(s2),
                    'end': _gs_time_str(e2),
                    'class_group': l2.get('class_group'),
                    'lesson_number': as_int(l2.get('lesson_number')),
                })

        return JsonResponse({'conflicts': conflicts, 'busy': len(conflicts) > 0})

                                                    
    try:
        target_lesson = Lesson.objects.filter(id=int(lesson_id)).first() if lesson_id else None
    except Exception:
        target_lesson = None

    def is_parallel_allowed(item):
        if not target_lesson:
            return False
        return (
            item.get('class_group') == target_lesson.class_group and
            item.get('lesson_number') == target_lesson.lesson_number
        )

    conflicts = []

                                         
    for l in _active_lessons().filter(
        classroom=classroom,
        day_of_week=day,
        start_time__lt=end_t,
        end_time__gt=start_t
    ).values('start_time', 'end_time', 'class_group', 'lesson_number', 'id'):
                                 
        if target_lesson and l.get('id') == target_lesson.id:
            continue
        if not is_parallel_allowed(l):
            conflicts.append({
                'type': 'room_lesson',
                'start': l['start_time'].strftime('%H:%M') if hasattr(l['start_time'], 'strftime') else str(l['start_time']),
                'end': l['end_time'].strftime('%H:%M') if hasattr(l['end_time'], 'strftime') else str(l['end_time']),
                'class_group': l.get('class_group'),
                'lesson_number': l.get('lesson_number'),
            })

                                                 
    for repl in Replacement.objects.filter(
        date=selected_date,
        replacement_classroom=classroom
    ).select_related('lesson'):
        l2 = repl.lesson
        if not l2:
            continue
        if target_lesson and l2.id == target_lesson.id:
            continue
        if l2.day_of_week != day:
            continue
        if l2.start_time < end_t and l2.end_time > start_t:
            item = {
                'class_group': l2.class_group,
                'lesson_number': l2.lesson_number
            }
            if not is_parallel_allowed(item):
                conflicts.append({
                    'type': 'room_replacement',
                    'start': l2.start_time.strftime('%H:%M') if hasattr(l2.start_time, 'strftime') else str(l2.start_time),
                    'end': l2.end_time.strftime('%H:%M') if hasattr(l2.end_time, 'strftime') else str(l2.end_time),
                    'class_group': l2.class_group,
                    'lesson_number': l2.lesson_number,
                })

    return JsonResponse({'conflicts': conflicts, 'busy': any(conflicts)})


def _effective_shift_for_class(class_group: str, stored_shift: int | None = None) -> int | None:
    """Return school-rule shift for a class label.

    For 1–4 -> 1st shift; 6–8 -> 2nd shift; 5/9/10/11 -> 1st shift.
    If class/grade can't be parsed, falls back to stored_shift.
    """
    grade = extract_grade(class_group)
    sh = infer_shift_by_grade(grade)
    return sh if sh in (1, 2) else stored_shift


def _effective_times_for_lesson(class_group: str, lesson_number: int, shift: int | None,
                               fallback_start, fallback_end):
    """Return (start,end) times for a lesson.

    Priority:
      1) ClassSchedule row for the same (lesson_number, shift) AND exact class_group match.
      2) ClassSchedule row for the same (lesson_number, shift) with the same grade.
      3) Fallback to Lesson.start_time/Lesson.end_time.

    Always returns a tuple (start_time, end_time).
    """
                     
    if not shift or not class_group or not lesson_number:
        return fallback_start, fallback_end

    grade = extract_grade(class_group)
    if grade is None:
        return fallback_start, fallback_end

    try:
        rows = ClassSchedule.objects.filter(lesson_number=lesson_number, shift=int(shift))
                           
        cls_norm = (class_group or '').strip().lower()
        for r in rows:
            if (r.class_group or '').strip().lower() == cls_norm:
                return (r.start_time or fallback_start, r.end_time or fallback_end)
                          
        for r in rows:
            if extract_grade(r.class_group) == grade:
                return (r.start_time or fallback_start, r.end_time or fallback_end)
    except Exception:
                                           
        pass

    return fallback_start, fallback_end


def _active_lessons():
    """QuerySet of currently active lessons (current schedule)."""
    return Lesson.objects.filter(is_active=True)


def _teacher_replacements():
    """Replacements for teacher substitutions (exclude cabinet-only replacements)."""
    return Replacement.objects.exclude(replacement_teacher_id=F("original_teacher_id"))


def _cabinet_replacements():
    """Replacements created for cabinet changes (teacher stays the same)."""
    return Replacement.objects.filter(replacement_teacher_id=F("original_teacher_id"))


@login_required
def calendar_view(request):
                                                                           
    if not (request.user.is_superuser or _is_guest_user(request.user) or _is_teacher_user(request.user) or getattr(request.user, 'can_calendar', False)):
        return HttpResponse("Forbidden", status=403)
    if _is_teacher_user(request.user):
        teacher = _resolve_teacher_for_user(request.user)
        return render(
            request,
            "teacher_calendar.html",
            {
                "teacher_name": teacher.full_name if teacher else (request.user.full_name or request.user.username),
                "teacher_found": bool(teacher),
            },
        )
    day = request.GET.get('day', 'пн')
    return render(request, 'calendar.html', {'day': day})


def _norm_name_key(value: str) -> str:
    return re.sub(r"[^0-9a-zа-яё]+", "", (value or "").casefold(), flags=re.IGNORECASE)


def _norm_text_key(value: str) -> str:
    return re.sub(r"[^0-9a-zа-яё]+", "", (value or "").casefold(), flags=re.IGNORECASE)


def _is_generic_vacancy_label(value: str | None) -> bool:
    v = " ".join((value or "").split()).strip()
    if not v:
        return False
    key = _norm_text_key(v)
    return key in {"вакансия", "vacancy"}


def _subject_matches(subject_from_docx: str, lesson_subject: str) -> bool:
    docx_key = _norm_text_key(subject_from_docx)
    lesson_key = _norm_text_key(lesson_subject)
    if not docx_key or not lesson_key:
        return False
    if docx_key == lesson_key:
        return True

    aliases = {
        "англязык": "английскийязык",
        "англ": "английскийязык",
        "матпракт": "математическийпрактикум",
        "читгр": "читательскаяграмотность",
        "ров": "разговорыоважном",
    }
    docx_norm = aliases.get(docx_key, docx_key)
    lesson_norm = aliases.get(lesson_key, lesson_key)
    return (
        docx_norm == lesson_norm
        or docx_norm in lesson_norm
        or lesson_norm in docx_norm
    )


def _parse_docx_rows(uploaded_file) -> list[dict]:
    doc = Document(uploaded_file)
    rows = []

    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join((c.text or "").split()) for c in row.cells]
            if len(cells) < 6:
                                                                                     
                continue

            c0 = (cells[0] or "").strip().casefold()
            if "номер" in c0 and "урок" in c0:
                continue
            if "смена" in c0:
                continue

            try:
                lesson_number = int(float((cells[0] or "").replace(",", ".")))
            except Exception:
                continue

            class_group = (cells[1] or "").strip()
            classroom = (cells[2] or "").strip()
            subject_name = (cells[3] or "").strip()
            original_teacher = (cells[4] or "").strip()
            replacement_teacher = (cells[5] or "").strip()

            if not (class_group and replacement_teacher):
                continue

            rows.append(
                {
                    "lesson_number": lesson_number,
                    "class_group": class_group,
                    "class_key": _norm_text_key(class_group),
                    "classroom": classroom,
                    "subject_name": subject_name,
                    "original_teacher": original_teacher,
                    "replacement_teacher": replacement_teacher,
                }
            )
    return rows


def _pick_best_lesson_for_docx_row(day_lessons: list[Lesson], row: dict) -> Lesson | None:
    row_class_key = row.get("class_key") or ""
    row_number = int(row.get("lesson_number") or 0)
    row_original_name = (row.get("original_teacher") or "").strip()
    row_subject = (row.get("subject_name") or "").strip()
    row_room = (row.get("classroom") or "").strip()

    initial = [
        l for l in day_lessons
        if _norm_text_key(l.class_group or "") == row_class_key and int(l.lesson_number or 0) == row_number
    ]
    if not initial:
        return None
    if len(initial) == 1:
        return initial[0]

    row_original_key = _norm_name_key(row_original_name)
    row_original_is_generic_vacancy = _is_generic_vacancy_label(row_original_name)
    row_room_key = _norm_text_key(row_room)
    best = None
    best_score = -1
    for lesson in initial:
        score = 0
        lesson_teacher_name = lesson.teacher.full_name if lesson.teacher else ""
        if row_original_key and _norm_name_key((lesson.teacher.full_name if lesson.teacher else "")) == row_original_key:
            score += 5
        if row_original_is_generic_vacancy and _is_vacancy_teacher_name(lesson_teacher_name):
            score += 5
        if row_subject and lesson.subject and _subject_matches(row_subject, lesson.subject.name):
            score += 3
        if row_room_key and _norm_text_key(lesson.classroom or "") == row_room_key:
            score += 1
        if score > best_score:
            best_score = score
            best = lesson

    if best_score <= 0:
        return None
    return best


def _get_or_create_teacher_by_name(raw_name: str, teacher_by_key: dict[str, list[Teacher]]) -> Teacher:
    full_name = " ".join((raw_name or "").split()).strip()
    if not full_name:
        raise ValueError("Пустое имя учителя")

    key = _norm_name_key(full_name)
    candidates = teacher_by_key.get(key, [])
    if len(candidates) == 1:
        return candidates[0]
    if len(candidates) > 1:
        for t in candidates:
            if (t.full_name or "").casefold() == full_name.casefold():
                return t
        return candidates[0]

    t = Teacher.objects.create(full_name=full_name, specialization="", hours_per_week=0)
    teacher_by_key.setdefault(key, []).append(t)
    return t


@login_required
@ratelimit(key="user_or_ip", rate="20/h", block=True, method="POST")
def import_replacements_docx(request):
    guest_forbidden = _deny_guest_write_json(request)
    if guest_forbidden:
        return guest_forbidden

    if request.method != "POST":
        return JsonResponse({"error": "Метод не поддерживается"}, status=405)

    if not (request.user.is_superuser or getattr(request.user, "can_calendar", False)):
        return JsonResponse({"error": "Forbidden"}, status=403)

    uploaded_file = request.FILES.get("file")
    date_str = (request.POST.get("date") or "").strip()
    replace_all = str(request.POST.get("replace_all", "1")).strip().lower() in {"1", "true", "yes", "y"}

    validation_error = _validate_docx_import_request(uploaded_file=uploaded_file, date_str=date_str)
    if validation_error:
        return JsonResponse({"error": validation_error}, status=400)

    try:
        target_date = datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception:
        return JsonResponse({"error": "Некорректная дата. Ожидается YYYY-MM-DD"}, status=400)

    file_bytes = uploaded_file.read()
    if not file_bytes:
        return JsonResponse({"error": "Файл пустой"}, status=400)

    try:
        job = DocxImportTask.objects.create(
            created_by=request.user,
            file_name=uploaded_file.name,
            date=target_date,
            replace_all=replace_all,
            status=DocxImportTask.STATUS_QUEUED,
        )
    except (OperationalError, ProgrammingError):
        return JsonResponse(
            {"error": "Таблица задач импорта не создана. Выполните миграции: python manage.py migrate"},
            status=500,
        )
    except Exception:
        logger.exception("Failed to create DOCX import job")
        return JsonResponse({"error": "Не удалось создать задачу импорта"}, status=500)

    try:
        from .tasks import import_replacements_docx_task

        async_result = import_replacements_docx_task.delay(
            job_id=str(job.id),
            file_bytes_b64=base64.b64encode(file_bytes).decode("ascii"),
            date_str=target_date.strftime("%Y-%m-%d"),
            replace_all=replace_all,
            actor_user_id=request.user.id,
        )
        job.celery_task_id = async_result.id or ""
        job.save(update_fields=["celery_task_id", "updated_at"])
    except Exception:
        logger.exception("Failed to enqueue DOCX import job")
        job.status = DocxImportTask.STATUS_FAILED
        job.error = "Не удалось поставить задачу в Celery. Проверьте redis/celery_worker."
        job.save(update_fields=["status", "error", "updated_at"])
        return JsonResponse({"error": job.error, "job_id": str(job.id)}, status=500)

    log_activity(
        request,
        "replacements_import_docx_enqueued",
        {
            "job_id": str(job.id),
            "date": target_date.strftime("%Y-%m-%d"),
            "file_name": uploaded_file.name,
            "replace_all": replace_all,
        },
    )

    return JsonResponse(
        {
            "status": "queued",
            "job_id": str(job.id),
            "date": target_date.strftime("%Y-%m-%d"),
        },
        status=202,
    )

    if not uploaded_file:
        return JsonResponse({"error": "Файл не передан"}, status=400)
    if uploaded_file.size > int(getattr(settings, "MAX_DOCX_UPLOAD_SIZE", 10 * 1024 * 1024)):
        return JsonResponse({"error": "Файл слишком большой"}, status=400)
    if not date_str:
        return JsonResponse({"error": "Параметр date обязателен"}, status=400)
    if not uploaded_file.name.lower().endswith(".docx"):
        return JsonResponse({"error": "Поддерживаются только .docx файлы"}, status=400)
    content_type = str(getattr(uploaded_file, "content_type", "") or "").lower()
    if content_type and content_type not in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",
    }:
        return JsonResponse({"error": "Некорректный тип файла"}, status=400)

    try:
        target_date = datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception:
        return JsonResponse({"error": "Некорректная дата. Ожидается YYYY-MM-DD"}, status=400)

    try:
        parsed_rows = _parse_docx_rows(uploaded_file)
    except Exception as exc:
        return JsonResponse({"error": f"Не удалось прочитать DOCX: {exc}"}, status=400)

    if not parsed_rows:
        return JsonResponse({"error": "В файле не найдены строки замещений"}, status=400)

    day_code = day_short_from_date(target_date)
    day_lessons = list(
        _active_lessons()
        .filter(day_of_week=day_code)
        .select_related("teacher", "subject")
    )

    all_teachers = list(Teacher.objects.all())
    teacher_by_key: dict[str, list[Teacher]] = {}
    for t in all_teachers:
        teacher_by_key.setdefault(_norm_name_key(t.full_name or ""), []).append(t)

    to_create_by_lesson: dict[int, dict] = {}
    unresolved = []
    same_teacher_skipped = 0

    for idx, row in enumerate(parsed_rows, start=1):
        lesson = _pick_best_lesson_for_docx_row(day_lessons, row)
        if not lesson:
            unresolved.append(
                {
                    "row": idx,
                    "class_group": row.get("class_group"),
                    "lesson_number": row.get("lesson_number"),
                    "original_teacher": row.get("original_teacher"),
                    "replacement_teacher": row.get("replacement_teacher"),
                }
            )
            continue

        try:
            original_name_from_docx = row.get("original_teacher") or ""
                                                                                                    
                                                                                               
            if _is_generic_vacancy_label(original_name_from_docx) and lesson.teacher and _is_vacancy_teacher_name(lesson.teacher.full_name):
                original_teacher = lesson.teacher
            else:
                original_teacher = _get_or_create_teacher_by_name(
                    original_name_from_docx or (lesson.teacher.full_name if lesson.teacher else ""),
                    teacher_by_key,
                )
            replacement_teacher = _get_or_create_teacher_by_name(
                row.get("replacement_teacher") or "",
                teacher_by_key,
            )
        except Exception:
            unresolved.append(
                {
                    "row": idx,
                    "class_group": row.get("class_group"),
                    "lesson_number": row.get("lesson_number"),
                    "error": "Не удалось сопоставить учителя",
                }
            )
            continue

        if replacement_teacher.id == original_teacher.id:
            same_teacher_skipped += 1
            continue

        replacement_room = None
        docx_room = (row.get("classroom") or "").strip()
        if docx_room and _norm_text_key(docx_room) != _norm_text_key(lesson.classroom or ""):
            replacement_room = docx_room

        to_create_by_lesson[lesson.id] = {
            "lesson": lesson,
            "original_teacher": original_teacher,
            "replacement_teacher": replacement_teacher,
            "replacement_classroom": replacement_room,
        }

    if not to_create_by_lesson:
        return JsonResponse(
            {
                "error": "Не удалось сопоставить ни одной строки с расписанием",
                "unresolved": unresolved[:50],
            },
            status=400,
        )

    with transaction.atomic():
        if replace_all:
            _teacher_replacements().filter(date=target_date).delete()

        created_count = 0
        for payload in to_create_by_lesson.values():
            Replacement.objects.create(
                lesson=payload["lesson"],
                date=target_date,
                original_teacher=payload["original_teacher"],
                replacement_teacher=payload["replacement_teacher"],
                confirmed=True,
                production_necessity=False,
                ignore_in_reports=False,
                replacement_classroom=payload["replacement_classroom"],
            )
            created_count += 1

    log_activity(request, "replacements_import_docx", {
        "date": target_date.strftime("%Y-%m-%d"),
        "file_name": uploaded_file.name,
        "parsed_rows": len(parsed_rows),
        "created": created_count,
        "unresolved": len(unresolved),
        "replace_all": replace_all,
    })

    return JsonResponse(
        {
            "status": "success",
            "date": target_date.strftime("%Y-%m-%d"),
            "parsed_rows": len(parsed_rows),
            "created": created_count,
            "skipped_same_teacher": same_teacher_skipped,
            "unresolved_count": len(unresolved),
            "unresolved": unresolved[:30],
        }
    )


def _validate_docx_import_request(*, uploaded_file, date_str: str) -> str | None:
    if not uploaded_file:
        return "Файл не передан"
    if uploaded_file.size > int(getattr(settings, "MAX_DOCX_UPLOAD_SIZE", 10 * 1024 * 1024)):
        return "Файл слишком большой"
    if not date_str:
        return "Параметр date обязателен"
    if not uploaded_file.name.lower().endswith(".docx"):
        return "Поддерживаются только .docx файлы"
    content_type = str(getattr(uploaded_file, "content_type", "") or "").lower()
    if content_type and content_type not in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",
    }:
        return "Некорректный тип файла"
    return None


class DocxImportProcessingError(Exception):
    def __init__(self, message: str, unresolved: list[dict] | None = None):
        super().__init__(message)
        self.message = message
        self.unresolved = unresolved or []


def _run_docx_import_core(*, file_bytes: bytes, target_date, replace_all: bool) -> dict:
    try:
        parsed_rows = _parse_docx_rows(io.BytesIO(file_bytes))
    except Exception as exc:
        raise DocxImportProcessingError(f"Не удалось прочитать DOCX: {exc}")

    if not parsed_rows:
        raise DocxImportProcessingError("В файле не найдены строки замещений")

    day_code = day_short_from_date(target_date)
    day_lessons = list(
        _active_lessons()
        .filter(day_of_week=day_code)
        .select_related("teacher", "subject")
    )

    all_teachers = list(Teacher.objects.all())
    teacher_by_key: dict[str, list[Teacher]] = {}
    for t in all_teachers:
        teacher_by_key.setdefault(_norm_name_key(t.full_name or ""), []).append(t)

    to_create_by_lesson: dict[int, dict] = {}
    unresolved = []
    same_teacher_skipped = 0

    for idx, row in enumerate(parsed_rows, start=1):
        lesson = _pick_best_lesson_for_docx_row(day_lessons, row)
        if not lesson:
            unresolved.append(
                {
                    "row": idx,
                    "class_group": row.get("class_group"),
                    "lesson_number": row.get("lesson_number"),
                    "original_teacher": row.get("original_teacher"),
                    "replacement_teacher": row.get("replacement_teacher"),
                }
            )
            continue

        try:
            original_name_from_docx = row.get("original_teacher") or ""
            if _is_generic_vacancy_label(original_name_from_docx) and lesson.teacher and _is_vacancy_teacher_name(lesson.teacher.full_name):
                original_teacher = lesson.teacher
            else:
                original_teacher = _get_or_create_teacher_by_name(
                    original_name_from_docx or (lesson.teacher.full_name if lesson.teacher else ""),
                    teacher_by_key,
                )
            replacement_teacher = _get_or_create_teacher_by_name(
                row.get("replacement_teacher") or "",
                teacher_by_key,
            )
        except Exception:
            unresolved.append(
                {
                    "row": idx,
                    "class_group": row.get("class_group"),
                    "lesson_number": row.get("lesson_number"),
                    "error": "Не удалось сопоставить учителя",
                }
            )
            continue

        if replacement_teacher.id == original_teacher.id:
            same_teacher_skipped += 1
            continue

        replacement_room = None
        docx_room = (row.get("classroom") or "").strip()
        if docx_room and _norm_text_key(docx_room) != _norm_text_key(lesson.classroom or ""):
            replacement_room = docx_room

        to_create_by_lesson[lesson.id] = {
            "lesson": lesson,
            "original_teacher": original_teacher,
            "replacement_teacher": replacement_teacher,
            "replacement_classroom": replacement_room,
        }

    if not to_create_by_lesson:
        raise DocxImportProcessingError(
            "Не удалось сопоставить ни одной строки с расписанием",
            unresolved=unresolved[:50],
        )

    with transaction.atomic():
        if replace_all:
            _teacher_replacements().filter(date=target_date).delete()

        created_count = 0
        for payload in to_create_by_lesson.values():
            Replacement.objects.create(
                lesson=payload["lesson"],
                date=target_date,
                original_teacher=payload["original_teacher"],
                replacement_teacher=payload["replacement_teacher"],
                confirmed=True,
                production_necessity=False,
                ignore_in_reports=False,
                replacement_classroom=payload["replacement_classroom"],
            )
            created_count += 1

    return {
        "date": target_date.strftime("%Y-%m-%d"),
        "parsed_rows": len(parsed_rows),
        "created": created_count,
        "skipped_same_teacher": same_teacher_skipped,
        "unresolved_count": len(unresolved),
        "unresolved": unresolved[:30],
    }


@login_required
@require_GET
def import_replacements_docx_status(request, job_id):
    try:
        job = DocxImportTask.objects.get(id=job_id)
    except DocxImportTask.DoesNotExist:
        return JsonResponse({"error": "Задача не найдена"}, status=404)
    except (OperationalError, ProgrammingError):
        return JsonResponse(
            {"error": "Таблица задач импорта не создана. Выполните миграции: python manage.py migrate"},
            status=500,
        )

    if not (
        request.user.is_superuser
        or getattr(request.user, "can_calendar", False)
        or (job.created_by_id and job.created_by_id == request.user.id)
    ):
        return JsonResponse({"error": "Forbidden"}, status=403)

    payload = {
        "job_id": str(job.id),
        "status": job.status,
        "date": job.date.strftime("%Y-%m-%d"),
        "created_at": job.created_at.isoformat(),
        "updated_at": job.updated_at.isoformat(),
    }
    if job.status == DocxImportTask.STATUS_SUCCESS:
        payload.update(
            {
                "parsed_rows": job.parsed_rows,
                "created": job.created_count,
                "skipped_same_teacher": job.skipped_same_teacher,
                "unresolved_count": job.unresolved_count,
                "unresolved": job.unresolved_preview or [],
            }
        )
    if job.status == DocxImportTask.STATUS_FAILED:
        payload["error"] = job.error or "Ошибка импорта"
        payload["unresolved"] = job.unresolved_preview or []
    return JsonResponse(payload)


@login_required
def activity_logs_view(request):
    """Admin/staff-only audit log viewer."""
                                                                                         
    if not (request.user.is_superuser or request.user.is_staff or getattr(request.user, 'can_logs', False)):
        return HttpResponse("Forbidden", status=403)

    action = (request.GET.get("action") or "").strip()
    user_q = (request.GET.get("user") or "").strip()

    qs = ActivityLog.objects.select_related("user").all()

    if action:
        qs = qs.filter(action=action)
    if user_q:
        qs = qs.filter(
            Q(user__username__icontains=user_q) |
            Q(user__full_name__icontains=user_q)
        )

    qs = qs.order_by("-created_at")[:500]
    actions = list(ActivityLog.objects.values_list("action", flat=True).distinct().order_by("action"))

    return render(request, "activity_logs.html", {
        "logs": qs,
        "actions": actions,
        "action_filter": action,
        "user_filter": user_q,
    })


@login_required
def save_replacements(request):
    from .services.replacements_heavy_service import save_replacements_service
    return save_replacements_service(request)

@login_required
@require_GET
def get_suggestions(request):
    from .services.replacements_heavy_service import get_suggestions_service
    return get_suggestions_service(request)

@login_required
def export_to_docx(request):
    date = request.GET.get('date', datetime.today().strftime('%Y-%m-%d'))

    replacements = _teacher_replacements().filter(date=date).exclude(ignore_in_reports=True).select_related(
        'lesson__subject',
        'original_teacher',
        'replacement_teacher'
    )

    doc = Document()
    doc.add_heading(f'Замены на {date}', level=0)

    def add_grouped_tables(repl_items):
        grouped = {}
        for repl in repl_items:
            subject = repl.get("subject") or "-"
            grouped.setdefault(subject, []).append(repl)

        for subject, items in grouped.items():
            p = doc.add_paragraph()
            run = p.add_run(subject.upper())
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            p.paragraph_format.space_after = Pt(4)

            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            table.autofit = True

            headers = ['Номер урока', 'Класс', 'Кабинет', 'Предмет', 'Отсутствует', 'Замена', 'Подпись']
            hdr = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr[i].text = header

            for repl in items:
                row = table.add_row().cells
                row[0].text = str(repl.get("lesson_number") or "-")
                row[1].text = repl.get("class_group") or "-"
                row[2].text = repl.get("classroom") or "---"
                row[3].text = ""
                psub = row[3].paragraphs[0]
                rsub = psub.add_run(subject)
                rsub.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                row[4].text = repl.get("original_teacher") or "-"
                row[5].text = repl.get("replacement_teacher") or "-"
                row[6].text = ""

    normal_items = []
    production_items = []
    for repl in replacements:
        item = {
            "lesson_number": repl.lesson.lesson_number if repl.lesson else None,
            "class_group": repl.lesson.class_group if repl.lesson else None,
            "classroom": repl.lesson.classroom if repl.lesson else None,
            "subject": repl.lesson.subject.name if repl.lesson and repl.lesson.subject else "-",
            "original_teacher": repl.original_teacher.full_name if repl.original_teacher else "-",
            "replacement_teacher": repl.replacement_teacher.full_name if repl.replacement_teacher else "-",
        }
        if getattr(repl, "production_necessity", False):
            production_items.append(item)
        else:
            normal_items.append(item)

    special_repls = SpecialReplacement.objects.filter(date=date).select_related("replacement_teacher", "original_teacher")
    for sr in special_repls:
        normal_items.append({
            "lesson_number": sr.lesson_number,
            "class_group": sr.class_group,
            "classroom": sr.classroom or "---",
            "subject": sr.subject_name or "-",
            "original_teacher": sr.original_teacher.full_name if sr.original_teacher else "-",
            "replacement_teacher": sr.replacement_teacher.full_name if sr.replacement_teacher else "-",
        })

    if normal_items:
        add_grouped_tables(normal_items)

    if production_items:
        if normal_items:
            doc.add_page_break()
        doc.add_heading("Замещения по производственной необходимости", level=1)
        add_grouped_tables(production_items)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="замены_{date}.docx"'
    doc.save(response)
    log_activity(request, "export_replacements_docx", {"date": date, "count": replacements.count()})
    return response

                                                

@login_required
def cabinet_replacement_view(request):
    """
    Страница для управления замещениями кабинетов. Доступна только суперпользователям
    или пользователям с правами на календарь или звонки (can_calendar, can_calls).
    """
    if not (request.user.is_superuser or _is_guest_user(request.user) or getattr(request.user, "can_calendar", False) or getattr(request.user, "can_calls", False)):
        return HttpResponse("Forbidden", status=403)
    return render(request, "cabinet_replacements.html", {})


@login_required
@require_GET
def cabinet_lessons(request):
    """
    API: Возвращает список уроков на выбранную дату в выбранном кабинете.
    В ответе передаются id, start, end, class, subject, teacher, number.
    """
    if not (request.user.is_superuser or _is_guest_user(request.user) or getattr(request.user, "can_calendar", False) or getattr(request.user, "can_calls", False)):
        return HttpResponse("Forbidden", status=403)

    date_str = request.GET.get("date")
    cabinet = (request.GET.get("cabinet") or "").strip()
    if not date_str or not cabinet:
        return JsonResponse({"lessons": []})
    try:
        selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception:
        return JsonResponse({"lessons": []})
    if _use_gsheets_backend():
        day_short = day_short_from_date(selected_date)
        teacher_map = _gs_teacher_map()
        lesson_map = _gs_lesson_map()
        subject_map = _gs_subject_map()
        result = []
        for l in _gs_active_lessons_rows():
            if str(l.get("classroom") or "") != cabinet:
                continue
            if str(l.get("day_of_week") or "") != day_short:
                continue
            tid = as_int(l.get("teacher_id"))
            sid = as_int(l.get("subject_id_subject"))
            result.append({
                "id": as_int(l.get("lesson_id")),
                "start": _gs_time_str(_gs_parse_time(l.get("start_time"))),
                "end": _gs_time_str(_gs_parse_time(l.get("end_time"))),
                "class": str(l.get("class_group") or ""),
                "subject": subject_map.get(sid, "") if sid is not None else "",
                "teacher": teacher_map.get(tid, "") if tid is not None else "",
                "number": as_int(l.get("lesson_number")),
            })
        result.sort(key=lambda x: ((x.get("number") or 0), (x.get("id") or 0)))
        return JsonResponse({"lessons": result})
    day_short = day_short_from_date(selected_date)
    lessons_qs = _active_lessons().filter(classroom=cabinet, day_of_week=day_short).select_related("teacher", "subject")
    result = []
    for l in lessons_qs:
        result.append({
            "id": l.id,
            "start": l.start_time.strftime("%H:%M") if l.start_time else "",
            "end": l.end_time.strftime("%H:%M") if l.end_time else "",
            "class": l.class_group or "",
            "subject": l.subject.name if l.subject else "",
            "teacher": (l.teacher.full_name if hasattr(l.teacher, "full_name") else (l.teacher.get_full_name() if l.teacher else "")),
            "number": l.lesson_number,
        })
    return JsonResponse({"lessons": result})


@login_required
def save_cabinet_replacements(request):
    """
    API: Сохраняет замены кабинетов. Предполагается, что учитель остаётся тем же,
    а меняется только кабинет. Получает список replacements, каждый с полями
    lesson_id, date, classroom (новый кабинет), confirmed (опционально).
    """
    guest_forbidden = _deny_guest_write_json(request)
    if guest_forbidden:
        return guest_forbidden

    if request.method != "POST":
        return JsonResponse({"error": "Метод не разрешён"}, status=405)
    try:
        data = json.loads(request.body)
    except Exception:
        return JsonResponse({"error": "Некорректный JSON"}, status=400)
    items = data.get("replacements", [])
    if not isinstance(items, list):
        return JsonResponse({"error": "replacements должен быть списком"}, status=400)
    if _use_gsheets_backend():
        try:
            store = _gs_store()
            lessons = _gs_active_lessons_rows()
            lessons_by_id = {as_int(l.get("lesson_id")): l for l in lessons if as_int(l.get("lesson_id")) is not None}
            repl_rows = store.get_table_dicts("replacements_replacement")

            for item in items:
                lesson_id = as_int(item.get("lesson_id"))
                date_str = item.get("date")
                replacement_room = (item.get("classroom") or "").strip() or None
                confirmed = bool(item.get("confirmed"))
                if not (lesson_id and date_str and replacement_room):
                    return JsonResponse({"error": "Некорректные данные"}, status=400)
                try:
                    selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
                except Exception:
                    return JsonResponse({"error": f"Некорректная дата: {date_str}"}, status=400)
                lesson = lessons_by_id.get(lesson_id)
                if not lesson:
                    return JsonResponse({"error": f"Урок с ID {lesson_id} не найден"}, status=404)

                l_start = _gs_parse_time(lesson.get("start_time"))
                l_end = _gs_parse_time(lesson.get("end_time"))
                if not l_start or not l_end:
                    return JsonResponse({"error": "Урок имеет некорректное время"}, status=400)
                lesson_day = str(lesson.get("day_of_week") or "")
                lesson_class = str(lesson.get("class_group") or "")
                lesson_num = as_int(lesson.get("lesson_number"))

                for rec in lessons:
                    if as_int(rec.get("lesson_id")) == lesson_id:
                        continue
                    if str(rec.get("classroom") or "") != replacement_room:
                        continue
                    if str(rec.get("day_of_week") or "") != lesson_day:
                        continue
                    rs = _gs_parse_time(rec.get("start_time"))
                    re_ = _gs_parse_time(rec.get("end_time"))
                    if not rs or not re_ or not overlaps(l_start, l_end, rs, re_):
                        continue
                    if str(rec.get("class_group") or "") == lesson_class and as_int(rec.get("lesson_number")) == lesson_num:
                        continue
                    if not confirmed:
                        return JsonResponse({
                            "error": (
                                f"Требуется подтверждение: выбранный кабинет {replacement_room} занят уроком "
                                f"в расписании."
                            )
                        }, status=400)

                for repl in repl_rows:
                    if str(repl.get("date") or "") != selected_date.strftime("%Y-%m-%d"):
                        continue
                    if str(repl.get("replacement_classroom") or "") != replacement_room:
                        continue
                    if as_int(repl.get("lesson_id")) == lesson_id:
                        continue
                    l2 = lessons_by_id.get(as_int(repl.get("lesson_id")))
                    if not l2 or str(l2.get("day_of_week") or "") != lesson_day:
                        continue
                    s2 = _gs_parse_time(l2.get("start_time"))
                    e2 = _gs_parse_time(l2.get("end_time"))
                    if not s2 or not e2 or not overlaps(l_start, l_end, s2, e2):
                        continue
                    if str(l2.get("class_group") or "") == lesson_class and as_int(l2.get("lesson_number")) == lesson_num:
                        continue
                    if not confirmed:
                        return JsonResponse({
                            "error": (
                                f"Требуется подтверждение: кабинет {replacement_room} уже используется "
                                "в другой замене."
                            )
                        }, status=400)

                idx = None
                for i, r in enumerate(repl_rows):
                    if as_int(r.get("lesson_id")) == lesson_id and str(r.get("date") or "") == selected_date.strftime("%Y-%m-%d"):
                        idx = i
                        break
                if idx is not None:
                    row = repl_rows[idx]
                else:
                    row = {"id": _gs_next_id(repl_rows, "id")}
                original_teacher_id = as_int(lesson.get("teacher_id"))
                row.update({
                    "lesson_id": lesson_id,
                    "date": selected_date.strftime("%Y-%m-%d"),
                    "original_teacher_id": original_teacher_id,
                    "replacement_teacher_id": original_teacher_id,
                    "replacement_classroom": replacement_room,
                    "confirmed": 1 if confirmed else 0,
                    "production_necessity": as_int(row.get("production_necessity"), 0) or 0,
                    "ignore_in_reports": as_int(row.get("ignore_in_reports"), 0) or 0,
                })
                if idx is None:
                    repl_rows.append(row)
                else:
                    repl_rows[idx] = row

            store.replace_table_dicts("replacements_replacement", repl_rows)
            return JsonResponse({"status": "success"})
        except Exception:
            return _api_error("Внутренняя ошибка сервера", status=500, code="internal_error")
    try:
        with transaction.atomic():
            for item in items:
                lesson_id = item.get("lesson_id")
                date_str = item.get("date")
                replacement_room = (item.get("classroom") or "").strip() or None
                confirmed = bool(item.get("confirmed"))
                if not (lesson_id and date_str and replacement_room):
                    return JsonResponse({"error": "Некорректные данные"}, status=400)
                try:
                    selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
                except Exception:
                    return JsonResponse({"error": f"Некорректная дата: {date_str}"}, status=400)
                lesson = Lesson.objects.filter(id=lesson_id).select_related("teacher").first()
                if not lesson:
                    return JsonResponse({"error": f"Урок с ID {lesson_id} не найден"}, status=404)
                original_teacher_id = lesson.teacher_id
                                                                                             
                for rec in _active_lessons().filter(
                    classroom=replacement_room,
                    day_of_week=lesson.day_of_week,
                    start_time__lt=lesson.end_time,
                    end_time__gt=lesson.start_time
                ).exclude(id=lesson.id).values("class_group", "lesson_number"):
                                                                          
                    if rec["class_group"] == lesson.class_group and rec["lesson_number"] == lesson.lesson_number:
                        continue
                    if not confirmed:
                        return JsonResponse({
                            "error": (
                                f"Требуется подтверждение: выбранный кабинет {replacement_room} занят уроком "
                                f"в расписании."
                            )
                        }, status=400)
                                                                             
                existing_repls = Replacement.objects.filter(
                    date=selected_date,
                    replacement_classroom=replacement_room
                ).exclude(lesson_id=lesson.id).select_related("lesson")
                for repl in existing_repls:
                    l2 = repl.lesson
                    if not l2 or l2.day_of_week != lesson.day_of_week:
                        continue
                    if overlaps(lesson.start_time, lesson.end_time, l2.start_time, l2.end_time):
                                                       
                        if l2.class_group == lesson.class_group and l2.lesson_number == lesson.lesson_number:
                            continue
                        if not confirmed:
                            return JsonResponse({
                                "error": (
                                    f"Требуется подтверждение: кабинет {replacement_room} уже используется "
                                    "в другой замене."
                                )
                            }, status=400)
                                         
                Replacement.objects.update_or_create(
                    lesson_id=lesson.id,
                    date=selected_date,
                    defaults={
                        "original_teacher_id": original_teacher_id,
                        "replacement_teacher_id": original_teacher_id,
                        "replacement_classroom": replacement_room,
                        "confirmed": confirmed,
                    }
                )
            return JsonResponse({"status": "success"})
    except Exception:
        return _api_error("Внутренняя ошибка сервера", status=500, code="internal_error")


@login_required
@require_GET
def export_cabinet_docx(request):
    """
    API: экспорт в DOCX отчёта о замене кабинета.
    Требуются параметры date (YYYY-MM-DD) и cabinet (номер/имя кабинета).
    """
    if not _can_calendar_read(request.user):
        return HttpResponse("Forbidden", status=403)
    date_str = request.GET.get("date")
    cabinet = (request.GET.get("cabinet") or "").strip()
    if not date_str or not cabinet:
        return HttpResponse("Недостаточно параметров", status=400)
    try:
        selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception:
        return HttpResponse(f"Некорректная дата: {date_str}", status=400)
    if _use_gsheets_backend():
        lessons = _gs_active_lessons_rows()
        lessons_by_id = {as_int(l.get("lesson_id")): l for l in lessons if as_int(l.get("lesson_id")) is not None}
        subject_map = _gs_subject_map()
        teacher_map = _gs_teacher_map()
        repl_rows = [
            r for r in _gs_store().get_table_dicts("replacements_replacement")
            if str(r.get("date") or "") == selected_date.strftime("%Y-%m-%d")
            and as_int(r.get("replacement_teacher_id")) == as_int(r.get("original_teacher_id"))
        ]
        doc = Document()
        doc.add_heading(f"Замены на {date_str}", level=0)

        def add_grouped_tables(repl_items):
            grouped = {}
            for repl in repl_items:
                subject = repl.get("subject") or "-"
                grouped.setdefault(subject, []).append(repl)

            for subject, items in grouped.items():
                p = doc.add_paragraph()
                run = p.add_run(subject.upper())
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                p.paragraph_format.space_after = Pt(4)

                table = doc.add_table(rows=1, cols=7)
                table.style = "Table Grid"
                table.autofit = True
                headers = ["Номер урока", "Класс", "Кабинет", "Предмет", "Отсутствует", "Замена", "Подпись"]
                hdr = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr[i].text = header
                for item in items:
                    row = table.add_row().cells
                    row[0].text = str(item.get("lesson_number") or "-")
                    row[1].text = item.get("class_group") or "-"
                    row[2].text = item.get("classroom") or "---"
                    row[3].text = ""
                    psub = row[3].paragraphs[0]
                    rsub = psub.add_run(subject)
                    rsub.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    row[4].text = item.get("original_teacher") or "-"
                    row[5].text = item.get("replacement_teacher") or "-"
                    row[6].text = ""

        items = []
        for r in repl_rows:
            lsn = lessons_by_id.get(as_int(r.get("lesson_id")))
            if not lsn or str(lsn.get("classroom") or "") != cabinet:
                continue
            original_id = as_int(r.get("original_teacher_id")) or as_int(lsn.get("teacher_id"))
            sid = as_int(lsn.get("subject_id_subject"))
            items.append({
                "lesson_number": as_int(lsn.get("lesson_number")),
                "class_group": str(lsn.get("class_group") or ""),
                "classroom": str(r.get("replacement_classroom") or lsn.get("classroom") or ""),
                "subject": subject_map.get(sid, "-") if sid is not None else "-",
                "original_teacher": teacher_map.get(original_id, "-") if original_id is not None else "-",
                "replacement_teacher": "-",
            })
        if items:
            add_grouped_tables(items)
        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = f'attachment; filename="замена_кабинета_{cabinet}_{date_str}.docx"'
        doc.save(response)
        log_activity(request, "export_cabinet_docx", {"date": date_str, "cabinet": cabinet, "count": len(items), "backend": "gsheets"})
        return response
    repls = _cabinet_replacements().filter(
        date=selected_date,
        lesson__classroom=cabinet
    ).select_related("lesson__subject", "lesson__teacher", "original_teacher", "replacement_teacher")
    doc = Document()
    doc.add_heading(f"Замены на {date_str}", level=0)

    def add_grouped_tables(repl_items):
        grouped = {}
        for repl in repl_items:
            subject = repl.get("subject") or "-"
            grouped.setdefault(subject, []).append(repl)

        for subject, items in grouped.items():
            p = doc.add_paragraph()
            run = p.add_run(subject.upper())
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            p.paragraph_format.space_after = Pt(4)

            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            table.autofit = True

            headers = ["Номер урока", "Класс", "Кабинет", "Предмет", "Отсутствует", "Замена", "Подпись"]
            hdr = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr[i].text = header

            for item in items:
                row = table.add_row().cells
                row[0].text = str(item.get("lesson_number") or "-")
                row[1].text = item.get("class_group") or "-"
                row[2].text = item.get("classroom") or "---"
                row[3].text = ""
                psub = row[3].paragraphs[0]
                rsub = psub.add_run(subject)
                rsub.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                row[4].text = item.get("original_teacher") or "-"
                row[5].text = item.get("replacement_teacher") or "-"
                row[6].text = ""

    items = []
    for r in repls:
        lsn = r.lesson
        if not lsn:
            continue
        if r.original_teacher:
            original_name = r.original_teacher.full_name
        elif lsn.teacher:
            original_name = (
                lsn.teacher.full_name if hasattr(lsn.teacher, "full_name")
                else (lsn.teacher.get_full_name() if lsn.teacher else "")
            )
        else:
            original_name = ""
        items.append({
            "lesson_number": lsn.lesson_number,
            "class_group": lsn.class_group,
            "classroom": r.replacement_classroom or lsn.classroom,
            "subject": lsn.subject.name if lsn.subject else "-",
            "original_teacher": original_name or "-",
            "replacement_teacher": "-",                                                    
        })

    if items:
        add_grouped_tables(items)
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="замена_кабинета_{cabinet}_{date_str}.docx"'
    doc.save(response)
    log_activity(request, "export_cabinet_docx", {"date": date_str, "cabinet": cabinet, "count": repls.count()})
    return response


logger = logging.getLogger(__name__)


@login_required
@require_GET
def teacher_lessons_view(request, teacher_id, day):
    if not _can_calendar_read(request.user):
        return HttpResponse("Forbidden", status=403)
    try:
        if _use_gsheets_backend():
            teacher_map = _gs_teacher_map()
            if int(teacher_id) not in teacher_map:
                return render(request, 'error.html', {'error': 'Teacher not found'})
            teacher = SimpleNamespace(id=int(teacher_id), full_name=teacher_map.get(int(teacher_id), ""))
            lessons_rows = [
                l for l in _gs_active_lessons_rows()
                if as_int(l.get("teacher_id")) == int(teacher_id) and str(l.get("day_of_week") or "") == str(day)
            ]
            subject_map = _gs_subject_map()
            lessons = []
            for l in lessons_rows:
                sid = as_int(l.get("subject_id_subject"))
                subject_obj = SimpleNamespace(name=subject_map.get(sid, "")) if sid is not None else None
                lessons.append(
                    SimpleNamespace(
                        id=as_int(l.get("lesson_id")),
                        lesson_number=as_int(l.get("lesson_number")),
                        start_time=_gs_parse_time(l.get("start_time")),
                        end_time=_gs_parse_time(l.get("end_time")),
                        class_group=str(l.get("class_group") or ""),
                        classroom=str(l.get("classroom") or ""),
                        shift=as_int(l.get("shift")),
                        day_of_week=str(l.get("day_of_week") or ""),
                        subject=subject_obj,
                        teacher=teacher,
                    )
                )
            lessons = sorted(lessons, key=lambda x: (x.lesson_number or 0, x.id or 0))
        else:
            teacher = get_object_or_404(Teacher, id=teacher_id)
            lessons = _active_lessons().filter(teacher=teacher, day_of_week=day)
        return render(
            request,
            'teacher_lessons.html',
            {
                'teacher': teacher,
                'lessons': lessons,
                'day': day
            }
        )
    except Exception:
        return render(request, "error.html", {"error": "Не удалось открыть расписание учителя"})



@login_required
@require_GET
def teacher_conflicts_api(request):
    """Возвращает конфликты занятости учителя по времени (уроки + уже сохранённые замещения).
    Используется на фронте для показа предупреждения и требования подтверждения."""
    if not _can_calendar_read(request.user):
        return HttpResponse("Forbidden", status=403)
    try:
        teacher_id = int(request.GET.get('teacher_id'))
        date_str = request.GET.get('date')
        day = request.GET.get('day')
        start_str = request.GET.get('start')
        end_str = request.GET.get('end')
        lesson_id = request.GET.get('lesson_id')

        if not (date_str and day and start_str and end_str):
            return JsonResponse({'conflicts': []})

        selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
        start_t = datetime.strptime(start_str, "%H:%M").time()
        end_t = datetime.strptime(end_str, "%H:%M").time()

        if _use_gsheets_backend():
            active_lessons = _gs_active_lessons_rows()
            lessons_by_id = {as_int(l.get("lesson_id")): l for l in active_lessons if as_int(l.get("lesson_id")) is not None}
            repl_rows = _gs_store().get_table_dicts("replacements_replacement")

            target_lesson = lessons_by_id.get(as_int(lesson_id)) if lesson_id else None

            def is_parallel_allowed(item):
                if not target_lesson:
                    return False
                return (
                    str(item.get('class_group') or "") == str(target_lesson.get("class_group") or "")
                    and as_int(item.get('lesson_number')) == as_int(target_lesson.get("lesson_number"))
                )

            conflicts = []
            try:
                target_grade = extract_grade(str(target_lesson.get("class_group") or "")) if target_lesson else None
                if target_grade is not None:
                    grades = set()
                    for l in active_lessons:
                        if as_int(l.get("teacher_id")) != teacher_id:
                            continue
                        g = extract_grade(str(l.get("class_group") or ""))
                        if g is not None:
                            grades.add(g)
                    lvl = "none"
                    if grades:
                        if max(grades) <= 4:
                            lvl = "1-4"
                        elif min(grades) >= 5:
                            lvl = "5-11"
                        else:
                            lvl = "1-11"
                    if target_grade <= 4 and lvl == "5-11":
                        conflicts.insert(0, {"type": "cross_level", "target_grade": int(target_grade)})
                    elif target_grade >= 5 and lvl == "1-4":
                        conflicts.insert(0, {"type": "cross_level", "target_grade": int(target_grade)})
            except Exception:
                pass

            for l in active_lessons:
                if as_int(l.get("teacher_id")) != teacher_id:
                    continue
                if str(l.get("day_of_week") or "") != str(day):
                    continue
                ls = _gs_parse_time(l.get("start_time"))
                le = _gs_parse_time(l.get("end_time"))
                if not ls or not le or not overlaps(start_t, end_t, ls, le):
                    continue
                if not is_parallel_allowed(l):
                    conflicts.append({
                        'type': 'lesson',
                        'start': _gs_time_str(ls),
                        'end': _gs_time_str(le),
                        'class_group': l.get('class_group'),
                        'lesson_number': as_int(l.get('lesson_number')),
                    })

            for r in repl_rows:
                if not _gs_is_teacher_replacement(r):
                    continue
                if str(r.get("date") or "") != selected_date.strftime("%Y-%m-%d"):
                    continue
                if as_int(r.get("replacement_teacher_id")) != teacher_id:
                    continue
                if lesson_id and as_int(r.get("lesson_id")) == as_int(lesson_id):
                    continue
                l2 = lessons_by_id.get(as_int(r.get("lesson_id")))
                if not l2:
                    continue
                if str(l2.get("day_of_week") or "") != str(day):
                    continue
                s2 = _gs_parse_time(l2.get("start_time"))
                e2 = _gs_parse_time(l2.get("end_time"))
                if not s2 or not e2 or not overlaps(start_t, end_t, s2, e2):
                    continue
                if not is_parallel_allowed(l2):
                    conflicts.append({
                        'type': 'replacement',
                        'start': _gs_time_str(s2),
                        'end': _gs_time_str(e2),
                        'class_group': l2.get('class_group'),
                        'lesson_number': as_int(l2.get('lesson_number')),
                    })

            try:
                target_shift = None
                if target_lesson:
                    target_shift = _effective_shift_for_class(
                        str(target_lesson.get("class_group") or ""),
                        as_int(target_lesson.get("shift")),
                    )
                if target_shift in (1, 2):
                    shift1_max_end = None
                    shift2_min_start = None
                    for row in _gs_store().get_table_dicts("replacements_class_schedule"):
                        sh = as_int(row.get("shift"))
                        s = _gs_parse_time(row.get("start_time"))
                        e = _gs_parse_time(row.get("end_time"))
                        if sh == 1 and e:
                            shift1_max_end = e if shift1_max_end is None or e > shift1_max_end else shift1_max_end
                        if sh == 2 and s:
                            shift2_min_start = s if shift2_min_start is None or s < shift2_min_start else shift2_min_start
                    boundary = shift2_min_start
                    if shift1_max_end and shift2_min_start:
                        try:
                            dt1 = datetime.combine(selected_date, shift1_max_end)
                            dt2 = datetime.combine(selected_date, shift2_min_start)
                            boundary = (dt1 + (dt2 - dt1) / 2).time() if dt2 > dt1 else shift2_min_start
                        except Exception:
                            boundary = shift2_min_start

                    has_shift = False
                    for l in active_lessons:
                        if as_int(l.get("teacher_id")) != teacher_id:
                            continue
                        if str(l.get("day_of_week") or "") != str(day):
                            continue
                        sh = _effective_shift_for_class(str(l.get("class_group") or ""), as_int(l.get("shift")))
                        st = _gs_parse_time(l.get("start_time"))
                        if sh not in (1, 2):
                            try:
                                if shift2_min_start and boundary and st and st >= boundary:
                                    sh = 2
                                else:
                                    sh = 1
                            except Exception:
                                sh = 1
                        if int(sh) == int(target_shift):
                            has_shift = True
                            break
                    if not has_shift:
                        conflicts.insert(0, {'type': 'no_lessons', 'shift': int(target_shift)})
            except Exception:
                pass

            return JsonResponse({'conflicts': conflicts, 'busy': any(c.get('type') in ('lesson', 'replacement') for c in conflicts)})

        target_lesson = None
        try:
            if lesson_id:
                target_lesson = Lesson.objects.filter(id=int(lesson_id)).first()
        except Exception:
            target_lesson = None

        def is_parallel_allowed(item):
            if not target_lesson:
                return False
            return (
                item.get('class_group') == target_lesson.class_group and
                item.get('lesson_number') == target_lesson.lesson_number
            )

                                                                       
        lesson_qs = _active_lessons().filter(
            teacher_id=teacher_id,
            day_of_week=day,
            start_time__lt=end_t,
            end_time__gt=start_t
        ).values('start_time', 'end_time', 'class_group', 'lesson_number')

                                                        
        repl_qs = _teacher_replacements().filter(
            date=selected_date,
            replacement_teacher_id=teacher_id
        ).select_related('lesson')

        conflicts = []

                                                               
        try:
            target_grade = extract_grade(target_lesson.class_group) if target_lesson else None
            if target_grade is not None:
                grades = set()
                for cls in _active_lessons().filter(teacher_id=teacher_id).values_list("class_group", flat=True):
                    g = extract_grade(cls)
                    if g is not None:
                        grades.add(g)
                lvl = "none"
                if grades:
                    if max(grades) <= 4:
                        lvl = "1-4"
                    elif min(grades) >= 5:
                        lvl = "5-11"
                    else:
                        lvl = "1-11"
                if target_grade <= 4 and lvl == "5-11":
                    conflicts.insert(0, {"type": "cross_level", "target_grade": int(target_grade)})
                elif target_grade >= 5 and lvl == "1-4":
                    conflicts.insert(0, {"type": "cross_level", "target_grade": int(target_grade)})
        except Exception:
            pass
        for l in lesson_qs:
            if not is_parallel_allowed(l):
                conflicts.append({
                    'type': 'lesson',
                    'start': l['start_time'].strftime('%H:%M'),
                    'end': l['end_time'].strftime('%H:%M'),
                    'class_group': l.get('class_group'),
                    'lesson_number': l.get('lesson_number'),
                })

        for r in repl_qs:
            if lesson_id and r.lesson_id == int(lesson_id):
                continue
            l2 = r.lesson
            if not l2:
                continue
            if l2.day_of_week != day:
                continue
                           
            if l2.start_time < end_t and l2.end_time > start_t:
                item = {
                    'class_group': l2.class_group,
                    'lesson_number': l2.lesson_number,
                }
                if not is_parallel_allowed(item):
                    conflicts.append({
                        'type': 'replacement',
                        'start': l2.start_time.strftime('%H:%M'),
                        'end': l2.end_time.strftime('%H:%M'),
                        'class_group': l2.class_group,
                        'lesson_number': l2.lesson_number,
                    })

                                                                                                             
        try:
            from django.db.models import Max, Min
            target_shift = None
            if target_lesson:
                target_shift = _effective_shift_for_class(target_lesson.class_group, getattr(target_lesson, 'shift', None))
            if target_shift in (1, 2):
                has_shift = False
                shift1_max_end = ClassSchedule.objects.filter(shift=1).aggregate(mx=Max("end_time"))["mx"]
                shift2_min_start = ClassSchedule.objects.filter(shift=2).aggregate(mn=Min("start_time"))["mn"]
                boundary = None
                if shift2_min_start and shift1_max_end:
                    try:
                        dt1 = datetime.combine(selected_date, shift1_max_end)
                        dt2 = datetime.combine(selected_date, shift2_min_start)
                        boundary = (dt1 + (dt2 - dt1) / 2).time() if dt2 > dt1 else shift2_min_start
                    except Exception:
                        boundary = shift2_min_start
                elif shift2_min_start:
                    boundary = shift2_min_start

                for cls, st, sh_stored in _active_lessons().filter(teacher_id=teacher_id, day_of_week=day).values_list('class_group', 'start_time', 'shift'):
                    sh = _effective_shift_for_class(cls, sh_stored)
                    if sh not in (1, 2):
                        try:
                            if shift2_min_start and boundary and st >= boundary:
                                sh = 2
                            else:
                                sh = 1
                        except Exception:
                            sh = 1
                    if int(sh) == int(target_shift):
                        has_shift = True
                        break

                if not has_shift:
                    conflicts.insert(0, {'type': 'no_lessons', 'shift': int(target_shift)})
        except Exception:
            pass

        return JsonResponse({'conflicts': conflicts, 'busy': any(c.get('type') in ('lesson','replacement') for c in conflicts)})

    except Exception:
        return _api_error(
            "Не удалось проверить конфликты учителя",
            status=500,
            code="internal_error",
            conflicts=[],
        )


@login_required
@require_GET
def teacher_search(request):
    """Поиск учителей для ручного выбора замещения (Select2).

    Если переданы контекстные параметры урока (date/day/start/end/lesson_id),
    возвращает подпись вида: "ФИО (Свободен)" или "ФИО (Урок)".

    Также помечает занятых как disabled (чтобы нельзя было выбрать),
    кроме кейса «вторая группа» (параллельный урок того же класса+номера).
    """
    if not _can_calendar_read(request.user):
        return JsonResponse({"results": []}, status=403)

    term = (request.GET.get('term') or '').strip()
    teacher_id_param = request.GET.get('teacher_id')
    date_str = request.GET.get('date')
    day = request.GET.get('day')
    start_str = request.GET.get('start')
    end_str = request.GET.get('end')
    lesson_id = request.GET.get('lesson_id')

    if _use_gsheets_backend():
        teacher_rows = _gs_store().get_table_dicts("replacements_teacher")
        if teacher_id_param:
            try:
                tid_target = int(teacher_id_param)
                teachers = [t for t in teacher_rows if as_int(t.get("teacher_id")) == tid_target]
            except Exception:
                teachers = []
        else:
            teachers = sorted(teacher_rows, key=lambda t: str(t.get("full_name") or "").casefold())

        if not all([date_str, day, start_str, end_str]):
            teachers = [t for t in teachers if _name_matches_term_ci(str(t.get("full_name") or ""), term)]
            if not teacher_id_param:
                teachers = teachers[:25]
            results = [{'id': as_int(t.get("teacher_id")), 'text': str(t.get("full_name") or "")} for t in teachers if as_int(t.get("teacher_id")) is not None]
            return JsonResponse({'results': results})

        try:
            selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
            start_t = datetime.strptime(start_str, "%H:%M").time()
            end_t = datetime.strptime(end_str, "%H:%M").time()
        except Exception:
            results = [{'id': as_int(t.get("teacher_id")), 'text': str(t.get("full_name") or "")} for t in teachers if as_int(t.get("teacher_id")) is not None]
            return JsonResponse({'results': results})

        teachers = [t for t in teachers if _name_matches_term_ci(str(t.get("full_name") or ""), term)]
        if not teacher_id_param:
            teachers = teachers[:25]

        active_lessons = _gs_active_lessons_rows()
        lessons_by_id = {as_int(l.get("lesson_id")): l for l in active_lessons if as_int(l.get("lesson_id")) is not None}
        target_lesson = lessons_by_id.get(as_int(lesson_id)) if lesson_id else None

        teacher_ids = [as_int(t.get("teacher_id")) for t in teachers if as_int(t.get("teacher_id")) is not None]
        target_grade = extract_grade(str(target_lesson.get("class_group") or "")) if target_lesson else None
        target_is_primary = (target_grade is not None and target_grade <= 4)

        shift1_max_end, shift2_min_start, boundary = _gs_shift_boundary(selected_date)
        target_shift = None
        if target_lesson:
            target_shift = _effective_shift_for_class(str(target_lesson.get("class_group") or ""), as_int(target_lesson.get("shift")))
        if target_shift is None:
            if shift2_min_start and boundary and start_t >= boundary:
                target_shift = 2
            else:
                target_shift = 1

        grade_map_all: dict[int, set[int]] = {}
        for l in active_lessons:
            tid = as_int(l.get("teacher_id"))
            if tid not in teacher_ids:
                continue
            g = extract_grade(str(l.get("class_group") or ""))
            if g is None:
                continue
            grade_map_all.setdefault(tid, set()).add(g)

        def teacher_level_all(t_id: int) -> str:
            grades = grade_map_all.get(t_id) or set()
            if not grades:
                return "none"
            if max(grades) <= 4:
                return "1-4"
            if min(grades) >= 5:
                return "5-11"
            return "1-11"

        def grade_ok_for_target(t_id: int) -> bool:
            lvl = teacher_level_all(t_id)
            if target_grade is None:
                return True
            if target_is_primary:
                return lvl in ("1-4", "1-11")
            return lvl in ("5-11", "1-11")

        def _lesson_shift_for(cls: str, st_time):
            g = extract_grade(cls)
            sh = infer_shift_by_grade(g)
            if sh in (1, 2):
                return int(sh)
            try:
                if shift2_min_start and boundary and st_time and st_time >= boundary:
                    return 2
            except Exception:
                pass
            return 1

        day_lessons = []
        for l in active_lessons:
            tid = as_int(l.get("teacher_id"))
            if tid not in teacher_ids:
                continue
            if str(l.get("day_of_week") or "") != str(day):
                continue
            day_lessons.append({
                "teacher_id": tid,
                "start_time": _gs_parse_time(l.get("start_time")),
                "end_time": _gs_parse_time(l.get("end_time")),
                "class_group": str(l.get("class_group") or ""),
                "lesson_number": as_int(l.get("lesson_number")),
            })

        repl_rows = _gs_store().get_table_dicts("replacements_replacement")
        day_repls = []
        for r in repl_rows:
            if not _gs_is_teacher_replacement(r):
                continue
            if str(r.get("date") or "") != selected_date.strftime("%Y-%m-%d"):
                continue
            tid = as_int(r.get("replacement_teacher_id"))
            if tid not in teacher_ids:
                continue
            l2 = lessons_by_id.get(as_int(r.get("lesson_id")))
            if not l2:
                continue
            day_repls.append({
                "teacher_id": tid,
                "start_time": _gs_parse_time(l2.get("start_time")),
                "end_time": _gs_parse_time(l2.get("end_time")),
                "class_group": str(l2.get("class_group") or ""),
                "lesson_number": as_int(l2.get("lesson_number")),
            })

        lesson_map = {}
        for l in day_lessons:
            lesson_map.setdefault(l["teacher_id"], []).append(l)

        repl_map = {}
        for r in day_repls:
            repl_map.setdefault(r["teacher_id"], []).append(r)

        def _is_parallel_allowed(item) -> bool:
            if not target_lesson:
                return False
            return (
                str(item.get('class_group') or "") == str(target_lesson.get("class_group") or "")
                and as_int(item.get('lesson_number')) == as_int(target_lesson.get("lesson_number"))
            )

        results = []
        for t in teachers:
            tid = as_int(t.get("teacher_id"))
            if tid is None:
                continue
            busy = False
            conflicts = []

            for l in lesson_map.get(tid, []):
                s = l.get("start_time")
                e = l.get("end_time")
                if not s or not e:
                    continue
                if overlaps(start_t, end_t, s, e) and not _is_parallel_allowed(l):
                    busy = True
                    conflicts.append({
                        'type': 'lesson',
                        'start': _gs_time_str(s),
                        'end': _gs_time_str(e),
                        'class_group': l.get('class_group'),
                        'lesson_number': l.get('lesson_number'),
                    })

            for r in repl_map.get(tid, []):
                s = r.get("start_time")
                e = r.get("end_time")
                if not s or not e:
                    continue
                if overlaps(start_t, end_t, s, e) and not _is_parallel_allowed(r):
                    busy = True
                    conflicts.append({
                        'type': 'replacement',
                        'start': _gs_time_str(s),
                        'end': _gs_time_str(e),
                        'class_group': r.get('class_group'),
                        'lesson_number': r.get('lesson_number'),
                    })

            status = "Свободен" if not busy else "Урок"
            name = str(t.get("full_name") or "")
            is_vacancy = _is_vacancy_teacher_name(name)
            grade_ok = grade_ok_for_target(tid)
            disabled = False

            no_lessons_shift = True
            for l in lesson_map.get(tid, []):
                sh = _lesson_shift_for(str(l.get('class_group') or ""), l.get('start_time'))
                if int(sh) == int(target_shift):
                    no_lessons_shift = False
                    break

            if is_vacancy:
                text = f"{name} (Вакансия)"
            else:
                note = " (в этот день нет уроков в эту смену)" if no_lessons_shift else ""
                lvl_warn = " — ⚠ другая ступень, нужно подтверждение" if not grade_ok else ""
                text = f"{name} ({status}){note}{lvl_warn}"

            results.append({
                'id': tid,
                'text': text,
                'disabled': disabled,
                'busy': bool(busy),
                'conflicts': conflicts,
                'no_lessons_shift': bool(no_lessons_shift),
                'grade_ok': bool(grade_ok),
            })

        return JsonResponse({'results': results})

                                                                      
                                                                                                  
    if teacher_id_param:
        try:
            teachers = Teacher.objects.filter(id=int(teacher_id_param))
        except Exception:
            teachers = Teacher.objects.none()
    else:
        teachers = Teacher.objects.order_by('full_name')

                                                       
    if not all([date_str, day, start_str, end_str]):
        teachers = [t for t in teachers if _name_matches_term_ci(t.full_name, term)]
        if not teacher_id_param:
            teachers = teachers[:25]
        results = [{'id': t.id, 'text': t.full_name} for t in teachers]
        return JsonResponse({'results': results})

    try:
        selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
        start_t = datetime.strptime(start_str, "%H:%M").time()
        end_t = datetime.strptime(end_str, "%H:%M").time()
    except Exception:
        results = [{'id': t.id, 'text': t.full_name} for t in teachers]
        return JsonResponse({'results': results})

    teachers = [t for t in teachers if _name_matches_term_ci(t.full_name, term)]
    if not teacher_id_param:
        teachers = teachers[:25]

    target_lesson = None
    try:
        if lesson_id:
            target_lesson = Lesson.objects.filter(id=int(lesson_id)).first()
    except Exception:
        target_lesson = None

    teacher_ids = [t.id for t in teachers]
                                                                                     
    target_grade = extract_grade(target_lesson.class_group) if target_lesson else None
    target_is_primary = (target_grade is not None and target_grade <= 4)

                                                                             
    from django.db.models import Max, Min
    shift1_max_end = None
    shift2_min_start = None
    boundary = None
    target_shift = None
    if target_lesson:
        target_shift = _effective_shift_for_class(target_lesson.class_group, getattr(target_lesson, 'shift', None))
    if target_shift is None:
                             
        shift1_max_end = ClassSchedule.objects.filter(shift=1).aggregate(mx=Max("end_time"))["mx"]
        shift2_min_start = ClassSchedule.objects.filter(shift=2).aggregate(mn=Min("start_time"))["mn"]
        boundary = None
        if shift2_min_start and shift1_max_end:
            try:
                dt1 = datetime.combine(selected_date, shift1_max_end)
                dt2 = datetime.combine(selected_date, shift2_min_start)
                boundary = (dt1 + (dt2 - dt1) / 2).time() if dt2 > dt1 else shift2_min_start
            except Exception:
                boundary = shift2_min_start
        elif shift2_min_start:
            boundary = shift2_min_start

        if shift2_min_start and boundary and start_t >= boundary:
            target_shift = 2
        else:
            target_shift = 1

                                                        
    grade_map_all: dict[int, set[int]] = {}
    for t_id, cls in _active_lessons().filter(teacher_id__in=teacher_ids).values_list("teacher_id", "class_group"):
        g = extract_grade(cls)
        if g is None:
            continue
        grade_map_all.setdefault(t_id, set()).add(g)

    def teacher_level_all(t_id: int) -> str:
        grades = grade_map_all.get(t_id) or set()
        if not grades:
            return "none"
        if max(grades) <= 4:
            return "1-4"
        if min(grades) >= 5:
            return "5-11"
        return "1-11"

    def grade_ok_for_target(t_id: int) -> bool:
        lvl = teacher_level_all(t_id)
        if target_grade is None:
            return True
        if target_is_primary:
            return lvl in ("1-4", "1-11")
              
        return lvl in ("5-11", "1-11")

                                            
    def _lesson_shift_for(cls: str, st_time):
        g = extract_grade(cls)
        sh = infer_shift_by_grade(g)
        if sh in (1, 2):
            return int(sh)
                             
        try:
            if shift2_min_start and boundary and st_time >= boundary:
                return 2
        except Exception:
            pass
        return 1


                                                          
    day_lessons = list(
        _active_lessons().filter(teacher_id__in=teacher_ids, day_of_week=day)
        .values('teacher_id', 'start_time', 'end_time', 'class_group', 'lesson_number')
    )

                                                    
    day_repls = list(
        _teacher_replacements().filter(date=selected_date, replacement_teacher_id__in=teacher_ids)
        .select_related('lesson')
    )

                        
    lesson_map = {}
    for l in day_lessons:
        lesson_map.setdefault(l['teacher_id'], []).append(l)

    repl_map = {}
    for r in day_repls:
        if not r.lesson:
            continue
        repl_map.setdefault(r.replacement_teacher_id, []).append(
            {
                'start_time': r.lesson.start_time,
                'end_time': r.lesson.end_time,
                'class_group': r.lesson.class_group,
                'lesson_number': r.lesson.lesson_number,
            }
        )

    def _is_parallel_allowed(item) -> bool:
        """Разрешаем пересечение только если это параллельная группа
        (тот же класс и тот же номер урока)."""
        if not target_lesson:
            return False
        return (
            item.get('class_group') == target_lesson.class_group
            and item.get('lesson_number') == target_lesson.lesson_number
        )

    results = []
    for t in teachers:
        busy = False
        parallel_only = True                                                  
        conflicts = []                                                      

                                          
        for l in lesson_map.get(t.id, []):
            if overlaps(start_t, end_t, l['start_time'], l['end_time']):
                if not _is_parallel_allowed(l):
                    busy = True
                    parallel_only = False
                    conflicts.append({
                        'type': 'lesson',
                        'start': l['start_time'].strftime('%H:%M') if hasattr(l['start_time'], 'strftime') else str(l['start_time']),
                        'end': l['end_time'].strftime('%H:%M') if hasattr(l['end_time'], 'strftime') else str(l['end_time']),
                        'class_group': l.get('class_group'),
                        'lesson_number': l.get('lesson_number'),
                    })

                                                     
        for r in repl_map.get(t.id, []):
            if overlaps(start_t, end_t, r['start_time'], r['end_time']):
                if not _is_parallel_allowed(r):
                    busy = True
                    parallel_only = False
                    conflicts.append({
                        'type': 'replacement',
                        'start': r['start_time'].strftime('%H:%M') if hasattr(r['start_time'], 'strftime') else str(r['start_time']),
                        'end': r['end_time'].strftime('%H:%M') if hasattr(r['end_time'], 'strftime') else str(r['end_time']),
                        'class_group': r.get('class_group'),
                        'lesson_number': r.get('lesson_number'),
                    })

        status = "Свободен" if not busy else "Урок"

                                                                                           
        is_vacancy = _is_vacancy_teacher_name(t.full_name)
        grade_ok = grade_ok_for_target(t.id)
        disabled = False

                                                                          
        no_lessons_shift = True
        for l in lesson_map.get(t.id, []):
            sh = _lesson_shift_for(l.get('class_group'), l.get('start_time'))
            if int(sh) == int(target_shift):
                no_lessons_shift = False
                break

        if is_vacancy:
            text = f"{t.full_name} (Вакансия)"
        else:
            note = " (в этот день нет уроков в эту смену)" if no_lessons_shift else ""
            lvl_warn = ""
            if not grade_ok:
                lvl_warn = " — ⚠ другая ступень, нужно подтверждение"
            text = f"{t.full_name} ({status}){note}{lvl_warn}"

        results.append({
            'id': t.id,
            'text': text,
            'disabled': disabled,
            'busy': bool(busy),
            'conflicts': conflicts,
            'no_lessons_shift': bool(no_lessons_shift),
            'grade_ok': bool(grade_ok),
        })

    return JsonResponse({'results': results})


@login_required
@require_GET
def teacher_search_all(request):
    """Teacher search for maintenance operations (includes vacancy teachers)."""
    if not (request.user.is_superuser or getattr(request.user, 'can_upload', False)):
        return HttpResponse("Forbidden", status=403)

    term = (request.GET.get('term') or '').strip()
    if _use_gsheets_backend():
        teachers = sorted(
            _gs_store().get_table_dicts("replacements_teacher"),
            key=lambda t: str(t.get("full_name") or "").casefold(),
        )
        teachers = [t for t in teachers if _name_matches_term_ci(str(t.get("full_name") or ""), term)][:50]
        return JsonResponse({
            'results': [{'id': as_int(t.get("teacher_id")), 'text': str(t.get("full_name") or "")} for t in teachers if as_int(t.get("teacher_id")) is not None]
        })
    teachers = Teacher.objects.order_by('full_name')
    teachers = [t for t in teachers if _name_matches_term_ci(t.full_name, term)][:50]
    return JsonResponse({
        'results': [{'id': t.id, 'text': t.full_name} for t in teachers]
    })


@login_required
@require_GET
def vacancy_teachers_for_date(request):
    """Return vacancy teachers that have lessons on selected date."""
    if not (request.user.is_superuser or _is_guest_user(request.user) or getattr(request.user, 'can_calendar', False)):
        return HttpResponse("Forbidden", status=403)

    date_str = request.GET.get("date")
    if not date_str:
        return JsonResponse({"teachers": []})

    try:
        selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception:
        return JsonResponse({"teachers": []})

    day = day_short_from_date(selected_date)
    if _use_gsheets_backend():
        teachers_map = _gs_teacher_map()
        counts: dict[int, int] = {}
        for l in _gs_active_lessons_rows():
            if str(l.get("day_of_week") or "") != str(day):
                continue
            tid = as_int(l.get("teacher_id"))
            if tid is None:
                continue
            if _is_vacancy_teacher_name(teachers_map.get(tid)):
                counts[tid] = counts.get(tid, 0) + 1
        teachers = [
            {"id": tid, "name": teachers_map.get(tid, ""), "lessons_count": cnt}
            for tid, cnt in sorted(counts.items(), key=lambda x: (teachers_map.get(x[0], "").casefold(), x[0]))
        ]
        return JsonResponse({"teachers": teachers})

    rows = (
        _active_lessons()
        .filter(day_of_week=day)
        .filter(
            Q(teacher__full_name__contains="Вакан")
            | Q(teacher__full_name__contains="вакан")
            | Q(teacher__full_name__icontains="vakans")
        )
        .values("teacher_id", "teacher__full_name")
        .annotate(lessons_count=Count("id"))
        .order_by("teacher__full_name")
    )

    teachers = [
        {
            "id": r["teacher_id"],
            "name": r["teacher__full_name"] or "",
            "lessons_count": r["lessons_count"],
        }
        for r in rows
        if r.get("teacher_id")
    ]
    return JsonResponse({"teachers": teachers})


@login_required
@require_GET
def my_replacements_for_day_api(request):
    if not _is_teacher_user(request.user):
        return HttpResponse("Forbidden", status=403)

    teacher = _resolve_teacher_for_user(request.user)
    if not teacher:
        return JsonResponse({"error": "Учитель для текущего пользователя не найден"}, status=400)

    date_str = (request.GET.get("date") or "").strip()
    if not date_str:
        return JsonResponse({"error": "Параметр date обязателен"}, status=400)

    try:
        selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception:
        return JsonResponse({"error": "Некорректная дата. Ожидается YYYY-MM-DD"}, status=400)

    if _use_gsheets_backend():
        teacher_map = _gs_teacher_map()
        lesson_map = _gs_lesson_map()
        items = []
        for repl in _gs_store().get_table_dicts("replacements_replacement"):
            if str(repl.get("date") or "") != selected_date.strftime("%Y-%m-%d"):
                continue
            if not _gs_is_teacher_replacement(repl):
                continue
            if as_int(repl.get("replacement_teacher_id")) != teacher.id:
                continue
            lesson = lesson_map.get(as_int(repl.get("lesson_id")) or -1) or {}
            items.append(
                {
                    "lesson_number": as_int(lesson.get("lesson_number")),
                    "class_group": str(lesson.get("class_group") or ""),
                    "subject_name": "",
                    "start_time": str(lesson.get("start_time") or ""),
                    "end_time": str(lesson.get("end_time") or ""),
                    "classroom": str(repl.get("replacement_classroom") or lesson.get("classroom") or ""),
                    "original_teacher": teacher_map.get(as_int(repl.get("original_teacher_id")) or -1, ""),
                    "replacement_teacher": teacher.full_name,
                    "is_special": False,
                }
            )
        items.sort(key=lambda x: (x.get("lesson_number") is None, x.get("lesson_number") or 999, x.get("class_group") or ""))
        return JsonResponse(
            {
                "teacher": teacher.full_name,
                "date": selected_date.strftime("%Y-%m-%d"),
                "items": items,
            }
        )

    items = []
    repl_qs = (
        _teacher_replacements()
        .filter(date=selected_date, replacement_teacher=teacher)
        .select_related("lesson__subject", "original_teacher", "replacement_teacher")
        .order_by("lesson__lesson_number", "lesson__class_group")
    )
    for repl in repl_qs:
        lesson = repl.lesson
        items.append(
            {
                "lesson_number": lesson.lesson_number if lesson else None,
                "class_group": lesson.class_group if lesson else "",
                "subject_name": (lesson.subject.name if lesson and lesson.subject else ""),
                "start_time": lesson.start_time.strftime("%H:%M") if lesson and lesson.start_time else "",
                "end_time": lesson.end_time.strftime("%H:%M") if lesson and lesson.end_time else "",
                "classroom": (repl.replacement_classroom or (lesson.classroom if lesson else "") or ""),
                "original_teacher": repl.original_teacher.full_name if repl.original_teacher else "",
                "replacement_teacher": repl.replacement_teacher.full_name if repl.replacement_teacher else "",
                "is_special": False,
            }
        )

    special_qs = (
        SpecialReplacement.objects.filter(date=selected_date, replacement_teacher=teacher)
        .select_related("original_teacher")
        .order_by("lesson_number", "class_group")
    )
    for sr in special_qs:
        items.append(
            {
                "lesson_number": sr.lesson_number,
                "class_group": sr.class_group or "",
                "subject_name": sr.subject_name or "",
                "start_time": sr.start_time.strftime("%H:%M") if sr.start_time else "",
                "end_time": sr.end_time.strftime("%H:%M") if sr.end_time else "",
                "classroom": sr.classroom or "",
                "original_teacher": sr.original_teacher.full_name if sr.original_teacher else "",
                "replacement_teacher": teacher.full_name,
                "is_special": True,
            }
        )

    items.sort(key=lambda x: (x.get("lesson_number") is None, x.get("lesson_number") or 999, x.get("class_group") or ""))
    return JsonResponse(
        {
            "teacher": teacher.full_name,
            "date": selected_date.strftime("%Y-%m-%d"),
            "items": items,
        }
    )


@login_required
@require_GET
def my_replacement_dates_for_month_api(request):
    if not _is_teacher_user(request.user):
        return HttpResponse("Forbidden", status=403)

    teacher = _resolve_teacher_for_user(request.user)
    if not teacher:
        return JsonResponse({"dates": []})

    year_raw = request.GET.get("year")
    month_raw = request.GET.get("month")
    try:
        year = int(year_raw)
        month = int(month_raw)
        if not (1 <= month <= 12):
            raise ValueError
    except Exception:
        return JsonResponse({"dates": []})

    if _use_gsheets_backend():
        dates = set()
        for r in _gs_store().get_table_dicts("replacements_replacement"):
            if not _gs_is_teacher_replacement(r):
                continue
            if as_int(r.get("replacement_teacher_id")) != teacher.id:
                continue
            ds = str(r.get("date") or "")
            try:
                d = datetime.strptime(ds, "%Y-%m-%d").date()
            except Exception:
                continue
            if d.year == year and d.month == month:
                dates.add(ds)
        for s in _gs_store().get_table_dicts("replacements_special_replacement"):
            if as_int(s.get("replacement_teacher_id")) != teacher.id:
                continue
            ds = str(s.get("date") or "")
            try:
                d = datetime.strptime(ds, "%Y-%m-%d").date()
            except Exception:
                continue
            if d.year == year and d.month == month:
                dates.add(ds)
        return JsonResponse({"dates": sorted(dates)})

    start = datetime(year, month, 1).date()
    if month == 12:
        end = datetime(year + 1, 1, 1).date() - timedelta(days=1)
    else:
        end = datetime(year, month + 1, 1).date() - timedelta(days=1)

    qs_regular = _teacher_replacements().filter(
        replacement_teacher=teacher,
        date__range=(start, end),
    ).values_list("date", flat=True).distinct()
    qs_special = SpecialReplacement.objects.filter(
        replacement_teacher=teacher,
        date__range=(start, end),
    ).values_list("date", flat=True).distinct()

    dates = sorted({d.strftime("%Y-%m-%d") for d in list(qs_regular) + list(qs_special)})
    return JsonResponse({"dates": dates})


@login_required
def get_saved_replacements(request):
    from .services.replacements_heavy_service import get_saved_replacements_service
    return get_saved_replacements_service(request)

@login_required
def teachers_overview_view(request):
    """Страница "Учителя": ФИО + предметы/уроки, которые ведёт учитель (через запятую)."""
                                                                  
    if not (request.user.is_superuser or getattr(request.user, 'can_teachers', False) or getattr(request.user, 'can_editor', False)):
        return HttpResponse("Forbidden", status=403)

    if _use_gsheets_backend():
        teachers_rows = sorted(
            _gs_store().get_table_dicts("replacements_teacher"),
            key=lambda t: str(t.get("full_name") or "").casefold(),
        )
        active_lessons = _gs_active_lessons_rows()
        subject_map = _gs_subject_map()
        subj_by_teacher: dict[int, set[str]] = {}
        for l in active_lessons:
            tid = as_int(l.get("teacher_id"))
            sid = as_int(l.get("subject_id_subject"))
            if tid is None or sid is None:
                continue
            sname = subject_map.get(sid, "")
            if sname:
                subj_by_teacher.setdefault(tid, set()).add(sname)

        rows = []
        for t in teachers_rows:
            tid = as_int(t.get("teacher_id"))
            if tid is None:
                continue
            teacher_obj = SimpleNamespace(id=tid, full_name=str(t.get("full_name") or ""))
            subj_names = sorted(subj_by_teacher.get(tid, set()))
            subjects_str = ", ".join(subj_names) if subj_names else "—"
            rows.append({"teacher": teacher_obj, "subjects": subjects_str})
        return render(request, 'teachers.html', {'rows': rows})

    teachers = Teacher.objects.all().order_by('full_name').prefetch_related(Prefetch('lesson_set', queryset=_active_lessons().select_related('subject')))

    rows = []
    for t in teachers:
        subj_names = sorted({(l.subject.name if l.subject else '') for l in t.lesson_set.all() if l.subject})
        subjects_str = ", ".join(subj_names) if subj_names else "—"
        rows.append({
            'teacher': t,
            'subjects': subjects_str,
        })

    return render(request, 'teachers.html', {
        'rows': rows,
    })


@login_required
def specializations_view(request):
                                                                             
    if not (request.user.is_superuser or getattr(request.user, 'can_editor', False)):
        return HttpResponse("Forbidden", status=403)

    if _use_gsheets_backend():
        teachers_rows = sorted(
            _gs_store().get_table_dicts("replacements_teacher"),
            key=lambda t: str(t.get("full_name") or "").casefold(),
        )
        teachers = []
        specializations: dict[int, set[str]] = {}
        for t in teachers_rows:
            tid = as_int(t.get("teacher_id"))
            if tid is None:
                continue
            teachers.append(SimpleNamespace(id=tid, full_name=str(t.get("full_name") or "")))
            spec = str(t.get("specialization") or "").strip()
            specializations[tid] = {s.strip() for s in spec.split(",") if s.strip()} if spec else set()

        by_name_min_id: dict[str, int] = {}
        for s in _gs_store().get_table_dicts("replacements_subject"):
            sid = as_int(s.get("id_subject"))
            name = str(s.get("name") or "").strip()
            if sid is None or not name:
                continue
            if name not in by_name_min_id or sid < by_name_min_id[name]:
                by_name_min_id[name] = sid
        subjects = [
            SimpleNamespace(id_subject=sid, name=name)
            for name, sid in sorted(by_name_min_id.items(), key=lambda x: x[0].casefold())
        ]
        return render(request, 'specializations.html', {
            'teachers': teachers,
            'subjects': subjects,
            'specializations': specializations,
        })

    teachers = Teacher.objects.all().prefetch_related('lesson_set')

    unique_subjects = Subject.objects.values('name').annotate(
        first_id=Min('id_subject')
    ).order_by('name')

    subject_ids = [s['first_id'] for s in unique_subjects]
    subjects = Subject.objects.filter(id_subject__in=subject_ids).order_by('name')

    specializations = {
        t.id: set(t.specialization.split(',')) if t.specialization else set()
        for t in teachers
    }

    return render(request, 'specializations.html', {
        'teachers': teachers,
        'subjects': subjects,
        'specializations': specializations
    })


@login_required
def update_specialization(request):
    guest_forbidden = _deny_guest_write_json(request)
    if guest_forbidden:
        return guest_forbidden

                                                                               
    if not (request.user.is_superuser or getattr(request.user, 'can_editor', False)):
        return JsonResponse({'error': 'Forbidden'}, status=403)

    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            if _use_gsheets_backend():
                teacher_id = as_int(data.get('teacher_id'))
                subject_id = str(data.get('subject_id'))
                if teacher_id is None or not subject_id:
                    return JsonResponse({'error': 'Некорректные данные'}, status=400)
                rows = _gs_store().get_table_dicts("replacements_teacher")
                idx = None
                for i, r in enumerate(rows):
                    if as_int(r.get("teacher_id")) == teacher_id:
                        idx = i
                        break
                if idx is None:
                    return JsonResponse({'error': 'Учитель не найден'}, status=404)
                teacher_row = rows[idx]
                current = {
                    s.strip() for s in str(teacher_row.get("specialization") or "").split(",") if s.strip()
                }
                if data.get('action') == 'add':
                    current.add(subject_id)
                else:
                    current.discard(subject_id)
                teacher_row["specialization"] = ",".join(sorted(current)) if current else ""
                rows[idx] = teacher_row
                _gs_store().replace_table_dicts("replacements_teacher", rows)
                log_activity(request, "specialization_update", {
                    "teacher_id": teacher_id,
                    "action": data.get('action'),
                    "subject_id": subject_id,
                    "backend": "gsheets",
                })
                return JsonResponse({'status': 'success'})

            teacher = Teacher.objects.get(id=data['teacher_id'])
            subject_id = str(data['subject_id'])

            current = set(teacher.specialization.split(',')) if teacher.specialization else set()

            if data['action'] == 'add':
                current.add(subject_id)
            else:
                current.discard(subject_id)

            teacher.specialization = ','.join(sorted(current)) if current else None
            teacher.save()

            log_activity(request, "specialization_update", {
                "teacher_id": teacher.id,
                "action": data.get('action'),
                "subject_id": subject_id,
            })

            return JsonResponse({'status': 'success'})
        except Exception:
            return _api_error("Некорректные данные запроса", status=400, code="bad_request")
    return JsonResponse({'error': 'Invalid method'}, status=405)


@login_required
@require_POST
def add_teacher_api(request):
    """Create a new teacher from the specialization editor page."""
                                                                           
    guest_forbidden = _deny_guest_write_json(request)
    if guest_forbidden:
        return guest_forbidden

    if not (request.user.is_superuser or getattr(request.user, 'can_editor', False)):
        return JsonResponse({'error': 'Forbidden'}, status=403)
    try:
        data = json.loads(request.body or b"{}")
    except Exception:
        data = {}

    full_name = (data.get('full_name') or '').strip()
    full_name = re.sub(r'\s+', ' ', full_name)

    if not full_name:
        return JsonResponse({'error': 'ФИО не может быть пустым'}, status=400)

    if _use_gsheets_backend():
        rows = _gs_store().get_table_dicts("replacements_teacher")
        exists = any(str(r.get("full_name") or "").strip().casefold() == full_name.casefold() for r in rows)
        if exists:
            return JsonResponse({'error': 'Такой учитель уже существует'}, status=409)
        new_id = _gs_next_id(rows, "teacher_id")
        rows.append({
            "teacher_id": new_id,
            "full_name": full_name,
            "specialization": "",
            "hours_per_week": 0,
        })
        _gs_store().replace_table_dicts("replacements_teacher", rows)
        log_activity(request, "teacher_add", {"teacher_id": new_id, "full_name": full_name, "backend": "gsheets"})
        return JsonResponse({'status': 'success', 'teacher': {'id': new_id, 'full_name': full_name}})

    try:
        with transaction.atomic():
            teacher = Teacher.objects.create(full_name=full_name)
    except IntegrityError:
        return JsonResponse({'error': 'Такой учитель уже существует'}, status=409)
    except Exception:
        return _api_error("Некорректные данные запроса", status=400, code="bad_request")

    log_activity(request, "teacher_add", {"teacher_id": teacher.id, "full_name": teacher.full_name})
    return JsonResponse({'status': 'success', 'teacher': {'id': teacher.id, 'full_name': teacher.full_name}})


@login_required
@require_POST
def delete_teacher_api(request, teacher_id: int):
    """Delete teacher added by mistake.

    Safety: do not allow deleting teachers that already have lessons or replacements.
    """
                                                                         
    guest_forbidden = _deny_guest_write_json(request)
    if guest_forbidden:
        return guest_forbidden

    if not (request.user.is_superuser or getattr(request.user, 'can_editor', False)):
        return JsonResponse({'error': 'Forbidden'}, status=403)

    if _use_gsheets_backend():
        teacher_rows = _gs_store().get_table_dicts("replacements_teacher")
        teacher_row = next((r for r in teacher_rows if as_int(r.get("teacher_id")) == int(teacher_id)), None)
        if teacher_row is None:
            return JsonResponse({'error': 'Учитель не найден'}, status=404)
        lessons_cnt = sum(1 for l in _gs_store().get_table_dicts("replacements_lesson") if as_int(l.get("teacher_id")) == int(teacher_id))
        repl_cnt = sum(
            1
            for r in _gs_store().get_table_dicts("replacements_replacement")
            if as_int(r.get("original_teacher_id")) == int(teacher_id) or as_int(r.get("replacement_teacher_id")) == int(teacher_id)
        )
        if lessons_cnt > 0 or repl_cnt > 0:
            return JsonResponse(
                {
                    'error': 'Нельзя удалить учителя: у него уже есть уроки/замещения',
                    'lessons': lessons_cnt,
                    'replacements': repl_cnt,
                },
                status=409
            )
        updated = [r for r in teacher_rows if as_int(r.get("teacher_id")) != int(teacher_id)]
        _gs_store().replace_table_dicts("replacements_teacher", updated)
        log_activity(request, "teacher_delete", {"teacher_id": teacher_id, "full_name": str(teacher_row.get("full_name") or ""), "backend": "gsheets"})
        return JsonResponse({'status': 'success'})

    teacher = get_object_or_404(Teacher, id=teacher_id)

    lessons_cnt = Lesson.objects.filter(teacher=teacher).count()
    repl_cnt = Replacement.objects.filter(Q(original_teacher=teacher) | Q(replacement_teacher=teacher)).count()

    if lessons_cnt > 0 or repl_cnt > 0:
        return JsonResponse(
            {
                'error': 'Нельзя удалить учителя: у него уже есть уроки/замещения',
                'lessons': lessons_cnt,
                'replacements': repl_cnt,
            },
            status=409
        )

    try:
        with transaction.atomic():
            teacher.delete()
    except Exception:
        return _api_error("Некорректные данные запроса", status=400, code="bad_request")

    log_activity(request, "teacher_delete", {"teacher_id": teacher_id, "full_name": teacher.full_name})
    return JsonResponse({'status': 'success'})


from django.http import JsonResponse
from django.views.decorators.http import require_GET
from .models import Lesson, Subject, Teacher


@require_GET
def get_lesson_by_id(request, lesson_id):
    if _use_gsheets_backend():
        lesson = None
        for l in _gs_active_lessons_rows():
            if as_int(l.get("lesson_id")) == int(lesson_id):
                lesson = l
                break
        if not lesson:
            return JsonResponse({"error": f"Урок с ID {lesson_id} не найден"}, status=404)
        subject_map = _gs_subject_map()
        teacher_map = _gs_teacher_map()
        eff_shift = _effective_shift_for_class(str(lesson.get("class_group") or ""), as_int(lesson.get("shift")))
        eff_start, eff_end = _effective_times_for_lesson(
            str(lesson.get("class_group") or ""),
            as_int(lesson.get("lesson_number")) or 0,
            eff_shift,
            _gs_parse_time(lesson.get("start_time")),
            _gs_parse_time(lesson.get("end_time")),
        )
        sid = as_int(lesson.get("subject_id_subject"))
        tid = as_int(lesson.get("teacher_id"))
        return JsonResponse({
            "lesson_id": as_int(lesson.get("lesson_id")),
            "subject": subject_map.get(sid, "Не указано") if sid is not None else "Не указано",
            "teacher": teacher_map.get(tid, "Не указан") if tid is not None else "Не указан",
            "lesson_number": as_int(lesson.get("lesson_number")),
            "start_time": eff_start.strftime("%H:%M") if eff_start else "--:--",
            "end_time": eff_end.strftime("%H:%M") if eff_end else "--:--",
            "classroom": str(lesson.get("classroom") or ""),
            "class_group": str(lesson.get("class_group") or ""),
            "shift": eff_shift,
            "day_of_week": str(lesson.get("day_of_week") or ""),
        })
    try:
        lesson = Lesson.objects.select_related("subject", "teacher").get(id=lesson_id)

        eff_shift = _effective_shift_for_class(lesson.class_group, lesson.shift)
        eff_start, eff_end = _effective_times_for_lesson(
            lesson.class_group,
            lesson.lesson_number,
            eff_shift,
            lesson.start_time,
            lesson.end_time,
        )
        return JsonResponse({
            "lesson_id": lesson.id,
            "subject": lesson.subject.name if lesson.subject else "Не указано",
            "teacher": lesson.teacher.full_name if lesson.teacher else "Не указан",
            "lesson_number": lesson.lesson_number,
            "start_time": eff_start.strftime("%H:%M") if eff_start else "--:--",
            "end_time": eff_end.strftime("%H:%M") if eff_end else "--:--",
            "classroom": lesson.classroom,
            "class_group": lesson.class_group,
            "shift": eff_shift,
            "day_of_week": lesson.day_of_week,
        })
    except Lesson.DoesNotExist:
        return JsonResponse({"error": f"Урок с ID {lesson_id} не найден"}, status=404)


@login_required
@require_GET
def teacher_details(request, teacher_id):
    if not _can_calendar_read(request.user):
        return HttpResponse("Forbidden", status=403)
    if _use_gsheets_backend():
        for t in _gs_store().get_table_dicts("replacements_teacher"):
            if as_int(t.get("teacher_id")) == int(teacher_id):
                return JsonResponse({"name": str(t.get("full_name") or "")})
        return JsonResponse({"error": "Teacher not found"}, status=404)
    try:
        teacher = Teacher.objects.get(id=teacher_id)
        return JsonResponse({"name": teacher.full_name})
    except Teacher.DoesNotExist:
        return JsonResponse({"error": "Teacher not found"}, status=404)


@login_required
@require_GET
def get_lessons(request, teacher_id, day):
    if not _can_calendar_read(request.user):
        return HttpResponse("Forbidden", status=403)
    if _use_gsheets_backend():
        subject_map = _gs_subject_map()
        lessons = [
            l for l in _gs_active_lessons_rows()
            if as_int(l.get("teacher_id")) == int(teacher_id) and str(l.get("day_of_week") or "") == str(day)
        ]
        lesson_data = []
        for lesson in lessons:
            eff_shift = _effective_shift_for_class(str(lesson.get("class_group") or ""), as_int(lesson.get("shift")))
            eff_start, eff_end = _effective_times_for_lesson(
                str(lesson.get("class_group") or ""),
                as_int(lesson.get("lesson_number")) or 0,
                eff_shift,
                _gs_parse_time(lesson.get("start_time")),
                _gs_parse_time(lesson.get("end_time")),
            )
            sid = as_int(lesson.get("subject_id_subject"))
            lesson_data.append({
                "id": as_int(lesson.get("lesson_id")),
                "number": as_int(lesson.get("lesson_number")),
                "start": eff_start.strftime("%H:%M") if eff_start else "--:--",
                "end": eff_end.strftime("%H:%M") if eff_end else "--:--",
                "subject": subject_map.get(sid, "") if sid is not None else "",
                "class": str(lesson.get("class_group") or ""),
                "shift": eff_shift,
            })
        lesson_data.sort(key=lambda x: ((x.get("number") or 0), (x.get("id") or 0)))
        return JsonResponse({"lessons": lesson_data})
    lessons = _active_lessons().filter(teacher_id=teacher_id, day_of_week=day)
    lesson_data = []
    for lesson in lessons:
        eff_shift = _effective_shift_for_class(lesson.class_group, lesson.shift)
        eff_start, eff_end = _effective_times_for_lesson(
            lesson.class_group,
            lesson.lesson_number,
            eff_shift,
            lesson.start_time,
            lesson.end_time,
        )
        lesson_data.append({
            "id": lesson.id,
            "number": lesson.lesson_number,
            "start": eff_start.strftime("%H:%M") if eff_start else "--:--",
            "end": eff_end.strftime("%H:%M") if eff_end else "--:--",
            "subject": str(lesson.subject),
            "class": lesson.class_group,
            "shift": eff_shift,
        })
    return JsonResponse({"lessons": lesson_data})


def get_teacher(request, teacher_id):
    if _use_gsheets_backend():
        name = _gs_teacher_map().get(int(teacher_id))
        if name is None:
            return JsonResponse({"error": "Teacher not found"}, status=404)
        return JsonResponse({"name": name})
    try:
        teacher = Teacher.objects.get(id=teacher_id)
        return JsonResponse({"name": teacher.full_name})
    except Teacher.DoesNotExist:
        return JsonResponse({"error": "Teacher not found"}, status=404)


@login_required
@require_GET
def get_lessons_by_id(request, lesson_id):
    if not _can_calendar_read(request.user):
        return HttpResponse("Forbidden", status=403)
    if _use_gsheets_backend():
        lesson = None
        for l in _gs_active_lessons_rows():
            if as_int(l.get("lesson_id")) == int(lesson_id):
                lesson = l
                break
        if not lesson:
            return JsonResponse({"error": f"Урок с ID {lesson_id} не найден"}, status=404)
        subject_map = _gs_subject_map()
        teacher_map = _gs_teacher_map()
        eff_shift = _effective_shift_for_class(str(lesson.get("class_group") or ""), as_int(lesson.get("shift")))
        eff_start, eff_end = _effective_times_for_lesson(
            str(lesson.get("class_group") or ""),
            as_int(lesson.get("lesson_number")) or 0,
            eff_shift,
            _gs_parse_time(lesson.get("start_time")),
            _gs_parse_time(lesson.get("end_time")),
        )
        sid = as_int(lesson.get("subject_id_subject"))
        tid = as_int(lesson.get("teacher_id"))
        return JsonResponse({
            "id": as_int(lesson.get("lesson_id")),
            "subject": subject_map.get(sid, "Не указано") if sid is not None else "Не указано",
            "teacher": teacher_map.get(tid, "Не указан") if tid is not None else "Не указан",
            "number": as_int(lesson.get("lesson_number")),
            "class": str(lesson.get("class_group") or ""),
            "room": str(lesson.get("classroom") or ""),
            "start": eff_start.strftime("%H:%M") if eff_start else "--:--",
            "end": eff_end.strftime("%H:%M") if eff_end else "--:--",
            "shift": eff_shift,
            "day_of_week": str(lesson.get("day_of_week") or ""),
        })
    try:
        lesson = Lesson.objects.select_related("subject", "teacher").get(id=lesson_id)

        eff_shift = _effective_shift_for_class(lesson.class_group, lesson.shift)
        eff_start, eff_end = _effective_times_for_lesson(
            lesson.class_group,
            lesson.lesson_number,
            eff_shift,
            lesson.start_time,
            lesson.end_time,
        )
        return JsonResponse({
            "id": lesson.id,
            "subject": lesson.subject.name if lesson.subject else "Не указано",
            "teacher": lesson.teacher.full_name if lesson.teacher else "Не указан",
            "number": lesson.lesson_number,
            "class": lesson.class_group,
            "room": lesson.classroom,
            "start": eff_start.strftime("%H:%M") if eff_start else "--:--",
            "end": eff_end.strftime("%H:%M") if eff_end else "--:--",
            "shift": eff_shift,
            "day_of_week": lesson.day_of_week,
        })
    except Lesson.DoesNotExist:
        return JsonResponse({"error": f"Урок с ID {lesson_id} не найден"}, status=404)
    except Exception:
        return _api_error("Внутренняя ошибка сервера", status=500, code="internal_error")


@login_required
def check_replacements_for_date(request):
    if not (request.user.is_superuser or _is_guest_user(request.user) or getattr(request.user, 'can_calendar', False)):
        return HttpResponse("Forbidden", status=403)

    date = request.GET.get('date')
    if not date:
        return JsonResponse({'error': 'Дата не указана'}, status=400)

    if _use_gsheets_backend():
        exists_regular = any(
            str(r.get("date") or "") == date and _gs_is_teacher_replacement(r)
            for r in _gs_store().get_table_dicts("replacements_replacement")
        )
        exists_special = any(
            str(r.get("date") or "") == date
            for r in _gs_store().get_table_dicts("replacements_special_replacement")
        )
        return JsonResponse({'exists': bool(exists_regular or exists_special)})

    exists = _teacher_replacements().filter(date=date).exists() or SpecialReplacement.objects.filter(date=date).exists()
    return JsonResponse({'exists': exists})


@login_required
def delete_replacements_for_date(request):
    guest_forbidden = _deny_guest_write_json(request)
    if guest_forbidden:
        return guest_forbidden

    if request.method != 'POST':
        return JsonResponse({'error': 'Метод не разрешен'}, status=405)

    try:
        data = json.loads(request.body)
        date = data.get('date')
        if not date:
            return JsonResponse({'error': 'Не указана дата'}, status=400)

        if _use_gsheets_backend():
            store = _gs_store()
            replacements_rows = store.get_table_dicts("replacements_replacement")
            special_rows = store.get_table_dicts("replacements_special_replacement")

            before_repl = len(replacements_rows)
            before_special = len(special_rows)

            kept_repl = [
                r for r in replacements_rows
                if not (str(r.get("date") or "") == date and _gs_is_teacher_replacement(r))
            ]
            kept_special = [r for r in special_rows if str(r.get("date") or "") != date]

            store.replace_table_dicts("replacements_replacement", kept_repl)
            store.replace_table_dicts("replacements_special_replacement", kept_special)

            deleted_count = (before_repl - len(kept_repl)) + (before_special - len(kept_special))
            return JsonResponse({'status': 'success', 'deleted_count': deleted_count})

        deleted, _ = _teacher_replacements().filter(date=date).delete()
        deleted_special, _ = SpecialReplacement.objects.filter(date=date).delete()
        return JsonResponse({'status': 'success', 'deleted_count': deleted + deleted_special})

    except Exception:
        return _api_error("Внутренняя ошибка сервера", status=500, code="internal_error")


@login_required
def replacement_dates_for_month(request):
    """Return a list of YYYY-MM-DD dates that have saved replacements.

    Used by the calendar UI to highlight dates with existing replacements.
    Query params:
      - year: int
      - month: int (1-12)
    """
    if not (request.user.is_superuser or _is_guest_user(request.user) or getattr(request.user, 'can_calendar', False)):
        return HttpResponse("Forbidden", status=403)

    try:
        year = int(request.GET.get('year'))
        month = int(request.GET.get('month'))
        if not (1 <= month <= 12):
            raise ValueError("month must be 1..12")

        last_day = _pycalendar.monthrange(year, month)[1]
        start = datetime(year, month, 1).date()
        end = datetime(year, month, last_day).date()

        if _use_gsheets_backend():
            dates_set: set[str] = set()
            for r in _gs_store().get_table_dicts("replacements_replacement"):
                if not _gs_is_teacher_replacement(r):
                    continue
                d_raw = str(r.get("date") or "")
                try:
                    d = datetime.strptime(d_raw, "%Y-%m-%d").date()
                except Exception:
                    continue
                if start <= d <= end:
                    dates_set.add(d_raw)
            for r in _gs_store().get_table_dicts("replacements_special_replacement"):
                d_raw = str(r.get("date") or "")
                try:
                    d = datetime.strptime(d_raw, "%Y-%m-%d").date()
                except Exception:
                    continue
                if start <= d <= end:
                    dates_set.add(d_raw)
            return JsonResponse({'dates': sorted(dates_set)})

        qs_repl = _teacher_replacements().filter(date__range=(start, end)).values_list('date', flat=True).distinct()
        qs_special = SpecialReplacement.objects.filter(date__range=(start, end)).values_list('date', flat=True).distinct()
        dates = sorted({d.strftime('%Y-%m-%d') for d in list(qs_repl) + list(qs_special)})
        return JsonResponse({'dates': dates})
    except Exception:
        return _api_error("Некорректные данные запроса", status=400, code="bad_request")


from datetime import datetime, timedelta
from .models import Lesson, Replacement, Teacher


@login_required
@require_GET
def teacher_hours(request, teacher_id):
    if not _can_calendar_read(request.user):
        return HttpResponse("Forbidden", status=403)
    try:
                                    
        selected_date = request.GET.get('date')
        if not selected_date:
            return JsonResponse({'error': 'Missing date parameter'}, status=400)

        selected_date = datetime.strptime(selected_date, "%Y-%m-%d").date()

        if _use_gsheets_backend():
            tid = int(teacher_id)
            teacher_exists = any(as_int(t.get("teacher_id")) == tid for t in _gs_store().get_table_dicts("replacements_teacher"))
            if not teacher_exists:
                return JsonResponse({'error': 'Teacher not found'}, status=404)
            start_of_week = selected_date - timedelta(days=selected_date.weekday())
            end_of_week = start_of_week + timedelta(days=6)
            base_lessons_count = sum(
                1
                for l in _gs_active_lessons_rows()
                if as_int(l.get("teacher_id")) == tid
            )
            replacement_count = 0
            for r in _gs_store().get_table_dicts("replacements_replacement"):
                if not _gs_is_teacher_replacement(r):
                    continue
                if as_int(r.get("replacement_teacher_id")) != tid:
                    continue
                d_raw = str(r.get("date") or "")
                try:
                    d = datetime.strptime(d_raw, "%Y-%m-%d").date()
                except Exception:
                    continue
                if start_of_week <= d <= end_of_week:
                    replacement_count += 1
            return JsonResponse({'base_hours': base_lessons_count, 'replacement_hours': replacement_count})

        teacher = Teacher.objects.get(id=teacher_id)

                                                          
        start_of_week = selected_date - timedelta(days=selected_date.weekday())               
        end_of_week = start_of_week + timedelta(days=6)               

                                      
        base_lessons_count = _active_lessons().filter(teacher=teacher).count()

                                      
        replacement_count = _teacher_replacements().filter(
            replacement_teacher=teacher,
            date__range=(start_of_week, end_of_week)
        ).count()

        return JsonResponse({
            'base_hours': base_lessons_count,
            'replacement_hours': replacement_count
        })

    except Teacher.DoesNotExist:
        return JsonResponse({'error': 'Teacher not found'}, status=404)
    except Exception:
        return _api_error("Внутренняя ошибка сервера", status=500, code="internal_error")


@login_required
@require_GET
def backend_health_api(request):
    """Health and consistency checks for current data backend."""
    if not (request.user.is_superuser or getattr(request.user, 'can_logs', False)):
        return HttpResponse("Forbidden", status=403)

    try:
        if _use_gsheets_backend():
            store = _gs_store()
            teachers = store.get_table_dicts("replacements_teacher")
            subjects = store.get_table_dicts("replacements_subject")
            lessons = store.get_table_dicts("replacements_lesson")
            repls = store.get_table_dicts("replacements_replacement")
            specials = store.get_table_dicts("replacements_special_replacement")
            schedule = store.get_table_dicts("replacements_class_schedule")

            issues = []
            warnings = []

            def find_duplicate_ids(rows: list[dict], id_key: str, table_name: str):
                seen = set()
                dups = set()
                for r in rows:
                    rid = as_int(r.get(id_key))
                    if rid is None:
                        continue
                    if rid in seen:
                        dups.add(rid)
                    else:
                        seen.add(rid)
                if dups:
                    issues.append(f"{table_name}: duplicate IDs in `{id_key}`: {sorted(list(dups))[:20]}")

            find_duplicate_ids(teachers, "teacher_id", "replacements_teacher")
            find_duplicate_ids(subjects, "id_subject", "replacements_subject")
            find_duplicate_ids(lessons, "lesson_id", "replacements_lesson")
            find_duplicate_ids(repls, "id", "replacements_replacement")
            find_duplicate_ids(specials, "id", "replacements_special_replacement")
            find_duplicate_ids(schedule, "id", "replacements_class_schedule")

            teacher_ids = {as_int(t.get("teacher_id")) for t in teachers if as_int(t.get("teacher_id")) is not None}
            subject_ids = {as_int(s.get("id_subject")) for s in subjects if as_int(s.get("id_subject")) is not None}
            lesson_ids = {as_int(l.get("lesson_id")) for l in lessons if as_int(l.get("lesson_id")) is not None}

            bad_lesson_teachers = [
                as_int(l.get("lesson_id"))
                for l in lessons
                if as_int(l.get("teacher_id")) not in teacher_ids
            ]
            if bad_lesson_teachers:
                issues.append(f"replacements_lesson: missing teacher refs for lessons: {bad_lesson_teachers[:30]}")

            bad_lesson_subjects = [
                as_int(l.get("lesson_id"))
                for l in lessons
                if as_int(l.get("subject_id_subject")) not in subject_ids
            ]
            if bad_lesson_subjects:
                issues.append(f"replacements_lesson: missing subject refs for lessons: {bad_lesson_subjects[:30]}")

            bad_repl_lessons = [
                as_int(r.get("id"))
                for r in repls
                if as_int(r.get("lesson_id")) not in lesson_ids
            ]
            if bad_repl_lessons:
                issues.append(f"replacements_replacement: missing lesson refs: {bad_repl_lessons[:30]}")

            bad_repl_teachers = [
                as_int(r.get("id"))
                for r in repls
                if as_int(r.get("original_teacher_id")) not in teacher_ids
                or as_int(r.get("replacement_teacher_id")) not in teacher_ids
            ]
            if bad_repl_teachers:
                issues.append(f"replacements_replacement: missing teacher refs: {bad_repl_teachers[:30]}")

            bad_special_refs = []
            for s in specials:
                sid = as_int(s.get("id"))
                lt = as_int(s.get("lesson_id"))
                rt = as_int(s.get("replacement_teacher_id"))
                ot = as_int(s.get("original_teacher_id"))
                if lt is not None and lt not in lesson_ids:
                    bad_special_refs.append(sid)
                    continue
                if rt is not None and rt not in teacher_ids:
                    bad_special_refs.append(sid)
                    continue
                if ot is not None and ot not in teacher_ids:
                    bad_special_refs.append(sid)
                    continue
            if bad_special_refs:
                issues.append(f"replacements_special_replacement: broken refs: {bad_special_refs[:30]}")

            bad_lesson_time_ids = []
            for l in lessons:
                lid = as_int(l.get("lesson_id"))
                st = _gs_parse_time(l.get("start_time"))
                en = _gs_parse_time(l.get("end_time"))
                if not st or not en or st >= en:
                    bad_lesson_time_ids.append(lid)
            if bad_lesson_time_ids:
                warnings.append(f"replacements_lesson: invalid times for lesson IDs: {bad_lesson_time_ids[:30]}")

            bad_schedule_time_ids = []
            for r in schedule:
                rid = as_int(r.get("id"))
                st = _gs_parse_time(r.get("start_time"))
                en = _gs_parse_time(r.get("end_time"))
                if not st or not en or st >= en:
                    bad_schedule_time_ids.append(rid)
            if bad_schedule_time_ids:
                warnings.append(f"replacements_class_schedule: invalid times for IDs: {bad_schedule_time_ids[:30]}")

            ok = len(issues) == 0
            return JsonResponse({
                "ok": ok,
                "backend": "gsheets",
                "checked_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "counts": {
                    "teachers": len(teachers),
                    "subjects": len(subjects),
                    "lessons": len(lessons),
                    "replacements": len(repls),
                    "special_replacements": len(specials),
                    "class_schedule": len(schedule),
                },
                "issues": issues,
                "warnings": warnings,
            })

                     
        teacher_ids = set(Teacher.objects.values_list("id", flat=True))
        subject_ids = set(Subject.objects.values_list("id_subject", flat=True))
        lesson_ids = set(Lesson.objects.values_list("id", flat=True))
        issues = []
        warnings = []

        missing_lesson_teacher = Lesson.objects.exclude(teacher_id__in=teacher_ids).values_list("id", flat=True)[:30]
        if missing_lesson_teacher:
            issues.append(f"replacements_lesson: missing teacher refs: {list(missing_lesson_teacher)}")

        missing_lesson_subject = Lesson.objects.exclude(subject_id__in=subject_ids).values_list("id", flat=True)[:30]
        if missing_lesson_subject:
            issues.append(f"replacements_lesson: missing subject refs: {list(missing_lesson_subject)}")

        missing_repl_lesson = Replacement.objects.exclude(lesson_id__in=lesson_ids).values_list("id", flat=True)[:30]
        if missing_repl_lesson:
            issues.append(f"replacements_replacement: missing lesson refs: {list(missing_repl_lesson)}")

        bad_lesson_time_ids = [
            l.id for l in Lesson.objects.all()[:10000]
            if not l.start_time or not l.end_time or l.start_time >= l.end_time
        ]
        if bad_lesson_time_ids:
            warnings.append(f"replacements_lesson: invalid times for lesson IDs: {bad_lesson_time_ids[:30]}")

        ok = len(issues) == 0
        return JsonResponse({
            "ok": ok,
            "backend": "sqlite",
            "checked_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "counts": {
                "teachers": Teacher.objects.count(),
                "subjects": Subject.objects.count(),
                "lessons": Lesson.objects.count(),
                "replacements": Replacement.objects.count(),
                "special_replacements": SpecialReplacement.objects.count(),
                "class_schedule": ClassSchedule.objects.count(),
            },
            "issues": issues,
            "warnings": warnings,
        })
    except (DatabaseError, ValueError, TypeError):
        return JsonResponse({
            "ok": False,
            "error": {
                "code": "internal_error",
                "message": "Не удалось выполнить проверку backend",
            },
            "backend": "gsheets" if _use_gsheets_backend() else "sqlite",
            "checked_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }, status=500)


def _stats_date_range(mode: str, date_str: str | None, date_from: str | None, date_to: str | None, month_str: str | None):
    """Resolve stats range and human-readable label from query params."""
    today = datetime.today().date()
    mode_norm = (mode or "day").strip().lower()

    if mode_norm == "period":
        if not (date_from and date_to):
            raise ValueError("Для режима period нужны date_from и date_to")
        start = datetime.strptime(date_from, "%Y-%m-%d").date()
        end = datetime.strptime(date_to, "%Y-%m-%d").date()
        if end < start:
            raise ValueError("date_to не может быть раньше date_from")
        label = f"Период: {start.strftime('%d.%m.%Y')} - {end.strftime('%d.%m.%Y')}"
        return mode_norm, start, end, label

    if mode_norm == "month":
        if month_str:
            selected = datetime.strptime(f"{month_str}-01", "%Y-%m-%d").date()
        elif date_str:
            selected = datetime.strptime(date_str, "%Y-%m-%d").date()
        else:
            selected = today
        start = selected.replace(day=1)
        if start.month == 12:
            next_month = start.replace(year=start.year + 1, month=1, day=1)
        else:
            next_month = start.replace(month=start.month + 1, day=1)
        end = next_month - timedelta(days=1)
        label = f"Месяц: {start.strftime('%m.%Y')}"
        return mode_norm, start, end, label

                   
    selected = datetime.strptime(date_str, "%Y-%m-%d").date() if date_str else today
    label = f"День: {selected.strftime('%d.%m.%Y')}"
    return "day", selected, selected, label


@login_required
def statistics_view(request):
    if not (request.user.is_superuser or _is_guest_user(request.user) or _is_teacher_user(request.user) or getattr(request.user, 'can_calendar', False) or getattr(request.user, 'can_logs', False)):
        return HttpResponse("Forbidden", status=403)
    if _is_teacher_user(request.user):
        teacher = _resolve_teacher_for_user(request.user)
        return render(
            request,
            "teacher_statistics.html",
            {
                "teacher_name": teacher.full_name if teacher else (request.user.full_name or request.user.username),
                "teacher_found": bool(teacher),
            },
        )
    return render(request, "statistics.html", {})


@login_required
@require_GET
def replacement_statistics_api(request):
    if not (request.user.is_superuser or _is_guest_user(request.user) or _is_teacher_user(request.user) or getattr(request.user, 'can_calendar', False) or getattr(request.user, 'can_logs', False)):
        return HttpResponse("Forbidden", status=403)

    mode = request.GET.get("mode", "day")
    date_str = request.GET.get("date")
    date_from = request.GET.get("date_from")
    date_to = request.GET.get("date_to")
    month_str = request.GET.get("month")

    try:
        mode_norm, start, end, range_label = _stats_date_range(mode, date_str, date_from, date_to, month_str)
    except Exception:
        return _api_error("Некорректные данные запроса", status=400, code="bad_request")

    scoped_teacher = _resolve_teacher_for_user(request.user) if _is_teacher_user(request.user) else None
    if _is_teacher_user(request.user) and not scoped_teacher:
        return JsonResponse({"error": "Учитель для текущего пользователя не найден"}, status=400)

    if _use_gsheets_backend():
        scope_suffix = f":teacher:{scoped_teacher.id}" if scoped_teacher else ":all"
        cache_key = (
            f"replacements:stats:gsheets:{mode_norm}:"
            f"{start.strftime('%Y-%m-%d')}:{end.strftime('%Y-%m-%d')}{scope_suffix}"
        )
        cached_payload = cache.get(cache_key)
        if cached_payload is not None:
            return JsonResponse(cached_payload)

        teacher_map = _gs_teacher_map()

        replacements_rows = []
        for r in _gs_store().get_table_dicts("replacements_replacement"):
            if not _gs_is_teacher_replacement(r):
                continue
            if as_bool(r.get("ignore_in_reports")):
                continue
            d_raw = str(r.get("date") or "")
            try:
                d = datetime.strptime(d_raw, "%Y-%m-%d").date()
            except Exception:
                continue
            if start <= d <= end:
                if scoped_teacher and as_int(r.get("replacement_teacher_id")) != scoped_teacher.id:
                    continue
                replacements_rows.append(r)

        special_rows = []
        for s in _gs_store().get_table_dicts("replacements_special_replacement"):
            d_raw = str(s.get("date") or "")
            try:
                d = datetime.strptime(d_raw, "%Y-%m-%d").date()
            except Exception:
                continue
            if start <= d <= end:
                if scoped_teacher and as_int(s.get("replacement_teacher_id")) != scoped_teacher.id:
                    continue
                special_rows.append(s)

        absent_teacher_ids = {tid for tid in (as_int(r.get("original_teacher_id")) for r in replacements_rows) if tid is not None}
        replacement_teacher_ids = {tid for tid in (as_int(r.get("replacement_teacher_id")) for r in replacements_rows) if tid is not None}
        replacement_teacher_ids.update(
            tid for tid in (as_int(r.get("replacement_teacher_id")) for r in special_rows) if tid is not None
        )

        by_replacing_teacher: dict[str, dict[str, int]] = {}
        for r in replacements_rows:
            tid = as_int(r.get("replacement_teacher_id"))
            if tid is None:
                continue
            name = teacher_map.get(tid) or f"ID:{tid}"
            bucket = by_replacing_teacher.setdefault(name, {"regular": 0, "special": 0})
            if as_bool(r.get("production_necessity")):
                bucket["special"] += 1
            else:
                bucket["regular"] += 1
                                                 
        for s in special_rows:
            tid = as_int(s.get("replacement_teacher_id"))
            if tid is None:
                continue
            name = teacher_map.get(tid) or f"ID:{tid}"
            bucket = by_replacing_teacher.setdefault(name, {"regular": 0, "special": 0})
            bucket["regular"] += 1

        by_absent_teacher: dict[str, int] = {}
        for r in replacements_rows:
            tid = as_int(r.get("original_teacher_id"))
            if tid is None:
                continue
            name = teacher_map.get(tid) or f"ID:{tid}"
            by_absent_teacher[name] = by_absent_teacher.get(name, 0) + 1

        top_replacing = sorted(
            by_replacing_teacher.items(),
            key=lambda x: (-(x[1]["regular"] + x[1]["special"]), x[0]),
        )[:10]
        top_absent = sorted(by_absent_teacher.items(), key=lambda x: (-x[1], x[0]))[:10]
        all_replacing = sorted(
            by_replacing_teacher.items(),
            key=lambda x: (-(x[1]["regular"] + x[1]["special"]), x[0]),
        )

        max_replacing = max([(c["regular"] + c["special"]) for _, c in top_replacing], default=0)
        max_absent = max([c for _, c in top_absent], default=0)
        max_all_replacing = max([(c["regular"] + c["special"]) for _, c in all_replacing], default=0)

        daily_stats: dict[str, dict] = {}
        daily_details: dict[str, list[dict]] = {}
        for r in replacements_rows:
            key = str(r.get("date") or "")
            item = daily_stats.setdefault(
                key, {"count": 0, "regular": 0, "special": 0, "absent_ids": set(), "replacing_ids": set()}
            )
            item["count"] += 1
            if as_bool(r.get("production_necessity")):
                item["special"] += 1
            else:
                item["regular"] += 1
            ot = as_int(r.get("original_teacher_id"))
            rt = as_int(r.get("replacement_teacher_id"))
            if ot is not None:
                item["absent_ids"].add(ot)
            if rt is not None:
                item["replacing_ids"].add(rt)
            if scoped_teacher:
                lesson = lesson_map.get(as_int(r.get("lesson_id")) or -1) or {}
                daily_details.setdefault(key, []).append(
                    {
                        "lesson_number": as_int(lesson.get("lesson_number")),
                        "class_group": str(lesson.get("class_group") or ""),
                        "original_teacher": teacher_map.get(ot, "") if ot is not None else "",
                        "is_special": bool(as_bool(r.get("production_necessity"))),
                    }
                )

                                                 
        for s in special_rows:
            key = str(s.get("date") or "")
            item = daily_stats.setdefault(
                key, {"count": 0, "regular": 0, "special": 0, "absent_ids": set(), "replacing_ids": set()}
            )
            item["count"] += 1
            item["regular"] += 1
            rt = as_int(s.get("replacement_teacher_id"))
            if rt is not None:
                item["replacing_ids"].add(rt)
            if scoped_teacher:
                ot = as_int(s.get("original_teacher_id"))
                daily_details.setdefault(key, []).append(
                    {
                        "lesson_number": as_int(s.get("lesson_number")),
                        "class_group": str(s.get("class_group") or ""),
                        "original_teacher": teacher_map.get(ot, "") if ot is not None else "",
                        "is_special": False,
                    }
                )

        daily_rows = []
        cur = start
        while cur <= end:
            k = cur.strftime("%Y-%m-%d")
            ds = daily_stats.get(k) or {}
            daily_rows.append(
                {
                    "date": k,
                    "label": cur.strftime("%d.%m"),
                    "count": int(ds.get("count", 0)),
                    "regular_replacements": int(ds.get("regular", 0)),
                    "special_replacements": int(ds.get("special", 0)),
                    "absent_teachers": int(len(ds.get("absent_ids", set()))),
                    "replacement_teachers": int(len(ds.get("replacing_ids", set()))),
                }
            )
            cur += timedelta(days=1)

        total_special = sum(1 for r in replacements_rows if as_bool(r.get("production_necessity")))
        total_regular = (len(replacements_rows) - total_special) + len(special_rows)
        payload = {
            "mode": mode_norm,
            "range": {
                "from": start.strftime("%Y-%m-%d"),
                "to": end.strftime("%Y-%m-%d"),
                "label": range_label,
            },
            "totals": {
                "replacement_lessons": int(total_regular + total_special),
                "regular_replacements": int(total_regular),
                "special_replacements": int(total_special),
                "absent_teachers": int(len(absent_teacher_ids)),
                "replacement_teachers": int(len(replacement_teacher_ids)),
            },
            "top_replacing_teachers": [
                {
                    "teacher": name,
                    "icon": get_icon_for_display_name(name),
                    "count": int(cnt["regular"] + cnt["special"]),
                    "regular_count": int(cnt["regular"]),
                    "special_count": int(cnt["special"]),
                    "share": int(round(((cnt["regular"] + cnt["special"]) / max_replacing) * 100))
                    if max_replacing
                    else 0,
                }
                for name, cnt in top_replacing
            ],
            "top_absent_teachers": [
                {
                    "teacher": name,
                    "icon": get_icon_for_display_name(name),
                    "count": int(cnt),
                    "share": int(round((cnt / max_absent) * 100)) if max_absent else 0,
                }
                for name, cnt in top_absent
            ],
            "all_replacing_teachers": [
                {
                    "teacher": name,
                    "icon": get_icon_for_display_name(name),
                    "count": int(cnt["regular"] + cnt["special"]),
                    "regular_count": int(cnt["regular"]),
                    "special_count": int(cnt["special"]),
                    "share": int(round(((cnt["regular"] + cnt["special"]) / max_all_replacing) * 100))
                    if max_all_replacing
                    else 0,
                }
                for name, cnt in all_replacing
            ],
            "daily": daily_rows,
        }
        if scoped_teacher:
            payload["teacher_scope"] = scoped_teacher.full_name
            payload["daily_details"] = daily_details
        cache.set(cache_key, payload, 600)
        return JsonResponse(payload)

    replacements_qs = (
        _teacher_replacements()
        .filter(date__range=(start, end))
        .exclude(ignore_in_reports=True)
        .select_related("replacement_teacher", "original_teacher", "lesson")
    )
    special_qs = (
        SpecialReplacement.objects
        .filter(date__range=(start, end))
        .select_related("replacement_teacher")
    )
    if scoped_teacher:
        replacements_qs = replacements_qs.filter(replacement_teacher=scoped_teacher)
        special_qs = special_qs.filter(replacement_teacher=scoped_teacher)

    absent_teacher_ids = set(
        replacements_qs.exclude(original_teacher_id__isnull=True).values_list("original_teacher_id", flat=True)
    )
    replacement_teacher_ids = set(
        replacements_qs.exclude(replacement_teacher_id__isnull=True).values_list("replacement_teacher_id", flat=True)
    )
    replacement_teacher_ids.update(
        special_qs.exclude(replacement_teacher_id__isnull=True).values_list("replacement_teacher_id", flat=True)
    )

    by_replacing_teacher: dict[str, dict[str, int]] = {}
    for name, is_production in replacements_qs.exclude(replacement_teacher__isnull=True).values_list(
        "replacement_teacher__full_name",
        "production_necessity",
    ):
        bucket = by_replacing_teacher.setdefault(name, {"regular": 0, "special": 0})
        if bool(is_production):
            bucket["special"] += 1
        else:
            bucket["regular"] += 1
                                             
    for name in special_qs.exclude(replacement_teacher__isnull=True).values_list("replacement_teacher__full_name", flat=True):
        bucket = by_replacing_teacher.setdefault(name, {"regular": 0, "special": 0})
        bucket["regular"] += 1

    by_absent_teacher: dict[str, int] = {}
    for name in replacements_qs.exclude(original_teacher__isnull=True).values_list("original_teacher__full_name", flat=True):
        by_absent_teacher[name] = by_absent_teacher.get(name, 0) + 1

    top_replacing = sorted(
        by_replacing_teacher.items(),
        key=lambda x: (-(x[1]["regular"] + x[1]["special"]), x[0]),
    )[:10]
    top_absent = sorted(by_absent_teacher.items(), key=lambda x: (-x[1], x[0]))[:10]
    all_replacing = sorted(
        by_replacing_teacher.items(),
        key=lambda x: (-(x[1]["regular"] + x[1]["special"]), x[0]),
    )

    max_replacing = max([(c["regular"] + c["special"]) for _, c in top_replacing], default=0)
    max_absent = max([c for _, c in top_absent], default=0)
    max_all_replacing = max([(c["regular"] + c["special"]) for _, c in all_replacing], default=0)

    daily_stats: dict[str, dict] = {}
    daily_details: dict[str, list[dict]] = {}

    for r in replacements_qs:
        key = r.date.strftime("%Y-%m-%d")
        item = daily_stats.setdefault(key, {
            "count": 0,
            "regular": 0,
            "special": 0,
            "absent_ids": set(),
            "replacing_ids": set(),
        })
        item["count"] += 1
        if bool(getattr(r, "production_necessity", False)):
            item["special"] += 1
        else:
            item["regular"] += 1
        if r.original_teacher_id:
            item["absent_ids"].add(int(r.original_teacher_id))
        if r.replacement_teacher_id:
            item["replacing_ids"].add(int(r.replacement_teacher_id))
        if scoped_teacher:
            daily_details.setdefault(key, []).append({
                "lesson_number": (r.lesson.lesson_number if r.lesson else None),
                "class_group": (r.lesson.class_group if r.lesson else ""),
                "original_teacher": (r.original_teacher.full_name if r.original_teacher else ""),
                "is_special": bool(getattr(r, "production_necessity", False)),
            })

                                             
    for s in special_qs:
        key = s.date.strftime("%Y-%m-%d")
        item = daily_stats.setdefault(key, {
            "count": 0,
            "regular": 0,
            "special": 0,
            "absent_ids": set(),
            "replacing_ids": set(),
        })
        item["count"] += 1
        item["regular"] += 1
        if s.replacement_teacher_id:
            item["replacing_ids"].add(int(s.replacement_teacher_id))
        if scoped_teacher:
            daily_details.setdefault(key, []).append({
                "lesson_number": s.lesson_number,
                "class_group": s.class_group or "",
                "original_teacher": (s.original_teacher.full_name if s.original_teacher else ""),
                "is_special": False,
            })

    daily_rows = []
    cur = start
    while cur <= end:
        k = cur.strftime("%Y-%m-%d")
        ds = daily_stats.get(k) or {}
        daily_rows.append({
            "date": k,
            "label": cur.strftime("%d.%m"),
            "count": int(ds.get("count", 0)),
            "regular_replacements": int(ds.get("regular", 0)),
            "special_replacements": int(ds.get("special", 0)),
            "absent_teachers": int(len(ds.get("absent_ids", set()))),
            "replacement_teachers": int(len(ds.get("replacing_ids", set()))),
        })
        cur += timedelta(days=1)

    total_special = replacements_qs.filter(production_necessity=True).count()
    total_regular = replacements_qs.filter(production_necessity=False).count() + special_qs.count()

    payload = {
        "mode": mode_norm,
        "range": {
            "from": start.strftime("%Y-%m-%d"),
            "to": end.strftime("%Y-%m-%d"),
            "label": range_label,
        },
        "totals": {
            "replacement_lessons": int(total_regular + total_special),
            "regular_replacements": int(total_regular),
            "special_replacements": int(total_special),
            "absent_teachers": int(len(absent_teacher_ids)),
            "replacement_teachers": int(len(replacement_teacher_ids)),
        },
        "top_replacing_teachers": [
            {
                "teacher": name,
                "icon": get_icon_for_display_name(name),
                "count": int(cnt["regular"] + cnt["special"]),
                "regular_count": int(cnt["regular"]),
                "special_count": int(cnt["special"]),
                "share": int(round(((cnt["regular"] + cnt["special"]) / max_replacing) * 100)) if max_replacing else 0,
            }
            for name, cnt in top_replacing
        ],
        "top_absent_teachers": [
            {
                "teacher": name,
                "icon": get_icon_for_display_name(name),
                "count": int(cnt),
                "share": int(round((cnt / max_absent) * 100)) if max_absent else 0,
            }
            for name, cnt in top_absent
        ],
        "all_replacing_teachers": [
            {
                "teacher": name,
                "icon": get_icon_for_display_name(name),
                "count": int(cnt["regular"] + cnt["special"]),
                "regular_count": int(cnt["regular"]),
                "special_count": int(cnt["special"]),
                "share": int(round(((cnt["regular"] + cnt["special"]) / max_all_replacing) * 100)) if max_all_replacing else 0,
            }
            for name, cnt in all_replacing
        ],
        "daily": daily_rows,
    }
    if scoped_teacher:
        payload["teacher_scope"] = scoped_teacher.full_name
        payload["daily_details"] = daily_details
    return JsonResponse(payload)


def _report_range_from_request(request):
    """Resolve report range with backward compatibility for old monthly `?date=` requests."""
    mode = request.GET.get("mode")
    date_str = request.GET.get("date")
    date_from = request.GET.get("date_from")
    date_to = request.GET.get("date_to")
    month_str = request.GET.get("month")

                                                                                
    if date_str and not mode and not date_from and not date_to and not month_str:
        selected = datetime.strptime(date_str, "%Y-%m-%d").date()
        start = selected.replace(day=1)
        if start.month == 12:
            next_month = start.replace(year=start.year + 1, month=1, day=1)
        else:
            next_month = start.replace(month=start.month + 1, day=1)
        end = next_month - timedelta(days=1)
        return "month", start, end, f"Месяц: {start.strftime('%m.%Y')}"

    mode_norm, start, end, label = _stats_date_range(
        mode or "day",
        date_str,
        date_from,
        date_to,
        month_str,
    )
    return mode_norm, start, end, label


@login_required
def replacement_summary_report(request):
    """Формирует приказ/отчёт по замещениям за МЕСЯЦ в формате *один в один* как шаблон.

    Приоритетный шаблон: «Приказ о замещении ФЕВРАЛЬ 2026.docx» в корне проекта.
    Fallback: replacements/report_templates/order_template.docx.
    """
    from collections import Counter, defaultdict
    import copy
    from pathlib import Path
    from urllib.parse import quote

    from django.conf import settings

    try:
        mode_norm, start, end, range_label = _report_range_from_request(request)

        if _use_gsheets_backend():
            teacher_map = _gs_teacher_map()
            lesson_map = _gs_lesson_map()
            subject_map = _gs_subject_map()

            grouped: dict[tuple[str, str], Counter] = defaultdict(Counter)
            vacancy_subjects: dict[tuple[str, str], set[str]] = defaultdict(set)
            for r in _gs_store().get_table_dicts("replacements_replacement"):
                if not _gs_is_teacher_replacement(r):
                    continue
                if as_bool(r.get("production_necessity")) or as_bool(r.get("ignore_in_reports")):
                    continue
                d_raw = str(r.get("date") or "")
                try:
                    d = datetime.strptime(d_raw, "%Y-%m-%d").date()
                except Exception:
                    continue
                if not (start <= d <= end):
                    continue

                repl_id = as_int(r.get("replacement_teacher_id"))
                orig_id = as_int(r.get("original_teacher_id"))
                lesson_id = as_int(r.get("lesson_id"))

                repl_name = teacher_map.get(repl_id, "—") if repl_id is not None else "—"
                orig_is_vacancy = orig_id is None or not teacher_map.get(orig_id)
                orig_name_base = teacher_map.get(orig_id, "Вакансия") if orig_id is not None else "Вакансия"

                lesson = lesson_map.get(lesson_id) if lesson_id is not None else None
                subject_id = as_int((lesson or {}).get("subject_id_subject"))
                subj = subject_map.get(subject_id, "Без предмета") if subject_id is not None else "Без предмета"

                grouped[(repl_name, orig_name_base)][subj] += 1
                if orig_is_vacancy and subj and subj != "Без предмета":
                    vacancy_subjects[(repl_name, orig_name_base)].add(subj)

            rows_data: list[tuple[str, list[str], str]] = []
            for (repl_name, orig_base), subj_counts in grouped.items():
                parts = []
                for subj, cnt in sorted(subj_counts.items(), key=lambda x: (-x[1], x[0])):
                    if subj in ("Без предмета", "-"):
                        parts.append(f"{cnt}")
                    else:
                        parts.append(f"{cnt} ({subj})")
                orig_name = orig_base
                if orig_base == "Вакансия":
                    subs = sorted(vacancy_subjects.get((repl_name, orig_base), set()))
                    if subs:
                        orig_name = f"Вакансия ({', '.join(subs)})"
                rows_data.append((repl_name, parts or ["—"], orig_name))

            if not rows_data:
                rows_data = [("—", ["—"], "—")]
            rows_data.sort(key=lambda x: (x[2].lower(), x[0].lower()))
        else:
                                              
            replacements = (
                _teacher_replacements()
                .filter(date__range=(start, end))
                .exclude(production_necessity=True)
                .exclude(ignore_in_reports=True)
                .select_related("replacement_teacher", "original_teacher", "lesson__subject")
            )
                                                                          
            special_replacements = (
                SpecialReplacement.objects
                .filter(date__range=(start, end))
                .select_related("replacement_teacher", "original_teacher")
            )

                                                                                                           
            grouped: dict[tuple[str, str], Counter] = defaultdict(Counter)
            vacancy_subjects: dict[tuple[str, str], set[str]] = defaultdict(set)

            for r in replacements:
                repl_name = (r.replacement_teacher.full_name if r.replacement_teacher else "—")
                orig_is_vacancy = (r.original_teacher is None)
                orig_name_base = (r.original_teacher.full_name if r.original_teacher else "Вакансия")

                subj = getattr(getattr(r.lesson, "subject", None), "name", None) or "Без предмета"
                grouped[(repl_name, orig_name_base)][subj] += 1
                if orig_is_vacancy and subj and subj != "Без предмета":
                    vacancy_subjects[(repl_name, orig_name_base)].add(subj)

            for s in special_replacements:
                repl_name = (s.replacement_teacher.full_name if s.replacement_teacher else "—")
                orig_is_vacancy = (s.original_teacher is None)
                orig_name_base = (s.original_teacher.full_name if s.original_teacher else "Вакансия")
                subj = (s.subject_name or "").strip() or "Без предмета"
                grouped[(repl_name, orig_name_base)][subj] += 1
                if orig_is_vacancy and subj and subj != "Без предмета":
                    vacancy_subjects[(repl_name, orig_name_base)].add(subj)

                                          
            rows_data: list[tuple[str, list[str], str]] = []
            for (repl_name, orig_base), subj_counts in grouped.items():
                parts = []
                for subj, cnt in sorted(subj_counts.items(), key=lambda x: (-x[1], x[0])):
                    if subj == "Без предмета" or subj == "-":
                        parts.append(f"{cnt}")
                    else:
                        parts.append(f"{cnt} ({subj})")

                                                                 
                orig_name = orig_base
                if orig_base == "Вакансия":
                    subs = sorted(vacancy_subjects.get((repl_name, orig_base), set()))
                    if subs:
                        orig_name = f"Вакансия ({', '.join(subs)})"

                rows_data.append((repl_name, parts, orig_name))

            if not rows_data:
                rows_data = [("—", ["—"], "—")]

                                                                                            
            rows_data.sort(key=lambda x: (x[2].lower(), x[0].lower()))

        template_candidates = [
            Path(settings.BASE_DIR) / "Приказ о замещении ФЕВРАЛЬ 2026.docx",
            Path(settings.BASE_DIR) / "replacements" / "report_templates" / "order_template.docx",
        ]
        template_path = next((p for p in template_candidates if p.exists()), None)
        if template_path is None:
            return JsonResponse({"error": "Не найден шаблон отчёта (Приказ о замещении ФЕВРАЛЬ 2026.docx / order_template.docx)"}, status=500)

        doc = Document(str(template_path))
        if not doc.tables:
            return JsonResponse({"error": "Шаблон отчёта не содержит таблицу"}, status=500)

        table = doc.tables[0]
        if len(table.rows) < 2:
            return JsonResponse({"error": "Шаблон отчёта должен содержать хотя бы одну строку-образец данных"}, status=500)

        style_row = table.rows[1]
        style_tr = copy.deepcopy(style_row._tr)

                                                                                          
                                                                           
        while len(table.rows) > 2:
            table._tbl.remove(table.rows[2]._tr)

        def _clear_paragraph(p):
                                                                 
            for r in list(p.runs):
                p._p.remove(r._r)

        def _apply_paragraph_format(dst_p, src_p):
            dst_p.style = src_p.style
            dst_p.alignment = src_p.alignment
                                                 
            try:
                dst_pf = dst_p.paragraph_format
                src_pf = src_p.paragraph_format
                dst_pf.left_indent = src_pf.left_indent
                dst_pf.right_indent = src_pf.right_indent
                dst_pf.first_line_indent = src_pf.first_line_indent
                dst_pf.space_before = src_pf.space_before
                dst_pf.space_after = src_pf.space_after
                dst_pf.line_spacing = src_pf.line_spacing
                dst_pf.keep_together = src_pf.keep_together
                dst_pf.keep_with_next = src_pf.keep_with_next
                dst_pf.widow_control = src_pf.widow_control
            except Exception:
                pass

        def _write_cell_lines(cell, lines: list[str], template_cell):
            """Записать в ячейку несколько строк как отдельные параграфы, сохранив стиль из шаблона."""
            if not lines:
                lines = ["—"]

            tmpl_ps = template_cell.paragraphs
            tmpl_base = tmpl_ps[0] if tmpl_ps else None

                                                      
            while len(cell.paragraphs) < len(lines):
                cell.add_paragraph()

                        
            for i, line in enumerate(lines):
                p = cell.paragraphs[i]
                src_p = tmpl_ps[i] if i < len(tmpl_ps) else tmpl_base
                if src_p:
                    _apply_paragraph_format(p, src_p)
                _clear_paragraph(p)
                run = p.add_run(line)
                                                    
                run.font.size = Pt(12)

                                      
            for j in range(len(cell.paragraphs) - 1, len(lines) - 1, -1):
                try:
                    cell._tc.remove(cell.paragraphs[j]._p)
                except Exception:
                                                         
                    _clear_paragraph(cell.paragraphs[j])

        for repl_name, parts, orig_name in rows_data:
            new_tr = copy.deepcopy(style_tr)
            table._tbl.append(new_tr)
            row = table.rows[-1]

                                                                                          
            _write_cell_lines(row.cells[0], [repl_name], style_row.cells[0])

            if len(parts) <= 2:
                _write_cell_lines(row.cells[1], [" + ".join(parts)], style_row.cells[1])
            else:
                _write_cell_lines(row.cells[1], parts, style_row.cells[1])

            _write_cell_lines(row.cells[2], [orig_name], style_row.cells[2])

                                                       
        try:
            table._tbl.remove(table.rows[1]._tr)
        except Exception:
            pass

                                           
        if mode_norm == "day":
            filename = f"Приказ о замещении за {start.strftime('%d.%m.%Y')}.docx"
            safe_ascii = f"replacement_order_day_{start.strftime('%Y_%m_%d')}.docx"
        elif mode_norm == "period":
            filename = f"Приказ о замещении за период {start.strftime('%d.%m.%Y')} - {end.strftime('%d.%m.%Y')}.docx"
            safe_ascii = f"replacement_order_period_{start.strftime('%Y_%m_%d')}_{end.strftime('%Y_%m_%d')}.docx"
        else:
            ru_months = {
                1: "ЯНВАРЬ", 2: "ФЕВРАЛЬ", 3: "МАРТ", 4: "АПРЕЛЬ", 5: "МАЙ", 6: "ИЮНЬ",
                7: "ИЮЛЬ", 8: "АВГУСТ", 9: "СЕНТЯБРЬ", 10: "ОКТЯБРЬ", 11: "НОЯБРЬ", 12: "ДЕКАБРЬ"
            }
            month_name = ru_months.get(start.month, str(start.month))
            filename = f"Приказ о замещении {month_name} {start.year}.docx"
            safe_ascii = f"replacement_order_{start.year}_{start.month:02d}.docx"

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = (
            f'attachment; filename="{safe_ascii}"; '
            f"filename*=UTF-8''{quote(filename)}"
        )
        doc.save(response)
        log_activity(
            request,
            "export_monthly_order_docx",
            {"mode": mode_norm, "range": range_label, "from": str(start), "to": str(end), "rows": len(rows_data)},
        )
        return response

    except Exception:
        return _api_error("Внутренняя ошибка сервера", status=500, code="internal_error")


@login_required
def replacement_daily_summary_docx(request):
    """DOCX summary: per-date counts by replacing teacher split by grade."""
    from urllib.parse import quote
    try:
        mode_norm, start, end, range_label = _report_range_from_request(request)
    except Exception:
        return _api_error("Некорректные данные запроса", status=400, code="bad_request")

    if _use_gsheets_backend():
        teacher_map = _gs_teacher_map()
        lesson_map = _gs_lesson_map()
        rows = []
        for r in _gs_store().get_table_dicts("replacements_replacement"):
            if not _gs_is_teacher_replacement(r):
                continue
            if as_bool(r.get("production_necessity")) or as_bool(r.get("ignore_in_reports")):
                continue
            d_raw = str(r.get("date") or "")
            try:
                d = datetime.strptime(d_raw, "%Y-%m-%d").date()
            except Exception:
                continue
            if not (start <= d <= end):
                continue
            rows.append(r)

        def short_teacher_name(full_name: str | None) -> str:
            parts = [p for p in (full_name or "").split() if p]
            if not parts:
                return "—"
            surname = parts[0]
            initials = "".join(f"{p[0].upper()}." for p in parts[1:3] if p)
            return f"{surname} {initials}".strip()

        day_map: dict[str, dict[str, dict[str, int]]] = {}
        teacher_map_by_day: dict[str, dict[str, dict[str, int]]] = {}

        for r in rows:
            day_key = str(r.get("date"))
            teacher_id = as_int(r.get("replacement_teacher_id"))
            if teacher_id is None:
                continue
            teacher_name = short_teacher_name(teacher_map.get(teacher_id))
            lesson_id = as_int(r.get("lesson_id"))
            lesson = lesson_map.get(lesson_id) if lesson_id is not None else None
            class_group = str((lesson or {}).get("class_group") or "")
            grade = extract_grade(class_group) if class_group else None
            is_primary = grade is not None and grade <= 4
            entry = day_map.setdefault(day_key, {}).setdefault(teacher_name, {"zs": 0, "zn": 0})
            t_entry = teacher_map_by_day.setdefault(teacher_name, {}).setdefault(day_key, {"zs": 0, "zn": 0})
            if is_primary:
                entry["zn"] += 1
                t_entry["zn"] += 1
            else:
                entry["zs"] += 1
                t_entry["zs"] += 1

        doc = Document()
        for day_key in sorted(day_map.keys(), key=lambda d: datetime.strptime(d, "%Y-%m-%d")):
            p_day = doc.add_paragraph()
            p_day.add_run(f"Дата: {datetime.strptime(day_key, '%Y-%m-%d').strftime('%d.%m.%Y')}").bold = True
            for teacher_name in sorted(day_map[day_key].keys(), key=lambda x: x.lower()):
                zs = day_map[day_key][teacher_name]["zs"]
                zn = day_map[day_key][teacher_name]["zn"]
                parts = []
                if zs > 0:
                    parts.append(f"{zs} (ЗС)")
                if zn > 0:
                    parts.append(f"{zn} (ЗН)")
                if not parts:
                    continue
                doc.add_paragraph(f"{teacher_name}: " + ", ".join(parts))
            doc.add_paragraph("")

        if not day_map:
            p_day = doc.add_paragraph()
            p_day.add_run(f"Дата: {start.strftime('%d.%m.%Y')}").bold = True
            doc.add_paragraph("—")

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        if mode_norm == "day":
            safe_ascii = f"replacement_summary_day_{start.strftime('%Y_%m_%d')}.docx"
            filename = f"Сводка замещений за {start.strftime('%d.%m.%Y')}.docx"
        elif mode_norm == "period":
            safe_ascii = f"replacement_summary_period_{start.strftime('%Y_%m_%d')}_{end.strftime('%Y_%m_%d')}.docx"
            filename = f"Сводка замещений за период {start.strftime('%d.%m.%Y')} - {end.strftime('%d.%m.%Y')}.docx"
        else:
            safe_ascii = f"replacement_summary_{start.strftime('%Y_%m')}.docx"
            filename = f"Сводка_замещений_{start.strftime('%m.%Y')}.docx"
        response["Content-Disposition"] = (
            f'attachment; filename="{safe_ascii}"; '
            f"filename*=UTF-8''{quote(filename)}"
        )
        doc.save(response)
        log_activity(
            request,
            "export_replacement_summary_docx",
            {"mode": mode_norm, "range": range_label, "from": str(start), "to": str(end), "teachers": len(teacher_map_by_day)},
        )
        return response

    reps = (
        _teacher_replacements()
        .filter(date__range=(start, end))
        .exclude(production_necessity=True)
        .exclude(ignore_in_reports=True)
        .select_related("replacement_teacher", "lesson")
    )
    special_reps = (
        SpecialReplacement.objects
        .filter(date__range=(start, end))
        .select_related("replacement_teacher")
    )

    def short_teacher_name(full_name: str | None) -> str:
        parts = [p for p in (full_name or "").split() if p]
        if not parts:
            return "—"
        surname = parts[0]
        initials = "".join(f"{p[0].upper()}." for p in parts[1:3] if p)
        return f"{surname} {initials}".strip()

                                                    
    day_map: dict[str, dict[str, dict[str, int]]] = {}
                                                    
    teacher_map: dict[str, dict[str, dict[str, int]]] = {}

    for r in reps:
        if not r.date:
            continue
        day_key = r.date.strftime('%d.%m.%Y')
        teacher = r.replacement_teacher
        if not teacher:
            continue
        teacher_name = short_teacher_name(teacher.full_name)
        grade = extract_grade(r.lesson.class_group) if r.lesson else None
        is_primary = grade is not None and grade <= 4
        entry = day_map.setdefault(day_key, {}).setdefault(teacher_name, {"zs": 0, "zn": 0})
        t_entry = teacher_map.setdefault(teacher_name, {}).setdefault(day_key, {"zs": 0, "zn": 0})
        if is_primary:
            entry["zn"] += 1
            t_entry["zn"] += 1
        else:
            entry["zs"] += 1
            t_entry["zs"] += 1

    for s in special_reps:
        if not s.date:
            continue
        teacher = s.replacement_teacher
        if not teacher:
            continue
        day_key = s.date.strftime('%d.%m.%Y')
        teacher_name = short_teacher_name(teacher.full_name)
        grade = extract_grade(s.class_group) if s.class_group else None
        is_primary = grade is not None and grade <= 4
        entry = day_map.setdefault(day_key, {}).setdefault(teacher_name, {"zs": 0, "zn": 0})
        t_entry = teacher_map.setdefault(teacher_name, {}).setdefault(day_key, {"zs": 0, "zn": 0})
        if is_primary:
            entry["zn"] += 1
            t_entry["zn"] += 1
        else:
            entry["zs"] += 1
            t_entry["zs"] += 1

                                
    doc = Document()
    for day_key in sorted(day_map.keys(), key=lambda d: datetime.strptime(d, "%d.%m.%Y")):
        p_day = doc.add_paragraph()
        p_day.add_run(f"Дата: {day_key}").bold = True
        for teacher_name in sorted(day_map[day_key].keys(), key=lambda x: x.lower()):
            zs = day_map[day_key][teacher_name]["zs"]
            zn = day_map[day_key][teacher_name]["zn"]
            parts = []
            if zs > 0:
                parts.append(f"{zs} (ЗС)")
            if zn > 0:
                parts.append(f"{zn} (ЗН)")
            if not parts:
                continue
            doc.add_paragraph(f"{teacher_name}: " + ", ".join(parts))
        doc.add_paragraph("")

    if not day_map:
        p_day = doc.add_paragraph()
        p_day.add_run(f"Дата: {start.strftime('%d.%m.%Y')}").bold = True
        doc.add_paragraph("—")

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    if mode_norm == "day":
        safe_ascii = f"replacement_summary_day_{start.strftime('%Y_%m_%d')}.docx"
        filename = f"Сводка замещений за {start.strftime('%d.%m.%Y')}.docx"
    elif mode_norm == "period":
        safe_ascii = f"replacement_summary_period_{start.strftime('%Y_%m_%d')}_{end.strftime('%Y_%m_%d')}.docx"
        filename = f"Сводка замещений за период {start.strftime('%d.%m.%Y')} - {end.strftime('%d.%m.%Y')}.docx"
    else:
        safe_ascii = f"replacement_summary_{start.strftime('%Y_%m')}.docx"
        filename = f"Сводка_замещений_{start.strftime('%m.%Y')}.docx"
    response["Content-Disposition"] = (
        f'attachment; filename="{safe_ascii}"; '
        f"filename*=UTF-8''{quote(filename)}"
    )
    doc.save(response)
    log_activity(
        request,
        "export_replacement_summary_docx",
        {"mode": mode_norm, "range": range_label, "from": str(start), "to": str(end), "teachers": len(teacher_map)},
    )
    return response


@login_required
def replacement_teacher_summary_docx(request):
    """DOCX summary: teacher -> dates with counts (example format requested by user)."""
    from urllib.parse import quote
    try:
        mode_norm, start, end, range_label = _report_range_from_request(request)
    except Exception:
        return _api_error("Некорректные данные запроса", status=400, code="bad_request")

    if _use_gsheets_backend():
        teacher_map = _gs_teacher_map()
        lesson_map = _gs_lesson_map()
        rows = []
        for r in _gs_store().get_table_dicts("replacements_replacement"):
            if not _gs_is_teacher_replacement(r):
                continue
            if as_bool(r.get("production_necessity")) or as_bool(r.get("ignore_in_reports")):
                continue
            d_raw = str(r.get("date") or "")
            try:
                d = datetime.strptime(d_raw, "%Y-%m-%d").date()
            except Exception:
                continue
            if start <= d <= end:
                rows.append(r)

        def teacher_name_with_initials(full_name: str | None) -> str:
            parts = [p for p in (full_name or "").split() if p]
            if not parts:
                return "—"
            surname = parts[0]
            initials = [f"{p[0].upper()}." for p in parts[1:3] if p]
            return f"{surname} {' '.join(initials)}".strip()

        def replacement_word(count: int) -> str:
            n = abs(count) % 100
            n1 = n % 10
            if 11 <= n <= 14:
                return "замещений"
            if n1 == 1:
                return "замещение"
            if 2 <= n1 <= 4:
                return "замещения"
            return "замещений"

        teacher_stats: dict[int, dict] = {}
        for r in rows:
            teacher_id = as_int(r.get("replacement_teacher_id"))
            if teacher_id is None:
                continue
            teacher_entry = teacher_stats.setdefault(
                teacher_id,
                {"name": teacher_name_with_initials(teacher_map.get(teacher_id)), "total": 0, "days": {}},
            )
            teacher_entry["total"] += 1
            day_key = datetime.strptime(str(r.get("date")), "%Y-%m-%d").strftime("%d.%m.%Y")
            lesson_id = as_int(r.get("lesson_id"))
            lesson = lesson_map.get(lesson_id) if lesson_id is not None else None
            class_group = str((lesson or {}).get("class_group") or "")
            grade = extract_grade(class_group) if class_group else None
            is_primary = grade is not None and grade <= 4
            d_entry = teacher_entry["days"].setdefault(day_key, {"zs": 0, "zn": 0})
            if is_primary:
                d_entry["zn"] += 1
            else:
                d_entry["zs"] += 1

        doc = Document()
        teachers_sorted = sorted(teacher_stats.values(), key=lambda x: x["name"].lower())
        for teacher_data in teachers_sorted:
            teacher_name = teacher_data["name"]
            total = teacher_data["total"]
            p_teacher = doc.add_paragraph()
            p_teacher.add_run(f"{teacher_name} ({total} {replacement_word(total)}):").bold = True

            has_days = False
            day_map = teacher_data["days"]
            for day_key in sorted(day_map.keys(), key=lambda d: datetime.strptime(d, "%d.%m.%Y")):
                zs = day_map[day_key]["zs"]
                zn = day_map[day_key]["zn"]
                parts = []
                if zs > 0:
                    parts.append(f"{zs} (ЗС)")
                if zn > 0:
                    parts.append(f"{zn} (ЗН)")
                if not parts:
                    continue
                has_days = True
                p_day = doc.add_paragraph()
                p_day.add_run(f"Дата: {day_key}").bold = True
                p_day.add_run(f" - {', '.join(parts)}")

            if not has_days:
                doc.add_paragraph("—")
            doc.add_paragraph("")

        if not teacher_stats:
            doc.add_paragraph("—")

        if mode_norm == "day":
            safe_ascii = f"replacement_summary_teachers_day_{start.strftime('%Y_%m_%d')}.docx"
            filename = f"Сводка по педагогам за {start.strftime('%d.%m.%Y')}.docx"
        elif mode_norm == "period":
            safe_ascii = f"replacement_summary_teachers_period_{start.strftime('%Y_%m_%d')}_{end.strftime('%Y_%m_%d')}.docx"
            filename = f"Сводка по педагогам за период {start.strftime('%d.%m.%Y')} - {end.strftime('%d.%m.%Y')}.docx"
        else:
            safe_ascii = f"replacement_summary_teachers_{start.strftime('%Y_%m')}.docx"
            filename = f"Сводка_по_педагогам_{start.strftime('%m.%Y')}.docx"

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = (
            f'attachment; filename="{safe_ascii}"; '
            f"filename*=UTF-8''{quote(filename)}"
        )
        doc.save(response)
        log_activity(
            request,
            "export_replacement_teacher_summary_docx",
            {"mode": mode_norm, "range": range_label, "from": str(start), "to": str(end), "teachers": len(teacher_stats)},
        )
        return response

    reps = (
        _teacher_replacements()
        .filter(date__range=(start, end))
        .exclude(production_necessity=True)
        .exclude(ignore_in_reports=True)
        .select_related("replacement_teacher", "lesson")
    )
    special_reps = (
        SpecialReplacement.objects
        .filter(date__range=(start, end))
        .select_related("replacement_teacher")
    )

    def teacher_name_with_initials(full_name: str | None) -> str:
        parts = [p for p in (full_name or "").split() if p]
        if not parts:
            return "—"
        surname = parts[0]
        initials = [f"{p[0].upper()}." for p in parts[1:3] if p]
        return f"{surname} {' '.join(initials)}".strip()

    def replacement_word(count: int) -> str:
        n = abs(count) % 100
        n1 = n % 10
        if 11 <= n <= 14:
            return "замещений"
        if n1 == 1:
            return "замещение"
        if 2 <= n1 <= 4:
            return "замещения"
        return "замещений"

    teacher_map: dict[int, dict] = {}
    for r in reps:
        if not r.date or not r.replacement_teacher:
            continue
        teacher_id = r.replacement_teacher_id
        teacher_entry = teacher_map.setdefault(
            teacher_id,
            {
                "name": teacher_name_with_initials(r.replacement_teacher.full_name),
                "total": 0,
                "days": {},
            },
        )
        teacher_entry["total"] += 1
        day_key = r.date.strftime('%d.%m.%Y')
        grade = extract_grade(r.lesson.class_group) if r.lesson else None
        is_primary = grade is not None and grade <= 4
        t_entry = teacher_entry["days"].setdefault(day_key, {"zs": 0, "zn": 0})
        if is_primary:
            t_entry["zn"] += 1
        else:
            t_entry["zs"] += 1

    for s in special_reps:
        if not s.date or not s.replacement_teacher:
            continue
        teacher_id = s.replacement_teacher_id
        teacher_entry = teacher_map.setdefault(
            teacher_id,
            {
                "name": teacher_name_with_initials(s.replacement_teacher.full_name),
                "total": 0,
                "days": {},
            },
        )
        teacher_entry["total"] += 1
        day_key = s.date.strftime('%d.%m.%Y')
        grade = extract_grade(s.class_group) if s.class_group else None
        is_primary = grade is not None and grade <= 4
        t_entry = teacher_entry["days"].setdefault(day_key, {"zs": 0, "zn": 0})
        if is_primary:
            t_entry["zn"] += 1
        else:
            t_entry["zs"] += 1

    doc = Document()
    teachers_sorted = sorted(teacher_map.values(), key=lambda x: x["name"].lower())
    for teacher_data in teachers_sorted:
        teacher_name = teacher_data["name"]
        total = teacher_data["total"]
        p_teacher = doc.add_paragraph()
        p_teacher.add_run(f"{teacher_name} ({total} {replacement_word(total)}):").bold = True

        has_days = False
        day_map = teacher_data["days"]
        for day_key in sorted(day_map.keys(), key=lambda d: datetime.strptime(d, "%d.%m.%Y")):
            zs = day_map[day_key]["zs"]
            zn = day_map[day_key]["zn"]
            parts = []
            if zs > 0:
                parts.append(f"{zs} (ЗС)")
            if zn > 0:
                parts.append(f"{zn} (ЗН)")
            if not parts:
                continue
            has_days = True
            p_day = doc.add_paragraph()
            p_day.add_run(f"Дата: {day_key}").bold = True
            p_day.add_run(f" - {', '.join(parts)}")

        if not has_days:
            doc.add_paragraph("—")
        doc.add_paragraph("")

    if not teacher_map:
        doc.add_paragraph("—")

    if mode_norm == "day":
        safe_ascii = f"replacement_summary_teachers_day_{start.strftime('%Y_%m_%d')}.docx"
        filename = f"Сводка по педагогам за {start.strftime('%d.%m.%Y')}.docx"
    elif mode_norm == "period":
        safe_ascii = f"replacement_summary_teachers_period_{start.strftime('%Y_%m_%d')}_{end.strftime('%Y_%m_%d')}.docx"
        filename = f"Сводка по педагогам за период {start.strftime('%d.%m.%Y')} - {end.strftime('%d.%m.%Y')}.docx"
    else:
        safe_ascii = f"replacement_summary_teachers_{start.strftime('%Y_%m')}.docx"
        filename = f"Сводка_по_педагогам_{start.strftime('%m.%Y')}.docx"

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = (
        f'attachment; filename="{safe_ascii}"; '
        f"filename*=UTF-8''{quote(filename)}"
    )
    doc.save(response)
    log_activity(
        request,
        "export_replacement_teacher_summary_docx",
        {"mode": mode_norm, "range": range_label, "from": str(start), "to": str(end), "teachers": len(teacher_map)},
    )
    return response


@login_required
def class_schedule_view(request):
    if request.method == "POST":
        guest_forbidden = _deny_guest_write_json(request)
        if guest_forbidden:
            return guest_forbidden

                                                                                 
    if not (request.user.is_superuser or getattr(request.user, 'can_calls', False)):
        return HttpResponse("Forbidden", status=403)

    if request.method == "POST":
        try:
            data = json.loads(request.body)
            shift = int(data.get("shift", 1))
            entries = data.get("schedule", []) or []

            if _use_gsheets_backend():
                rows = _gs_store().get_table_dicts("replacements_class_schedule")
                for entry in entries:
                    class_group = entry.get("class")
                    lesson_number = as_int(entry.get("lesson"))
                    start_time = entry.get("start")
                    end_time = entry.get("end")
                    if class_group and lesson_number is not None and start_time and end_time:
                        idx = None
                        for i, r in enumerate(rows):
                            if (
                                str(r.get("class_group") or "") == str(class_group)
                                and as_int(r.get("lesson_number")) == lesson_number
                                and as_int(r.get("shift")) == int(shift)
                            ):
                                idx = i
                                break
                        if idx is None:
                            row = {"id": _gs_next_id(rows, "id")}
                        else:
                            row = rows[idx]
                        row.update({
                            "class_group": str(class_group),
                            "lesson_number": lesson_number,
                            "shift": int(shift),
                            "start_time": _gs_time_str(_gs_parse_time(start_time) or start_time),
                            "end_time": _gs_time_str(_gs_parse_time(end_time) or end_time),
                        })
                        if idx is None:
                            rows.append(row)
                        else:
                            rows[idx] = row
                _gs_store().replace_table_dicts("replacements_class_schedule", rows)
                log_activity(request, "class_schedule_update", {
                    "shift": shift,
                    "entries": len(entries),
                    "backend": "gsheets",
                })
                return JsonResponse({"status": "success"})

            for entry in entries:
                class_group = entry.get("class")
                lesson_number = entry.get("lesson")
                start_time = entry.get("start")
                end_time = entry.get("end")

                if class_group and start_time and end_time:
                    ClassSchedule.objects.update_or_create(
                        class_group=class_group,
                        lesson_number=lesson_number,
                        shift=shift,
                        defaults={"start_time": start_time, "end_time": end_time}
                    )

            log_activity(request, "class_schedule_update", {
                "shift": shift,
                "entries": len(entries),
            })
            return JsonResponse({"status": "success"})

        except Exception:
            return _api_error("Не удалось сохранить расписание звонков", status=500, code="internal_error")

    else:
        if _use_gsheets_backend():
            schedule = sorted(
                _gs_store().get_table_dicts("replacements_class_schedule"),
                key=lambda r: (
                    str(r.get("class_group") or ""),
                    as_int(r.get("lesson_number"), 0) or 0,
                    as_int(r.get("shift"), 0) or 0,
                ),
            )
        else:
            schedule = ClassSchedule.objects.all().order_by("class_group", "lesson_number", "shift")

                                                                      
        grouped = defaultdict(list)

                                     
        class_map = defaultdict(dict)                                           
        shifts = {}

        for entry in schedule:
            if _use_gsheets_backend():
                cls = str(entry.get("class_group") or "")
                lnum = as_int(entry.get("lesson_number"))
                st = _gs_parse_time(entry.get("start_time"))
                en = _gs_parse_time(entry.get("end_time"))
                sh = as_int(entry.get("shift"))
            else:
                cls = entry.class_group
                lnum = entry.lesson_number
                st = entry.start_time
                en = entry.end_time
                sh = entry.shift
            if not cls or lnum is None:
                continue
            class_map[cls][lnum] = (st, en)
            shifts[cls] = sh

                                                                 
        for class_group, times in class_map.items():
            key = (shifts[class_group], tuple(sorted(times.items())))
            grouped[key].append(class_group)

        class_list = [str(i) for i in range(1, 12)]
        lesson_numbers = list(range(9))

        def class_sort_key(cls: str):
            grade = extract_grade(cls)
            suffix = re.sub(r"^\\d+", "", cls or "").strip().lower()
            return (grade if grade is not None else 99, suffix, cls or "")

                                                                        
        sorted_grouped = []
        for key, classes in grouped.items():
            classes_sorted = sorted(classes, key=class_sort_key)
            sorted_grouped.append((key, classes_sorted))

        def group_sort_key(item):
            _key, classes = item
            return (class_sort_key(classes[0]) if classes else (99, "", ""), _key[0])

        sorted_grouped.sort(key=group_sort_key)

                                                             
        primary_groups = []
        senior_groups = []
        for item in sorted_grouped:
            _key, classes = item
            first_cls = classes[0] if classes else ""
            grade = extract_grade(first_cls)
            if grade is not None and grade <= 4:
                primary_groups.append(item)
            else:
                senior_groups.append(item)

        return render(request, "class_schedule.html", {
            "grouped_schedule_primary": primary_groups,
            "grouped_schedule_senior": senior_groups,
            "class_list": class_list,
            "lesson_numbers": lesson_numbers,
        })


from django.shortcuts import render
from .models import ClassSchedule
from .pars import parse_schedule_file, save_lessons_to_db

@login_required
def upload(request):
    """Upload and (re)build the *active* schedule from a saved HTML file.

    Safe by design:
    - Does NOT touch ClassSchedule (bells)
    - Does NOT delete Lesson rows (keeps Replacement history intact)
    - Simply deactivates previous active lessons and creates/reactivates new ones
    """
                                            
    if not (request.user.is_superuser or getattr(request.user, 'can_upload', False)):
        return HttpResponse("Forbidden", status=403)

    if _use_gsheets_backend():
        active_count = sum(1 for l in _gs_store().get_table_dicts("replacements_lesson") if as_bool(l.get("is_active")))
    else:
        active_count = Lesson.objects.filter(is_active=True).count()

    if request.method == 'POST':
        uploaded_file = request.FILES.get('schedule_file')
        if not uploaded_file:
            messages.error(
                request,
                'Файл не получен. Проверьте, что выбран HTML и форма отправляется с enctype="multipart/form-data".'
            )
            return render(request, 'upload.html', {'active_count': active_count})
        if uploaded_file.size > int(getattr(settings, "MAX_SCHEDULE_UPLOAD_SIZE", 10 * 1024 * 1024)):
            messages.error(request, "Файл слишком большой.")
            return render(request, 'upload.html', {'active_count': active_count})
        filename = str(getattr(uploaded_file, "name", "") or "").lower()
        if not (filename.endswith(".html") or filename.endswith(".htm")):
            messages.error(request, "Поддерживаются только HTML/HTM файлы.")
            return render(request, 'upload.html', {'active_count': active_count})

        raw = uploaded_file.read()

                                                                         
        html_content = None
        for enc in ("utf-8-sig", "utf-8", "cp1251", "windows-1251"):
            try:
                html_content = raw.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        if html_content is None:
            html_content = raw.decode("utf-8", errors="ignore")

        if _use_gsheets_backend():
            shifts = []
            for r in _gs_store().get_table_dicts("replacements_class_schedule"):
                shifts.append({
                    "shift": as_int(r.get("shift")),
                    "start_time": _gs_parse_time(r.get("start_time")),
                    "end_time": _gs_parse_time(r.get("end_time")),
                    "lesson_number": as_int(r.get("lesson_number")),
                    "class_group": str(r.get("class_group") or ""),
                })
        else:
            shifts = list(ClassSchedule.objects.values(
                'shift', 'start_time', 'end_time', 'lesson_number', 'class_group'
            ))

        parsed_data = parse_schedule_file(html_content, shifts)

        if not parsed_data:
            messages.error(
                request,
                'Парсер не нашёл уроков в файле. Обычно это происходит, если сохранён HTML другого типа (не "Расписание школы"), '
                'или структура страницы отличается, или не заполнено "Расписание звонков" (меню: Расписание уроков по классам).'
            )
            log_activity(request, "schedule_upload_failed", {"reason": "no_lessons_found", "filename": getattr(uploaded_file, "name", "")})
            return render(request, 'upload.html', {'active_count': active_count})

        try:
            class_count = len({str(r.get('class_group')) for r in parsed_data if r.get('class_group')})
        except Exception:
            class_count = None

        if _use_gsheets_backend():
            deactivated, _saved = _gs_save_lessons_from_parsed(parsed_data)
        else:
            with transaction.atomic():
                deactivated = Lesson.objects.filter(is_active=True).update(is_active=False)
                save_lessons_to_db(parsed_data)

        messages.success(request, f'Готово! Сохранено записей: {len(parsed_data)}')
        log_activity(request, "schedule_upload", {
            "filename": getattr(uploaded_file, "name", ""),
            "saved": len(parsed_data),
            "class_count": class_count,
            "deactivated_active_lessons": deactivated,
        })

        return render(request, 'result.html', {
            'result': parsed_data,
            'class_count': class_count,
        })

    return render(request, 'upload.html', {'active_count': active_count})


@login_required
def upload_schedule_view(request):
    """View active schedule grouped by classes with inline teacher reassignment."""
    if not (request.user.is_superuser or getattr(request.user, 'can_upload', False)):
        return HttpResponse("Forbidden", status=403)
    return render(request, "upload_schedule.html", {})


@login_required
@require_GET
def upload_schedule_api(request):
    """Return active schedule grouped by class for the Upload page."""
    if not (request.user.is_superuser or getattr(request.user, 'can_upload', False)):
        return HttpResponse("Forbidden", status=403)

    day_order = {"пн": 0, "вт": 1, "ср": 2, "чт": 3, "пт": 4, "сб": 5, "вс": 6}

    if _use_gsheets_backend():
        lessons = _gs_active_lessons_rows()
        teacher_map = _gs_teacher_map()
        subject_map = _gs_subject_map()

        def class_sort_key(cls):
            grade = extract_grade(cls)
            suffix = re.sub(r"^\\d+", "", cls or "")
            return (grade if grade is not None else 99, suffix, cls or "")

        grouped = defaultdict(list)
        for l in lessons:
            grouped[str(l.get("class_group") or "")].append(l)

        classes_payload = []
        for class_group in sorted(grouped.keys(), key=class_sort_key):
            items = sorted(
                grouped[class_group],
                key=lambda x: (day_order.get(str(x.get("day_of_week") or ""), 99), as_int(x.get("lesson_number"), 0) or 0),
            )
            lessons_payload = []
            for l in items:
                tid = as_int(l.get("teacher_id"))
                sid = as_int(l.get("subject_id_subject"))
                lessons_payload.append({
                    "id": as_int(l.get("lesson_id")),
                    "day": str(l.get("day_of_week") or ""),
                    "lesson_number": as_int(l.get("lesson_number")),
                    "start": _gs_time_str(_gs_parse_time(l.get("start_time"))),
                    "end": _gs_time_str(_gs_parse_time(l.get("end_time"))),
                    "class_group": str(l.get("class_group") or ""),
                    "classroom": str(l.get("classroom") or ""),
                    "subject": subject_map.get(sid, "") if sid is not None else "",
                    "teacher_id": tid,
                    "teacher_name": teacher_map.get(tid, "") if tid is not None else "",
                })
            classes_payload.append({"class_group": class_group, "lessons": lessons_payload})
        return JsonResponse({"classes": classes_payload})

    lessons = _active_lessons().select_related("teacher", "subject").all()

    def class_sort_key(cls):
        grade = extract_grade(cls)
        suffix = re.sub(r"^\\d+", "", cls or "")
        return (grade if grade is not None else 99, suffix, cls or "")

    grouped = defaultdict(list)
    for l in lessons:
        grouped[l.class_group].append(l)

    classes_payload = []
    for class_group in sorted(grouped.keys(), key=class_sort_key):
        items = sorted(
            grouped[class_group],
            key=lambda x: (day_order.get(x.day_of_week, 99), x.lesson_number or 0),
        )
        lessons_payload = []
        for l in items:
            lessons_payload.append({
                "id": l.id,
                "day": l.day_of_week,
                "lesson_number": l.lesson_number,
                "start": l.start_time.strftime("%H:%M") if l.start_time else "",
                "end": l.end_time.strftime("%H:%M") if l.end_time else "",
                "class_group": l.class_group,
                "classroom": l.classroom or "",
                "subject": l.subject.name if l.subject else "",
                "teacher_id": l.teacher.id if l.teacher else None,
                "teacher_name": l.teacher.full_name if l.teacher else "",
            })
        classes_payload.append({
            "class_group": class_group,
            "lessons": lessons_payload,
        })

    return JsonResponse({"classes": classes_payload})


@login_required
@require_GET
def special_replacement_options(request):
    """Return distinct class groups and subjects from active lessons."""
    if not (request.user.is_superuser or _is_guest_user(request.user) or getattr(request.user, 'can_calendar', False)):
        return HttpResponse("Forbidden", status=403)

    if _use_gsheets_backend():
        active_lessons = _gs_active_lessons_rows()
        subject_map = _gs_subject_map()
        classes = sorted({str(l.get("class_group") or "") for l in active_lessons if str(l.get("class_group") or "").strip()})
        subjects = sorted({
            subject_map.get(as_int(l.get("subject_id_subject")), "")
            for l in active_lessons
            if as_int(l.get("subject_id_subject")) in subject_map
        })
        subjects = [s for s in subjects if s]
        return JsonResponse({"classes": classes, "subjects": subjects})

    classes = sorted({c for c in _active_lessons().values_list("class_group", flat=True) if c})
    subjects = sorted({s for s in _active_lessons().values_list("subject__name", flat=True) if s})
    return JsonResponse({"classes": classes, "subjects": subjects})


@login_required
@require_GET
def special_replacement_lessons(request):
    """Return lessons for selected class+subject on a specific date."""
    if not (request.user.is_superuser or _is_guest_user(request.user) or getattr(request.user, 'can_calendar', False)):
        return HttpResponse("Forbidden", status=403)

    date_str = request.GET.get("date")
    class_group = (request.GET.get("class_group") or "").strip()
    subject_name = (request.GET.get("subject") or "").strip()
    if not (date_str and class_group and subject_name):
        return JsonResponse({"lessons": []})
    try:
        selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception:
        return JsonResponse({"lessons": []})

    if _use_gsheets_backend():
        day_short = day_short_from_date(selected_date)
        subject_id = _gs_subject_id_by_name(subject_name)
        if not subject_id:
            return JsonResponse({"lessons": []})
        lessons = []
        for l in _gs_active_lessons_rows():
            if str(l.get("class_group") or "") != class_group:
                continue
            if str(l.get("day_of_week") or "") != day_short:
                continue
            if as_int(l.get("subject_id_subject")) != subject_id:
                continue
            lessons.append({
                "id": as_int(l.get("lesson_id")),
                "lesson_number": as_int(l.get("lesson_number")),
                "start": _gs_time_str(_gs_parse_time(l.get("start_time"))),
                "end": _gs_time_str(_gs_parse_time(l.get("end_time"))),
                "classroom": str(l.get("classroom") or ""),
            })
        lessons.sort(key=lambda x: (x.get("lesson_number") or 0, x.get("id") or 0))
        return JsonResponse({"lessons": lessons})

    day_short = day_short_from_date(selected_date)
    lessons_qs = _active_lessons().select_related("subject").filter(
        class_group=class_group,
        day_of_week=day_short,
        subject__name=subject_name,
    ).order_by("lesson_number")

    lessons = []
    for l in lessons_qs:
        lessons.append({
            "id": l.id,
            "lesson_number": l.lesson_number,
            "start": l.start_time.strftime("%H:%M") if l.start_time else "",
            "end": l.end_time.strftime("%H:%M") if l.end_time else "",
            "classroom": l.classroom or "",
        })
    return JsonResponse({"lessons": lessons})


@login_required
@require_GET
def special_replacement_time(request):
    """Return shift and bell times for manual special replacement by class and lesson number."""
    if not (request.user.is_superuser or _is_guest_user(request.user) or getattr(request.user, 'can_calendar', False)):
        return HttpResponse("Forbidden", status=403)

    class_group = (request.GET.get("class_group") or "").strip()
    lesson_number_raw = request.GET.get("lesson_number")
    if not class_group or not lesson_number_raw:
        return JsonResponse({"shift": None, "start": "", "end": "", "found": False})

    try:
        lesson_number = int(lesson_number_raw)
    except Exception:
        return JsonResponse({"shift": None, "start": "", "end": "", "found": False})

    if _use_gsheets_backend():
        shift = _effective_shift_for_class(class_group, None)
                                                                        
        rows = _gs_store().get_table_dicts("replacements_class_schedule")
        start_t = end_t = None
        for r in rows:
            if as_int(r.get("shift")) != as_int(shift):
                continue
            if as_int(r.get("lesson_number")) != lesson_number:
                continue
            if str(r.get("class_group") or "") == class_group:
                start_t = _gs_parse_time(r.get("start_time"))
                end_t = _gs_parse_time(r.get("end_time"))
                break
        if not (start_t and end_t):
            grade = extract_grade(class_group)
            for r in rows:
                if as_int(r.get("shift")) != as_int(shift):
                    continue
                if as_int(r.get("lesson_number")) != lesson_number:
                    continue
                if extract_grade(str(r.get("class_group") or "")) == grade:
                    start_t = _gs_parse_time(r.get("start_time"))
                    end_t = _gs_parse_time(r.get("end_time"))
                    break
        found = bool(start_t and end_t)
        return JsonResponse({
            "shift": int(shift) if shift in (1, 2) else None,
            "start": _gs_time_str(start_t) if start_t else "",
            "end": _gs_time_str(end_t) if end_t else "",
            "found": found,
        })

    shift = _effective_shift_for_class(class_group, None)
    start_t, end_t = _effective_times_for_lesson(
        class_group=class_group,
        lesson_number=lesson_number,
        shift=shift,
        fallback_start=None,
        fallback_end=None,
    )
    found = bool(start_t and end_t)
    return JsonResponse({
        "shift": int(shift) if shift in (1, 2) else None,
        "start": start_t.strftime("%H:%M") if start_t else "",
        "end": end_t.strftime("%H:%M") if end_t else "",
        "found": found,
    })


@login_required
@require_POST
def update_lesson_teacher(request, lesson_id):
    """Update teacher either for one lesson or for all subjects of the same class."""
    guest_forbidden = _deny_guest_write_json(request)
    if guest_forbidden:
        return guest_forbidden

    if not (request.user.is_superuser or getattr(request.user, 'can_upload', False)):
        return HttpResponse("Forbidden", status=403)
    try:
        data = json.loads(request.body)
    except Exception:
        return JsonResponse({"error": "Некорректный JSON"}, status=400)

    teacher_id = data.get("teacher_id")
    apply_all_subjects = bool(data.get("apply_all_subjects"))
    if not teacher_id:
        return JsonResponse({"error": "teacher_id обязателен"}, status=400)

    if _use_gsheets_backend():
        try:
            teacher_id = int(teacher_id)
        except Exception:
            return JsonResponse({"error": "teacher_id должен быть числом"}, status=400)
        teacher_rows = _gs_store().get_table_dicts("replacements_teacher")
        teacher = next((t for t in teacher_rows if as_int(t.get("teacher_id")) == teacher_id), None)
        if not teacher:
            return JsonResponse({"error": "Учитель не найден"}, status=404)
        lessons_rows = _gs_store().get_table_dicts("replacements_lesson")
        lesson = next((l for l in lessons_rows if as_int(l.get("lesson_id")) == int(lesson_id)), None)
        if not lesson:
            return JsonResponse({"error": "Урок не найден"}, status=404)

        if apply_all_subjects:
            target_indices = [
                i for i, l in enumerate(lessons_rows)
                if as_bool(l.get("is_active"))
                and str(l.get("class_group") or "") == str(lesson.get("class_group") or "")
                and as_int(l.get("subject_id_subject")) == as_int(lesson.get("subject_id_subject"))
                and as_int(l.get("teacher_id")) == as_int(lesson.get("teacher_id"))
            ]
            mode = "class_same_subject_same_teacher"
        else:
            target_indices = [i for i, l in enumerate(lessons_rows) if as_int(l.get("lesson_id")) == as_int(lesson.get("lesson_id"))]
            mode = "single_lesson"

        affected_count = len(target_indices)
        old_teacher_ids = sorted({as_int(lessons_rows[i].get("teacher_id")) for i in target_indices if as_int(lessons_rows[i].get("teacher_id")) is not None})
        updated_count = 0
        for i in target_indices:
            if as_int(lessons_rows[i].get("teacher_id")) != teacher_id:
                lessons_rows[i]["teacher_id"] = teacher_id
                updated_count += 1

        _gs_store().replace_table_dicts("replacements_lesson", lessons_rows)
        log_activity(request, "lesson_teacher_update", {
            "lesson_id": lesson_id,
            "from_teacher_ids": old_teacher_ids,
            "to_teacher_id": teacher_id,
            "class_group": str(lesson.get("class_group") or ""),
            "subject_id": as_int(lesson.get("subject_id_subject")),
            "mode": mode,
            "updated_count": updated_count,
            "affected_count": affected_count,
            "backend": "gsheets",
        })
        return JsonResponse({
            "status": "success",
            "teacher_name": str(teacher.get("full_name") or ""),
            "mode": mode,
            "updated_count": updated_count,
            "affected_count": affected_count,
        })

    lesson = get_object_or_404(Lesson, id=lesson_id)
    teacher = get_object_or_404(Teacher, id=teacher_id)

    if apply_all_subjects:
        target_qs = _active_lessons().filter(
            class_group=lesson.class_group,
            subject_id=lesson.subject_id,
            teacher_id=lesson.teacher_id,
        )
        mode = "class_same_subject_same_teacher"
    else:
        target_qs = Lesson.objects.filter(id=lesson.id)
        mode = "single_lesson"

    affected_count = target_qs.count()
    old_teacher_ids = sorted(set(target_qs.values_list("teacher_id", flat=True)))
    updated_count = target_qs.exclude(teacher_id=teacher.id).update(teacher=teacher)

    log_activity(request, "lesson_teacher_update", {
        "lesson_id": lesson_id,
        "from_teacher_ids": old_teacher_ids,
        "to_teacher_id": teacher.id,
        "class_group": lesson.class_group,
        "subject_id": lesson.subject_id,
        "mode": mode,
        "updated_count": updated_count,
        "affected_count": affected_count,
    })
    return JsonResponse({
        "status": "success",
        "teacher_name": teacher.full_name,
        "mode": mode,
        "updated_count": updated_count,
        "affected_count": affected_count,
    })


@login_required
@require_POST
def reassign_teacher_lessons(request):
    """Move all lessons from one teacher to another."""
    guest_forbidden = _deny_guest_write_json(request)
    if guest_forbidden:
        return guest_forbidden

    if not (request.user.is_superuser or getattr(request.user, 'can_upload', False)):
        return HttpResponse("Forbidden", status=403)

    try:
        data = json.loads(request.body or b"{}")
    except Exception:
        return JsonResponse({"error": "Некорректный JSON"}, status=400)

    from_teacher_id = data.get("from_teacher_id")
    to_teacher_id = data.get("to_teacher_id")
    active_only = bool(data.get("active_only", True))

    if not from_teacher_id or not to_teacher_id:
        return JsonResponse({"error": "from_teacher_id и to_teacher_id обязательны"}, status=400)

    try:
        from_teacher_id = int(from_teacher_id)
        to_teacher_id = int(to_teacher_id)
    except Exception:
        return JsonResponse({"error": "ID учителей должны быть числами"}, status=400)

    if from_teacher_id == to_teacher_id:
        return JsonResponse({"error": "Выберите двух разных учителей"}, status=400)

    if _use_gsheets_backend():
        teacher_rows = _gs_store().get_table_dicts("replacements_teacher")
        from_teacher = next((t for t in teacher_rows if as_int(t.get("teacher_id")) == from_teacher_id), None)
        to_teacher = next((t for t in teacher_rows if as_int(t.get("teacher_id")) == to_teacher_id), None)
        if not from_teacher or not to_teacher:
            return JsonResponse({"error": "Учитель не найден"}, status=404)

        lessons_rows = _gs_store().get_table_dicts("replacements_lesson")
        target_indices = [
            i for i, l in enumerate(lessons_rows)
            if as_int(l.get("teacher_id")) == from_teacher_id and (not active_only or as_bool(l.get("is_active")))
        ]
        affected_count = len(target_indices)
        updated_count = 0
        for i in target_indices:
            lessons_rows[i]["teacher_id"] = to_teacher_id
            updated_count += 1
        _gs_store().replace_table_dicts("replacements_lesson", lessons_rows)
        log_activity(request, "lesson_teacher_reassign_bulk", {
            "from_teacher_id": from_teacher_id,
            "to_teacher_id": to_teacher_id,
            "from_teacher_name": str(from_teacher.get("full_name") or ""),
            "to_teacher_name": str(to_teacher.get("full_name") or ""),
            "active_only": active_only,
            "affected_count": affected_count,
            "updated_count": updated_count,
            "backend": "gsheets",
        })
        return JsonResponse({
            "status": "success",
            "from_teacher_name": str(from_teacher.get("full_name") or ""),
            "to_teacher_name": str(to_teacher.get("full_name") or ""),
            "active_only": active_only,
            "affected_count": affected_count,
            "updated_count": updated_count,
        })

    from_teacher = Teacher.objects.filter(id=from_teacher_id).first()
    to_teacher = Teacher.objects.filter(id=to_teacher_id).first()
    if not from_teacher or not to_teacher:
        return JsonResponse({"error": "Учитель не найден"}, status=404)

    lessons_qs = Lesson.objects.filter(teacher_id=from_teacher_id)
    if active_only:
        lessons_qs = lessons_qs.filter(is_active=True)

    affected_count = lessons_qs.count()
    updated_count = lessons_qs.update(teacher_id=to_teacher_id)

    log_activity(request, "lesson_teacher_reassign_bulk", {
        "from_teacher_id": from_teacher_id,
        "to_teacher_id": to_teacher_id,
        "from_teacher_name": from_teacher.full_name,
        "to_teacher_name": to_teacher.full_name,
        "active_only": active_only,
        "affected_count": affected_count,
        "updated_count": updated_count,
    })

    return JsonResponse({
        "status": "success",
        "from_teacher_name": from_teacher.full_name,
        "to_teacher_name": to_teacher.full_name,
        "active_only": active_only,
        "affected_count": affected_count,
        "updated_count": updated_count,
    })


@login_required
@require_POST
def clear_schedule(request):
    """Deactivate the currently active schedule without touching bells or replacement history."""
    if _is_guest_user(request.user):
        return HttpResponse("Forbidden", status=403)

    if not (request.user.is_superuser or getattr(request.user, 'can_upload', False)):
        return HttpResponse("Forbidden", status=403)

    if _use_gsheets_backend():
        rows = _gs_store().get_table_dicts("replacements_lesson")
        deactivated = 0
        for r in rows:
            if as_bool(r.get("is_active")):
                r["is_active"] = 0
                deactivated += 1
        _gs_store().replace_table_dicts("replacements_lesson", rows)
        messages.success(request, f'Расписание очищено (деактивировано уроков: {deactivated}).')
        log_activity(request, "schedule_clear", {"deactivated_active_lessons": deactivated, "backend": "gsheets"})
        return redirect('/upload/')
    deactivated = Lesson.objects.filter(is_active=True).update(is_active=False)
    messages.success(request, f'Расписание очищено (деактивировано уроков: {deactivated}).')
    log_activity(request, "schedule_clear", {"deactivated_active_lessons": deactivated})
    return redirect('/upload/')
